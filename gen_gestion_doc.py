"""
gen_gestion_doc.py
Genera: GESTION_REQUISITOS_PROPUESTA_VALOR_actualizado.docx
Mantiene la estructura original y agrega 3 nuevas simulaciones
de cambio (SC-002, SC-003, SC-004) para funcionalidades futuras
(v2.0 / v3.0) que NO afectan el MVP1 ya implementado.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests, zlib, io

# ══════════════════════════════════════════════════════════════════════════════
# PLANTUML helpers
# ══════════════════════════════════════════════════════════════════════════════
def _encode6bit(b):
    if b < 10: return chr(48 + b)
    b -= 10
    if b < 26: return chr(65 + b)
    b -= 26
    if b < 26: return chr(97 + b)
    b -= 26
    return '-' if b == 0 else '_'

def _append3bytes(b1, b2, b3):
    return (
        _encode6bit(b1 >> 2) +
        _encode6bit(((b1 & 0x3) << 4) | (b2 >> 4)) +
        _encode6bit(((b2 & 0xF) << 2) | (b3 >> 6)) +
        _encode6bit(b3 & 0x3F)
    )

def plantuml_encode(text):
    data = zlib.compress(text.encode('utf-8'))[2:-4]
    res = ""
    for i in range(0, len(data), 3):
        if i + 2 < len(data): res += _append3bytes(data[i], data[i+1], data[i+2])
        elif i + 1 < len(data): res += _append3bytes(data[i], data[i+1], 0)
        else: res += _append3bytes(data[i], 0, 0)
    return res

def fetch_diagram(uml_code, width_inches=5.8):
    try:
        enc = plantuml_encode(uml_code)
        url = f"http://www.plantuml.com/plantuml/png/{enc}"
        r = requests.get(url, timeout=25)
        if r.status_code == 200 and r.content[:4] == b'\x89PNG':
            return io.BytesIO(r.content), width_inches
    except Exception:
        pass
    return None

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, cols, bg='1E3A5F', fg='FFFFFF'):
    for cell, txt in zip(table.rows[0].cells, cols):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        set_cell_bg(cell, bg)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_diagram(doc, uml_code, caption, width=5.8):
    result = fetch_diagram(uml_code, width)
    if result:
        img_stream, w = result
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(w))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(f'[Diagrama: {caption}]').italic = True
    doc.add_paragraph()

def two_col_table(doc, rows_data, key_bg='D0D8E4'):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    for i, (key, val) in enumerate(rows_data):
        ck = tbl.rows[i].cells[0]
        cv = tbl.rows[i].cells[1]
        ck.text = key
        ck.paragraphs[0].runs[0].bold = True
        set_cell_bg(ck, key_bg)
        cv.text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMAS UML
# ══════════════════════════════════════════════════════════════════════════════

# Proceso general de gestión de cambios (Actividad)
UML_PROCESO_CAMBIOS = """@startuml
skinparam activity {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}
skinparam roundcorner 15

title Proceso Formal de Gestion de Cambios\\nOLYMPUS CORE

start
:Stakeholder identifica necesidad de cambio;
:Registrar Solicitud de Cambio (SC)\\ncon ID unico, motivo, solicitante e impacto esperado;

:Analisis de impacto tecnico y de negocio;
note right
  Impacto en: RF, RNF, RD, RQ-CAL
  Modulos afectados
  Costo estimado (dias-persona)
  Version objetivo (2.0 o 3.0)
end note

if (Impacto en MVP1 congelado?) then (si)
  :Documentar como riesgo;
  :Reclasificar a version futura;
  stop
else (no)
endif

:Revision por CCB\\n(4 integrantes del equipo);

if (CCB aprueba?) then (no)
  :Rechazar con justificacion documentada;
  :Actualizar bitacora de cambios;
  stop
else (si)
endif

:Planificar implementacion\\n(actividades, responsables, duracion);
:Implementar en rama feature/ del repositorio;
:Pruebas de integracion y regresion;

if (Pruebas exitosas?) then (no)
  :Revertir cambio;
  :Notificar al CCB;
  stop
else (si)
endif

:Despliegue a produccion;
:Actualizar: Matriz de trazabilidad,\\nDocumentacion y Bitacora de cambios;
stop
@enduml"""

# SC-002: Integración Pasarela de Pagos — Secuencia
UML_SEC_PAGOS = """@startuml
skinparam sequenceArrowThickness 2
skinparam roundcorner 10

title SC-002: Flujo de Pago Online con Pasarela Externa\\n(RF-20, RF-21 — Version 2.0)

actor "Socio" as SOC
participant "Frontend\\nReact" as UI
participant "API Backend\\nNode/Express" as API
participant "Pasarela de Pagos\\n(Culqi / MercadoPago)" as PAY
database "PostgreSQL" as DB

SOC -> UI: Seleccionar plan y hacer clic en 'Pagar Online'
UI -> API: POST /pagos/iniciar\\n{socio_id, plan_id, monto}

API -> DB: Verificar socio y plan activo
DB --> API: OK

API -> PAY: Crear cargo\\n{amount, currency, source_id, email}
PAY --> API: {charge_id, status: 'paid'}

alt Pago exitoso
  API -> DB: INSERT pagos (monto, metodo='online', estado='completado')
  API -> DB: INSERT/UPDATE membresias (estado='activa', fecha_fin)
  API --> UI: 200 OK {nueva_fecha_fin, mensaje_exito}
  UI --> SOC: Confirmacion de pago y nueva fecha de vencimiento
  note right
    RF-20: Pago online con tarjeta
    RF-21: Actualizacion automatica
    de membresia
  end note
else Pago fallido
  API --> UI: 402 {error: 'PAGO_RECHAZADO', motivo}
  UI --> SOC: Mostrar error y sugerir metodo alternativo
end
@enduml"""

# SC-003: App Móvil — Componentes
UML_COMP_APP_MOVIL = """@startuml
skinparam componentStyle uml2
skinparam component {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}

title SC-003: Arquitectura App Movil — Version 3.0\\n(RF-27, RF-28, RF-29)

node "Dispositivo Movil\\n(iOS / Android)" {
  component "App Movil\\nReact Native" as APP {
    [Pantalla QR del Socio]
    [Mis Reservas]
    [Mi Progreso]
    [Notificaciones Push]
  }
  note right of APP
    RF-27: Ver QR desde movil
    RF-28: Gestionar reservas
    RF-29: Ver progreso fisico
  end note
}

node "Vercel (CDN)" {
  [API REST existente\\nNode.js/Express] as API
  note right of API
    La app movil reutiliza
    los mismos endpoints
    del MVP1 (sin cambios)
  end note
}

node "Neon (DBaaS)" {
  database "PostgreSQL" as DB
}

node "Firebase / OneSignal" {
  [Servicio Push\\nNotificaciones] as PUSH
  note right of PUSH
    RF-22: Notificaciones
    de vencimiento QR
  end note
}

APP <--> API : HTTPS REST API (JWT)
API <--> DB : TCP/SSL
API --> PUSH : Triggear notificacion
PUSH --> APP : Push notification
@enduml"""

# SC-004: Módulo Nutricional — Clases
UML_CLASE_NUTRICION = """@startuml
hide circle
skinparam classAttributeIconSize 0
skinparam class {
  BackgroundColor #F0F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}

title SC-004: Modelo de Datos — Modulo de Planes Nutricionales\\n(RF-30, RF-31, RF-32 — Version 3.0)

class Socios {
  + id : UUID
  + nombre : String
  + estado : Enum
  -- ya existe en MVP1 --
}

class FichasMedicas {
  + id : UUID
  + objetivo : String
  + restricciones_alimentarias : String [NEW]
  + alergias : String [NEW]
  -- RF-30: extension de ficha medica --
}

class PlanesNutricionales {
  + id : UUID
  + nombre : String
  + calorias_diarias : Int
  + descripcion : Text
  + objetivo : Enum {bajar_peso,ganar_masa,mantenimiento}
  + creado_por : UUID (Entrenador/Nutricionista)
  + fecha_creacion : Date
  -- RF-31: catalogo de planes nutricionales --
}

class AsignacionNutricional {
  + id : UUID
  + socio_id : UUID
  + plan_id : UUID
  + fecha_inicio : Date
  + fecha_fin : Date
  + observaciones : Text
  -- RF-32: asignacion a socio --
}

class RegistroComidas {
  + id : UUID
  + asignacion_id : UUID
  + fecha : Date
  + descripcion : String
  + calorias_consumidas : Int
  -- RF-33: seguimiento diario --
}

Socios "1" -- "1" FichasMedicas : complementa >
Socios "1" -- "*" AsignacionNutricional : recibe >
PlanesNutricionales "1" -- "*" AsignacionNutricional : < define
AsignacionNutricional "1" -- "*" RegistroComidas : tiene >

note bottom of PlanesNutricionales
  RF-30: No requiere modificar
  tablas existentes del MVP1.
  Se agregan tablas nuevas.
end note
@enduml"""

# Estado del ciclo de vida de una SC
UML_ESTADO_SC = """@startuml
skinparam state {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}
skinparam roundcorner 15

title Ciclo de Vida de una Solicitud de Cambio (SC)\\nOLYMPUS CORE — Proceso CCB

[*] --> Borrador : Stakeholder registra SC

Borrador : SC redactada, sin revisar
Borrador --> EnRevision : Analista recibe y evalua

EnRevision : Analisis de impacto tecnico\\ny de negocio en curso
EnRevision --> Aprobada : CCB aprueba (4/4 votos)
EnRevision --> Rechazada : CCB rechaza

Aprobada : Asignada a version objetivo
Aprobada --> EnImplementacion : Sprint planificado

Rechazada : Justificacion documentada
Rechazada --> [*]

EnImplementacion : Desarrollo en rama feature/
EnImplementacion --> EnPruebas : PR aprobado y merge

EnPruebas : QA: pruebas funcionales\\ny de regresion
EnPruebas --> Cerrada : Pruebas exitosas, desplegado
EnPruebas --> EnImplementacion : Pruebas fallidas\\n(revertir)

Cerrada : Documentacion actualizada
Cerrada --> [*]

note right of Aprobada
  SC-002: v2.0
  SC-003: v3.0
  SC-004: v3.0
end note
@enduml"""

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCIÓN DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── PORTADA ────────────────────────────────────────────────────────────────────
for text, size, bold, color in [
    ('UNIVERSIDAD NACIONAL MAYOR DE SAN MARCOS', 14, True,  (0x1E,0x3A,0x5F)),
    ('FACULTAD DE SISTEMAS E INFORMÁTICA',       13, False, (0x1E,0x3A,0x5F)),
    ('',11,False,(0,0,0)),
    ('ASIGNATURA', 11, False, (0,0,0)),
    ('Ingeniería de Requisitos', 13, True, (0x1E,0x3A,0x5F)),
    ('',11,False,(0,0,0)),
    ('DOCENTE', 11, False, (0,0,0)),
    ('Rodríguez Rodríguez, Ciro', 12, False, (0,0,0)),
    ('',11,False,(0,0,0)),
    ('INTEGRANTES GRUPO 5:', 11, True, (0,0,0)),
    ('Alva Chacon, Jose Benjamin          24200045', 11, False, (0,0,0)),
    ('Cordova Guerra, Josue Rodrigo       24200155', 11, False, (0,0,0)),
    ('Sandoval Dominguez, Erick Marco     24200172', 11, False, (0,0,0)),
    ('Melendez Bustamante, Alvaro Mathias 24200166', 11, False, (0,0,0)),
    ('',11,False,(0,0,0)),
    ('Proyecto: OLYMPUS CORE — Sistema Gestor de Gimnasio', 11, False, (0,0,0)),
    ('2026', 12, False, (0,0,0)),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)

doc.add_page_break()

# ── ÍNDICE ─────────────────────────────────────────────────────────────────────
doc.add_heading('Índice', 1)
indice = [
    '4. Gestión de Requisitos',
    '    4.1. Línea base: Declaración de los requisitos congelados',
    '    4.2. Matriz de trazabilidad (actualizada)',
    '    4.3. Gestión de cambios (Simulaciones)',
    '        4.3.1. Proceso de gestión de cambios adoptado',
    '        4.3.2. SC-001 — Notificaciones por WhatsApp (v2.0)',
    '        4.3.3. SC-002 — Integración con Pasarela de Pagos Online (v2.0)',
    '        4.3.4. SC-003 — Aplicación Móvil para Socios (v3.0)',
    '        4.3.5. SC-004 — Módulo de Planes Nutricionales (v3.0)',
    '    4.4. Resumen de la Bitácora de Cambios',
    '5. Propuesta de Valor del Proyecto',
]
for line in indice:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5 if line.startswith('    ') else 0)
    p.add_run(line).font.size = Pt(11)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 4: Gestión de Requisitos
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Gestión de Requisitos', 1)

# 4.1 Línea base
doc.add_heading('4.1. Línea Base: Declaración de los Requisitos Congelados', 2)
doc.add_paragraph(
    'La línea base del MVP1 de OLYMPUS CORE está conformada por los siguientes módulos aprobados '
    'e implementados. Cualquier modificación a estos módulos requiere evaluación del CCB y '
    'justificación de impacto. Las simulaciones de cambio de este documento NO afectan esta '
    'línea base — están diseñadas para versiones 2.0 y 3.0 del sistema:'
)
lb_t = doc.add_table(rows=1, cols=3)
lb_t.style = 'Table Grid'
header_row(lb_t, ['Módulo MVP1 (Congelado)', 'RF asociados', 'Estado'])
lb_data = [
    ('Autenticación y Control de Acceso (RBAC)', 'RF-01, RF-10, RF-11',       'CONGELADO ✓'),
    ('Gestión de Socios',                        'RF-02, RF-08, RF-15',       'CONGELADO ✓'),
    ('Gestión de Membresías y Planes',           'RF-03, RF-04, RF-05',       'CONGELADO ✓'),
    ('Control de Acceso QR',                     'RF-06, RF-07',              'CONGELADO ✓'),
    ('Dashboard Administrativo',                 'RF-07',                     'CONGELADO ✓'),
    ('Reserva y Control de Clases Grupales',     'RF-09, RF-16, RF-17, RF-18','CONGELADO ✓'),
    ('Sistema de Penalizaciones (Strikes)',       'RF-14, RF-17',              'CONGELADO ✓'),
    ('Gestión de Rutinas y Ficha Médica',        'RF-12, RF-13',              'CONGELADO ✓'),
    ('Historial de Pagos / Perfil del Socio',    'RF-19',                     'CONGELADO ✓'),
]
for row in lb_data:
    r = lb_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
    set_cell_bg(r.cells[2], 'D4EDDA')
doc.add_paragraph()
doc.add_paragraph(
    'Nota: Cualquier cambio solicitado que afecte un módulo de la tabla anterior '
    'será rechazado en esta fase o diferido a una versión posterior con análisis de '
    'impacto completo.'
)

# 4.2 Matriz de trazabilidad actualizada
doc.add_heading('4.2. Matriz de Trazabilidad (Actualizada — 34 RF)', 2)
doc.add_paragraph(
    'La siguiente matriz incluye los 34 RF del proyecto, diferenciando los implementados en '
    'el MVP1 (Must Have / Should Have) de los planificados para versiones futuras '
    '(Could Have / Won\'t Have):'
)
mt = doc.add_table(rows=1, cols=5)
mt.style = 'Table Grid'
header_row(mt, ['ID', 'Requisito', 'Caso de Uso', 'Módulo', 'Versión'])

rf_data = [
    # MVP1 — Must Have
    ('RF-01',  'Inicio de sesión con RBAC',               'CU-01: Iniciar sesión',            'Seguridad',        'MVP1'),
    ('RF-02',  'Registrar socio',                         'CU-02: Registrar socio',           'Recepción',        'MVP1'),
    ('RF-03',  'Gestionar planes de membresía',           'CU-03: Gestionar planes',          'Membresías',       'MVP1'),
    ('RF-04',  'Visualizar estado de membresía',          'CU-04: Ver membresía',             'Membresías',       'MVP1'),
    ('RF-05',  'Registrar pago de membresía',             'CU-05: Registrar pago',            'Recepción',        'MVP1'),
    ('RF-06',  'Validar acceso por código QR',            'CU-06: Validar QR',                'Acceso',           'MVP1'),
    ('RF-07',  'Dashboard administrativo',                'CU-07: Ver dashboard',             'Administración',   'MVP1'),
    ('RF-08',  'Listar socios con filtros',               'CU-08: Listar socios',             'Recepción',        'MVP1'),
    ('RF-09',  'Reservar clase grupal',                   'CU-09: Reservar clase',            'Reservas',         'MVP1'),
    ('RF-10',  'Visualizar perfil y código QR',           'CU-10: Ver perfil',                'Socio',            'MVP1'),
    ('RF-11',  'Gestión de roles y usuarios',             'CU-11: Gestionar usuarios',        'Seguridad',        'MVP1'),
    # MVP1 — Should Have
    ('RF-12',  'Gestión de rutinas de entrenamiento',     'CU-12: Gestionar rutinas',         'Entrenadores',     'MVP1'),
    ('RF-13',  'Registro de progreso físico',             'CU-13: Registrar progreso',        'Entrenadores',     'MVP1'),
    ('RF-14',  'Penalización automática por strike',      'CU-14: Penalizar inasistencia',    'Reservas',         'MVP1'),
    ('RF-15',  'Edición de información del socio',        'CU-15: Editar socio',              'Recepción',        'MVP1'),
    ('RF-16',  'Gestión de clases grupales',              'CU-16: Gestionar clases',          'Entrenadores',     'MVP1'),
    ('RF-17',  'Justificación de strike por Admin',       'CU-17: Justificar strike',         'Administración',   'MVP1'),
    ('RF-18',  'Control de aforo en tiempo real',         'CU-18: Controlar aforo',           'Reservas',         'MVP1'),
    ('RF-19',  'Historial de pagos del socio',            'CU-19: Ver historial',             'Membresías',       'MVP1'),
    # v2.0 — Could Have
    ('RF-20',  'Pago online con tarjeta (SC-002)',        'CU-20: Pagar online',              'Membresías',       'v2.0'),
    ('RF-21',  'Actualización automática post-pago',      'CU-21: Activar membresía auto',    'Membresías',       'v2.0'),
    ('RF-22',  'Notificaciones por WhatsApp (SC-001)',    'CU-22: Notificar vencimiento',     'Notificaciones',   'v2.0'),
    ('RF-23',  'Recordatorio por correo electrónico',     'CU-23: Enviar correo',             'Notificaciones',   'v2.0'),
    ('RF-24',  'Módulo de promociones y convenios',       'CU-24: Gestionar promociones',     'Administración',   'v2.0'),
    ('RF-25',  'Reportes de ingresos exportables (PDF)',  'CU-25: Exportar reporte',          'Administración',   'v2.0'),
    ('RF-26',  'Búsqueda avanzada de socios',             'CU-26: Buscar socio avanzado',     'Recepción',        'v2.0'),
    # v3.0 — Could Have / Won't Have
    ('RF-27',  'App móvil — Ver QR (SC-003)',             'CU-27: Ver QR en móvil',           'App Móvil',        'v3.0'),
    ('RF-28',  'App móvil — Gestionar reservas',          'CU-28: Reservar desde móvil',      'App Móvil',        'v3.0'),
    ('RF-29',  'App móvil — Ver progreso físico',         'CU-29: Ver progreso en móvil',     'App Móvil',        'v3.0'),
    ('RF-30',  'Ficha médica ampliada con alergias',      'CU-30: Editar ficha médica',       'Entrenadores',     'v3.0'),
    ('RF-31',  'Catálogo de planes nutricionales (SC-004)','CU-31: Gestionar nutrición',     'Nutrición',        'v3.0'),
    ('RF-32',  'Asignación de plan nutricional',          'CU-32: Asignar dieta',             'Nutrición',        'v3.0'),
    ('RF-33',  'Registro de comidas diarias',             'CU-33: Registrar dieta diaria',    'Nutrición',        'v3.0'),
    ('RF-34',  'Integración con wearables/smartwatch',   'CU-34: Sincronizar dispositivo',   'Integraciones',    'v3.0'),
]
VERSION_BG = {'MVP1': 'D4EDDA', 'v2.0': 'FFF3CD', 'v3.0': 'D0D8E4'}
for row in rf_data:
    r = mt.add_row()
    for c, t in zip(r.cells, row): c.text = t
    bg = VERSION_BG.get(row[4], 'FFFFFF')
    set_cell_bg(r.cells[4], bg)
doc.add_paragraph()
doc.add_paragraph('Leyenda: 🟢 MVP1 (congelado) | 🟡 v2.0 (planificado) | 🔵 v3.0 (futuro)')

# 4.3 Gestión de Cambios
doc.add_heading('4.3. Gestión de Cambios (Simulaciones)', 2)
doc.add_paragraph(
    'La gestión de cambios en OLYMPUS CORE se maneja siguiendo un proceso formal estructurado '
    'en seis etapas: identificación, evaluación, aprobación, implementación, '
    'verificación/validación y documentación. Esto asegura que cualquier modificación a los '
    'requisitos congelados en la línea base (sección 4.1) sea controlada, trazable y no '
    'comprometa el alcance del MVP1.\n\n'
    'A continuación se presentan cuatro simulaciones de solicitudes de cambio, de las cuales '
    'la SC-001 fue definida originalmente y las SC-002, SC-003 y SC-004 son nuevas propuestas '
    'generadas en esta actualización.'
)

# 4.3.1 Proceso formal
doc.add_heading('4.3.1. Proceso de Gestión de Cambios Adoptado', 3)
proc_t = doc.add_table(rows=1, cols=2)
proc_t.style = 'Table Grid'
header_row(proc_t, ['Etapa', 'Descripción'])
etapas = [
    ('1. Identificación',          'Registro formal de la solicitud de cambio (formulario SC), incluyendo motivo, solicitante e impacto esperado. Se asigna ID único (SC-XXX).'),
    ('2. Evaluación',              'Análisis de impacto técnico y de negocio; consulta con miembros del equipo. Se verifica si el cambio afecta el MVP1 congelado.'),
    ('3. Aprobación',              'Revisión por el Comité de Control de Cambios (CCB) del grupo, integrado por los 4 integrantes. Se requiere consenso para aprobar.'),
    ('4. Implementación',          'Planificación y ejecución del cambio en rama feature/ del repositorio Git, actualizando el módulo afectado sin romper el MVP1.'),
    ('5. Verificación y validación','Pruebas funcionales y de regresión para confirmar que el cambio cumple lo esperado sin afectar otros módulos del sistema.'),
    ('6. Documentación',           'Actualización de la matriz de trazabilidad, documentación técnica y bitácora de cambios. Cierre formal de la SC.'),
]
for e in etapas:
    r = proc_t.add_row()
    for c, t in zip(r.cells, e): c.text = t
doc.add_paragraph()

doc.add_paragraph('El siguiente diagrama resume el proceso completo de gestión de cambios:')
add_diagram(doc, UML_PROCESO_CAMBIOS,
            'Figura 4.3.1: Diagrama de Actividad — Proceso Formal de Gestión de Cambios CCB', 5.5)

doc.add_paragraph('El siguiente diagrama muestra el ciclo de vida de cualquier Solicitud de Cambio (SC):')
add_diagram(doc, UML_ESTADO_SC,
            'Figura 4.3.1b: Diagrama de Estado — Ciclo de Vida de una Solicitud de Cambio (SC)', 5.5)

# ─────────────────────────────────────────────
# SC-001 (original — condensado pero completo)
# ─────────────────────────────────────────────
doc.add_heading('4.3.2. SC-001 — Notificaciones Automáticas por WhatsApp', 3)
two_col_table(doc, [
    ('ID Solicitud',           'SC-001'),
    ('Fecha',                  '18/06/2026'),
    ('Solicitante',            'Administración del gimnasio (stakeholder)'),
    ('RF asociado',            'RF-22 — Notificaciones por WhatsApp (v2.0)'),
    ('Descripción del cambio', 'Implementar notificaciones automáticas por WhatsApp para recordar el vencimiento de membresías a los socios con 7 días de anticipación.'),
    ('Motivo / Justificación', 'Reducir la tasa de socios que no renuevan por falta de recordatorio, y disminuir la carga de la recepción al eliminar llamadas manuales.'),
    ('Impacto esperado',       'Positivo: mejora en retención de socios y reducción de trabajo manual. Negativo: costo por integración con API externa y dependencia de tercero.'),
    ('Clasificación',          'Nueva funcionalidad'),
    ('Prioridad percibida',    'Media'),
    ('Versión objetivo',       'v2.0 — NO afecta el MVP1 congelado'),
])

doc.add_paragraph('Análisis de impacto técnico:')
for b in [
    'Módulo de Membresías: lógica de alertas para disparar mensajes vía WhatsApp Business API / Twilio.',
    'Base de datos: almacenar estado de envío de notificación por socio (nueva tabla notificaciones).',
    'Seguridad: manejo seguro de credenciales de API externa y consentimiento explícito del socio.',
    'Costo estimado: desarrollo (3 días) + pruebas de integración (2 días) + despliegue (1 día) ≈ 1 semana.',
]:
    add_bullet(doc, b)

doc.add_paragraph('Riesgos identificados:')
for r in [
    'Falla o límite de envíos de la API externa (Twilio / WhatsApp Business).',
    'Números de teléfono inválidos o no verificados en la base de datos.',
    'Incremento del costo si el volumen de socios crece (costo por mensaje).',
    'Posible incumplimiento de normativa de protección de datos sin gestión del consentimiento.',
]:
    add_bullet(doc, r)

two_col_table(doc, [
    ('Acta CCB — SC-001',      ''),
    ('Participantes',          'Alva Chacon, Cordova Guerra, Sandoval Dominguez, Melendez Bustamante'),
    ('Decisión',               'APROBADO CON CONDICIONES — Implementar en versión 2.0'),
    ('Justificación',          'El MVP1 debe priorizar los módulos críticos ya definidos. La integración con WhatsApp introduce una dependencia externa no indispensable para el lanzamiento inicial.'),
    ('Condiciones',            '1. Contar con presupuesto aprobado para el servicio de mensajería. 2. Obtener consentimiento explícito del socio. 3. Definir canal de respaldo (correo electrónico).'),
])

imp1_t = doc.add_table(rows=1, cols=3)
imp1_t.style = 'Table Grid'
header_row(imp1_t, ['Actividad', 'Responsable', 'Duración estimada'])
for row in [
    ('Configuración de cuenta API de mensajería', 'Equipo backend', '1 día'),
    ('Extensión del módulo de alertas (RF-22)',   'Equipo backend', '2 días'),
    ('Pruebas de integración',                    'Equipo QA',      '1 día'),
    ('Despliegue en ambiente de pruebas',         'DevOps',         '1 día'),
]:
    r = imp1_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

two_col_table(doc, [
    ('Verificación y validación', ''),
    ('Prueba funcional',  'Enviar notificación de prueba a número verificado y confirmar recepción.'),
    ('Prueba de integridad', 'Verificar que la fecha de vencimiento en el mensaje coincida con la BD.'),
    ('Criterio de aceptación', '100% de socios con membresía próxima a vencer (7 días) reciben la notificación correctamente.'),
])

two_col_table(doc, [
    ('ID Cambio',                'SC-001'),
    ('Fecha de implementación',  'Pendiente (versión 2.0)'),
    ('Descripción',              'Notificaciones automáticas por WhatsApp para vencimiento de membresías'),
    ('Impacto',                  'Medio — nuevo módulo de integración externa'),
    ('Resultado de validación',  'Pendiente'),
    ('Aprobado por',             'CCB del equipo'),
    ('Documentación actualizada','Sí — Matriz de trazabilidad actualizada, RF-22 creado'),
])

# ─────────────────────────────────────────────
# SC-002 NUEVA
# ─────────────────────────────────────────────
doc.add_heading('4.3.3. SC-002 — Integración con Pasarela de Pagos Online', 3)
two_col_table(doc, [
    ('ID Solicitud',           'SC-002'),
    ('Fecha',                  '25/06/2026'),
    ('Solicitante',            'Dueño/Administrador del gimnasio'),
    ('RF asociados',           'RF-20 (Pago online con tarjeta) / RF-21 (Actualización automática de membresía)'),
    ('Descripción del cambio', 'Permitir que los socios paguen su renovación de membresía en línea directamente desde la aplicación web usando tarjeta de crédito/débito, mediante integración con una pasarela de pagos (Culqi o MercadoPago).'),
    ('Motivo / Justificación', 'Actualmente el pago se realiza solo en caja física. Muchos socios solicitan la opción de pagar desde casa para evitar desplazamientos y agilizar la renovación. Esto reduciría la carga operativa de recepción y aumentaría la retención de socios.'),
    ('Impacto esperado',       'Positivo: incremento del 20-30% en renovaciones en línea según benchmarking. Negativo: comisión por transacción (3-4% por pago), implementación de seguridad PCI-DSS.'),
    ('Clasificación',          'Nueva funcionalidad'),
    ('Prioridad percibida',    'Alta'),
    ('Versión objetivo',       'v2.0 — NO afecta el MVP1 congelado'),
])

doc.add_paragraph('Análisis de impacto técnico:')
for b in [
    'Módulo de Membresías: nuevo endpoint POST /pagos/online que recibe el source_id del formulario de Culqi/MercadoPago.',
    'Integración con API externa: Culqi (Perú) o MercadoPago con SDK de Node.js. Manejo de webhooks para confirmación asíncrona.',
    'Base de datos: la tabla pagos existente se extiende con columnas charge_id y metodo_pago=\'online\'.',
    'Seguridad: no se almacenan datos de tarjeta en la BD (tokenización en el proveedor). Se requiere HTTPS y cumplimiento PCI-DSS básico.',
    'Frontend: nuevo componente de formulario de pago con SDK de Culqi (iframe embebido). No se modifica la UI del MVP1.',
    'Costo estimado: desarrollo (4 días) + pruebas de integración (2 días) + despliegue (1 día) ≈ 1.5 semanas.',
]:
    add_bullet(doc, b)

doc.add_paragraph('Análisis Costo-Beneficio:')
for b in [
    'Costo de desarrollo: ≈ S/. 2,000 (configuración inicial, pruebas y despliegue).',
    'Comisión por transacción: 3.5% por pago online (Culqi Perú).',
    'Beneficio estimado: reducción del 25% en socios que no renuevan por inconveniencia del pago presencial.',
    'ROI estimado: recuperación de la inversión en ≈ 2 meses con 50 renovaciones online/mes.',
]:
    add_bullet(doc, b)

doc.add_paragraph('Riesgos identificados:')
for r in [
    'Falla del servicio de la pasarela (Culqi/MercadoPago) durante pico de renovaciones.',
    'Chargebacks o pagos fraudulentos que requieran proceso de disputa.',
    'Rechazo de tarjetas por configuración incorrecta del sistema antifraude.',
    'Necesidad de verificación de identidad del titular de la tarjeta.',
]:
    add_bullet(doc, r)

doc.add_paragraph('El siguiente diagrama ilustra el flujo de interacción del sistema durante el proceso de pago online:')
add_diagram(doc, UML_SEC_PAGOS,
            'Figura 4.3.3: Diagrama de Secuencia — Integración con Pasarela de Pagos Online (SC-002)', 6.0)

two_col_table(doc, [
    ('Acta CCB — SC-002',      ''),
    ('Participantes',          'Alva Chacon, Cordova Guerra, Sandoval Dominguez, Melendez Bustamante'),
    ('Decisión',               'APROBADO — Implementar en versión 2.0'),
    ('Justificación',          'La integración con pasarela es independiente del MVP1. Solo añade un nuevo endpoint y extiende la tabla de pagos sin modificar la lógica de membresías existente. Impacto mínimo en el MVP1.'),
    ('Condiciones',            '1. Usar Culqi (proveedor peruano con soporte local). 2. No almacenar datos de tarjeta en la BD. 3. Implementar webhook de confirmación asíncrona para evitar estados inconsistentes.'),
])

imp2_t = doc.add_table(rows=1, cols=3)
imp2_t.style = 'Table Grid'
header_row(imp2_t, ['Actividad', 'Responsable', 'Duración estimada'])
for row in [
    ('Configuración de cuenta Culqi y credenciales API',      'Equipo backend', '0.5 día'),
    ('Desarrollo endpoint POST /pagos/online y webhook',      'Equipo backend', '3 días'),
    ('Desarrollo componente frontend (formulario Culqi)',     'Equipo frontend','1.5 días'),
    ('Pruebas de integración en ambiente sandbox',            'Equipo QA',      '2 días'),
    ('Despliegue y prueba en producción con monto mínimo',    'DevOps',         '0.5 día'),
]:
    r = imp2_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

two_col_table(doc, [
    ('Verificación y validación', ''),
    ('Prueba funcional',   'Realizar pago de prueba en sandbox con tarjeta de test. Verificar que la membresía se activa automáticamente.'),
    ('Prueba de regresión','Verificar que el flujo de pago en caja física (RF-05) no sea afectado por el nuevo endpoint.'),
    ('Criterio de aceptación', '1. El pago online actualiza la membresía en < 5 segundos tras la confirmación. 2. Un pago fallido no modifica el estado de la membresía. 3. El log de pagos refleja el método \'online\' correctamente.'),
])

two_col_table(doc, [
    ('ID Cambio',                'SC-002'),
    ('Fecha de implementación',  'Planificado Q3 2026 (versión 2.0)'),
    ('Descripción',              'Integración con pasarela de pagos Culqi para renovaciones online de membresía'),
    ('Impacto',                  'Alto — nuevo módulo de integración externa con impacto en experiencia del socio'),
    ('Resultado de validación',  'Pendiente — ambiente sandbox configurado'),
    ('Aprobado por',             'CCB del equipo (consenso 4/4)'),
    ('Documentación actualizada','Sí — RF-20 y RF-21 añadidos a la matriz de trazabilidad'),
])

# ─────────────────────────────────────────────
# SC-003 NUEVA
# ─────────────────────────────────────────────
doc.add_heading('4.3.4. SC-003 — Aplicación Móvil para Socios', 3)
two_col_table(doc, [
    ('ID Solicitud',           'SC-003'),
    ('Fecha',                  '02/07/2026'),
    ('Solicitante',            'Socios del gimnasio (resultado del cuestionario digital — 78% solicitaron app móvil)'),
    ('RF asociados',           'RF-27 (App Móvil — Ver QR) / RF-28 (Reservar desde móvil) / RF-29 (Ver progreso físico móvil)'),
    ('Descripción del cambio', 'Desarrollar una aplicación móvil nativa multiplataforma (iOS y Android) usando React Native que permita a los socios consultar su código QR, gestionar reservas de clases y ver su progreso físico desde su smartphone sin necesidad de abrir el navegador.'),
    ('Motivo / Justificación', 'El 78% de los socios en el cuestionario digital indicaron que prefieren una app móvil sobre acceder desde el navegador. El QR en el teléfono elimina la necesidad de imprimir o buscar el enlace. Además, las reservas desde el móvil aumentarían la tasa de ocupación de clases.'),
    ('Impacto esperado',       'Positivo: mejora significativa en la experiencia del socio, mayor adopción del sistema. Negativo: esfuerzo de desarrollo considerable (estimado 6-8 semanas), mantenimiento de dos clientes (web + móvil).'),
    ('Clasificación',          'Nuevo canal (aplicación móvil)'),
    ('Prioridad percibida',    'Alta'),
    ('Versión objetivo',       'v3.0 — NO afecta el MVP1 congelado. Reutiliza los endpoints REST existentes.'),
])

doc.add_paragraph('Análisis de impacto técnico:')
for b in [
    'Backend: la app móvil reutiliza los mismos endpoints REST del MVP1 sin necesidad de crear nuevos (GET /socios/:id/qr, POST /reservas, GET /progreso). Impacto en backend: NULO.',
    'Frontend web: no requiere ninguna modificación. La app móvil es un cliente independiente.',
    'App móvil: desarrollada con React Native (permite reutilizar lógica de la app web React). Pantallas: QR, Mis Reservas, Mi Progreso, Notificaciones Push.',
    'Notificaciones Push: integración con Firebase Cloud Messaging (FCM) o OneSignal para notificaciones de vencimiento (complementa SC-001).',
    'Autenticación: usa el mismo JWT del MVP1. Se añade gestión de sesión persistente en el dispositivo (AsyncStorage).',
    'Costo estimado: desarrollo (6 semanas) + pruebas en dispositivos (1 semana) + publicación en tiendas (0.5 semanas) ≈ 7.5 semanas.',
]:
    add_bullet(doc, b)

doc.add_paragraph('Análisis Costo-Beneficio:')
for b in [
    'Costo de desarrollo: ≈ S/. 8,000 (React Native + pruebas en dispositivos físicos).',
    'Costo publicación en tiendas: Google Play (USD 25 único) + Apple App Store (USD 99/año).',
    'Beneficio estimado: incremento del 40% en uso del sistema por parte de socios (mayor engagement).',
    'Beneficio colateral: las reservas desde móvil aumentan la tasa de ocupación de clases grupales.',
]:
    add_bullet(doc, b)

doc.add_paragraph('Riesgos identificados:')
for r in [
    'Fragmentación de dispositivos Android (múltiples versiones de SO, tamaños de pantalla).',
    'Proceso de aprobación de Apple App Store puede tomar 1-2 semanas.',
    'Mantenimiento paralelo de dos clientes (web y móvil) incrementa la deuda técnica.',
    'Rendimiento del QR en pantallas de baja resolución puede dificultar el escaneo.',
]:
    add_bullet(doc, r)

doc.add_paragraph('El siguiente diagrama muestra la arquitectura de la app móvil y su relación con el backend existente:')
add_diagram(doc, UML_COMP_APP_MOVIL,
            'Figura 4.3.4: Diagrama de Componentes — Arquitectura App Móvil (SC-003, v3.0)', 6.0)

two_col_table(doc, [
    ('Acta CCB — SC-003',      ''),
    ('Participantes',          'Alva Chacon, Cordova Guerra, Sandoval Dominguez, Melendez Bustamante'),
    ('Decisión',               'APROBADO CON CONDICIONES — Implementar en versión 3.0'),
    ('Justificación',          'La app móvil no modifica el backend del MVP1 (reutiliza los mismos endpoints). Sin embargo, el esfuerzo de desarrollo es considerable, por lo que se difiere a v3.0 para garantizar la estabilidad del MVP1 primero.'),
    ('Condiciones',            '1. Completar y estabilizar el MVP1 antes de iniciar el desarrollo. 2. Usar React Native para maximizar la reutilización de código. 3. Iniciar con Android (Google Play) antes que iOS para reducir fricción de publicación.'),
])

imp3_t = doc.add_table(rows=1, cols=3)
imp3_t.style = 'Table Grid'
header_row(imp3_t, ['Actividad', 'Responsable', 'Duración estimada'])
for row in [
    ('Configuración del proyecto React Native',     'Equipo móvil',   '0.5 semana'),
    ('Desarrollo pantalla QR y autenticación',      'Equipo móvil',   '1 semana'),
    ('Desarrollo pantalla Mis Reservas',            'Equipo móvil',   '1.5 semanas'),
    ('Desarrollo pantalla Mi Progreso',             'Equipo móvil',   '1 semana'),
    ('Integración de notificaciones Push (FCM)',    'Equipo backend',  '0.5 semana'),
    ('Pruebas en dispositivos físicos',             'Equipo QA',       '1 semana'),
    ('Publicación en Google Play',                  'DevOps',          '0.5 semana'),
]:
    r = imp3_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

two_col_table(doc, [
    ('ID Cambio',                'SC-003'),
    ('Fecha de implementación',  'Planificado Q1 2027 (versión 3.0)'),
    ('Descripción',              'App móvil iOS/Android para socios con QR, reservas y progreso físico'),
    ('Impacto',                  'Alto — nuevo canal de acceso, sin impacto en MVP1'),
    ('Resultado de validación',  'Pendiente'),
    ('Aprobado por',             'CCB del equipo (consenso 4/4)'),
    ('Documentación actualizada','Sí — RF-27, RF-28, RF-29 añadidos a la matriz de trazabilidad'),
])

# ─────────────────────────────────────────────
# SC-004 NUEVA
# ─────────────────────────────────────────────
doc.add_heading('4.3.5. SC-004 — Módulo de Planes Nutricionales', 3)
two_col_table(doc, [
    ('ID Solicitud',           'SC-004'),
    ('Fecha',                  '04/07/2026'),
    ('Solicitante',            'Entrenadores del gimnasio + Socios (resultado de entrevista con entrenadores jefes)'),
    ('RF asociados',           'RF-30 (Ficha médica ampliada) / RF-31 (Catálogo nutricional) / RF-32 (Asignación) / RF-33 (Registro diario)'),
    ('Descripción del cambio', 'Incorporar un módulo de planes nutricionales personalizados al sistema, que permita a los entrenadores/nutricionistas crear planes de alimentación, asignarlos a socios y que los socios puedan registrar su ingesta diaria. La ficha médica se extendería con campos de restricciones alimentarias y alergias.'),
    ('Motivo / Justificación', 'El 65% de los socios en el cuestionario indicaron interés en un plan nutricional. El gimnasio desea diferenciarse de la competencia ofreciendo un servicio integral de entrenamiento + nutrición. Actualmente este servicio se ofrece manualmente en cuadernos o WhatsApp.'),
    ('Impacto esperado',       'Positivo: diferenciación del servicio, mayor fidelización de socios, posibilidad de cobrar un plus por el servicio nutricional. Negativo: requiere definir si el nutricionista es un nuevo rol en el sistema (RBAC) o si el entrenador asume esa función.'),
    ('Clasificación',          'Nuevo módulo del sistema'),
    ('Prioridad percibida',    'Media'),
    ('Versión objetivo',       'v3.0 — NO afecta el MVP1 congelado. Se añaden nuevas tablas sin modificar las existentes.'),
])

doc.add_paragraph('Análisis de impacto técnico:')
for b in [
    'Base de datos: se crean 3 tablas nuevas (planes_nutricionales, asignaciones_nutricionales, registros_comidas) sin modificar ninguna tabla existente del MVP1.',
    'Tabla fichas_medicas: se añaden 2 columnas nuevas (restricciones_alimentarias, alergias) como nullable para no romper los registros existentes.',
    'Backend: 4 nuevos controllers y sus endpoints (GET/POST /nutricion/planes, POST /nutricion/asignaciones, POST /nutricion/registros-comidas).',
    'Frontend: nueva sección "Nutrición" en el panel del entrenador y nueva pestaña en el perfil del socio.',
    'Roles: se evalúa crear rol "nutricionista" o asignar permisos al rol "entrenador". Decisión a tomar en diseño detallado.',
    'Costo estimado: desarrollo BD (1 día) + API (3 días) + UI (4 días) + pruebas (2 días) ≈ 2 semanas.',
]:
    add_bullet(doc, b)

doc.add_paragraph('Análisis Costo-Beneficio:')
for b in [
    'Costo de desarrollo: ≈ S/. 3,500 (backend + frontend + pruebas).',
    'Beneficio estimado: monetización adicional por servicio nutricional premium (S/. 50/mes por socio interesado).',
    'Beneficio colateral: mayor retención de socios al ofrecer servicio integral de entrenamiento + nutrición.',
    'ROI estimado: con 30 socios que contraten el servicio premium, recuperación en < 3 meses.',
]:
    add_bullet(doc, b)

doc.add_paragraph('Riesgos identificados:')
for r in [
    'Definición ambigua de responsabilidades: entrenador vs. nutricionista certificado (implicancias legales).',
    'Calidad de los datos ingresados: los registros de comidas dependen de la disciplina del socio.',
    'Complejidad de los planes: planes nutricionales médicamente correctos requieren expertise especializado.',
    'Volumen de datos: registros diarios de comidas pueden crecer rápidamente (optimización de BD necesaria).',
]:
    add_bullet(doc, r)

doc.add_paragraph('El siguiente diagrama muestra el modelo de datos del nuevo módulo nutricional:')
add_diagram(doc, UML_CLASE_NUTRICION,
            'Figura 4.3.5: Diagrama de Clases — Modelo de Datos del Módulo Nutricional (SC-004, v3.0)', 6.0)

two_col_table(doc, [
    ('Acta CCB — SC-004',      ''),
    ('Participantes',          'Alva Chacon, Cordova Guerra, Sandoval Dominguez, Melendez Bustamante'),
    ('Decisión',               'APROBADO CON CONDICIONES — Implementar en versión 3.0'),
    ('Justificación',          'El módulo no modifica ninguna tabla existente del MVP1 (solo agrega columnas nullable y tablas nuevas). Sin embargo, requiere resolver primero la definición del rol nutricionista y la estrategia de precios del servicio con el cliente.'),
    ('Condiciones',            '1. Definir con el cliente si se crea rol "nutricionista" o se usa el rol "entrenador". 2. Los nuevos campos en fichas_medicas deben ser nullable para no romper registros existentes. 3. Consultar con un nutricionista certificado para validar la estructura de los planes antes de implementar.'),
])

imp4_t = doc.add_table(rows=1, cols=3)
imp4_t.style = 'Table Grid'
header_row(imp4_t, ['Actividad', 'Responsable', 'Duración estimada'])
for row in [
    ('Migración de BD: nuevas tablas y columnas nullable', 'Equipo backend',  '1 día'),
    ('API: endpoints de planes y asignaciones nutricionales', 'Equipo backend', '3 días'),
    ('Frontend: sección Nutrición en panel del entrenador', 'Equipo frontend', '3 días'),
    ('Frontend: pestaña nutrición en perfil del socio',    'Equipo frontend', '1.5 días'),
    ('Pruebas funcionales y de regresión',                 'Equipo QA',       '2 días'),
]:
    r = imp4_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

two_col_table(doc, [
    ('ID Cambio',                'SC-004'),
    ('Fecha de implementación',  'Planificado Q2 2027 (versión 3.0)'),
    ('Descripción',              'Módulo de planes nutricionales personalizados: catálogo, asignación y registro diario de comidas'),
    ('Impacto',                  'Medio — nuevo módulo independiente, sin impacto en MVP1'),
    ('Resultado de validación',  'Pendiente — sujeto a definición del rol nutricionista'),
    ('Aprobado por',             'CCB del equipo (consenso 4/4)'),
    ('Documentación actualizada','Sí — RF-30 a RF-33 añadidos a la matriz de trazabilidad'),
])

# 4.4 Resumen Bitácora
doc.add_heading('4.4. Resumen de la Bitácora de Cambios', 2)
doc.add_paragraph('La siguiente tabla consolida todas las solicitudes de cambio gestionadas hasta la fecha:')
bit_t = doc.add_table(rows=1, cols=7)
bit_t.style = 'Table Grid'
header_row(bit_t, ['SC-ID', 'Descripción', 'Solicitante', 'Prioridad', 'Versión', 'Decisión CCB', 'Estado'])
bitacora = [
    ('SC-001', 'Notificaciones WhatsApp membresía',    'Administración',        'Media', 'v2.0', 'Aprobado',            'Pendiente implementación'),
    ('SC-002', 'Pasarela de pagos online (Culqi)',     'Dueño/Administrador',   'Alta',  'v2.0', 'Aprobado',            'Pendiente implementación'),
    ('SC-003', 'Aplicación móvil iOS/Android',         'Socios (encuesta 78%)', 'Alta',  'v3.0', 'Aprobado c/cond.',    'Diferido a v3.0'),
    ('SC-004', 'Módulo de planes nutricionales',       'Entrenadores + Socios', 'Media', 'v3.0', 'Aprobado c/cond.',    'Diferido a v3.0'),
]
BIT_COLORS = {'v2.0': 'FFF3CD', 'v3.0': 'D0D8E4'}
for row in bitacora:
    r = bit_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
    set_cell_bg(r.cells[4], BIT_COLORS.get(row[4], 'FFFFFF'))
doc.add_paragraph()

doc.add_paragraph(
    'Nota: Ninguna de las 4 solicitudes de cambio aprobadas modifica módulos '
    'del MVP1 congelado. SC-001 y SC-002 están planificadas para la versión 2.0, '
    'mientras que SC-003 y SC-004 se desarrollarán en la versión 3.0.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 5: Propuesta de Valor
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Propuesta de Valor del Proyecto', 1)
doc.add_paragraph(
    'OLYMPUS CORE no solo busca digitalizar procesos administrativos, sino transformar '
    'la experiencia operativa y de autogestión dentro del gimnasio mediante automatización, '
    'control inteligente y herramientas modernas. La propuesta de valor se articula en '
    'tres horizontes temporales:'
)

doc.add_heading('5.1. Valor Inmediato — MVP1 (Implementado)', 2)
for v in [
    'Reducción del tiempo de atención en recepción: de 60-90 segundos a < 5 segundos mediante validación QR automática.',
    'Eliminación de pérdidas por accesos indebidos: el sistema bloquea automáticamente membresías vencidas sin intervención humana.',
    'Control financiero centralizado: todos los pagos, membresías y estados de socios en una sola plataforma accesible desde cualquier navegador.',
    'Reducción de errores humanos: el sistema de reservas con validación de aforo y penalizaciones automáticas elimina inconsistencias del proceso manual.',
    'Centralización de toda la información del negocio: socios, membresías, clases, pagos y progreso físico en una BD PostgreSQL estructurada.',
    'Control de acceso RBAC: cada actor accede solo a los módulos que le corresponden, mejorando la seguridad operativa.',
]:
    add_bullet(doc, v)

doc.add_heading('5.2. Valor a Mediano Plazo — Versión 2.0', 2)
for v in [
    'SC-001: Incremento del 15-20% en la tasa de renovación de membresías gracias a notificaciones automáticas de vencimiento por WhatsApp.',
    'SC-002: Habilitación del pago online 24/7 — los socios pueden renovar su membresía sin visitar el gimnasio, incrementando la retención.',
    'Módulo de reportes exportables (RF-25): el administrador puede generar informes de ingresos en PDF para análisis financiero mensual.',
    'Módulo de promociones y convenios (RF-24): capacidad de aplicar descuentos corporativos controlados, abriendo un nuevo segmento de mercado.',
]:
    add_bullet(doc, v)

doc.add_heading('5.3. Valor a Largo Plazo — Versión 3.0', 2)
for v in [
    'SC-003: Aplicación móvil iOS/Android — convierte OLYMPUS CORE en un gimnasio tecnológicamente diferenciado. El socio lleva el gimnasio en su bolsillo.',
    'SC-004: Módulo nutricional — servicio integral entrenamiento + nutrición que permite al gimnasio posicionarse como un centro de salud completo y cobrar un servicio premium adicional.',
    'Integración con wearables (RF-34): sincronización con smartwatches para importar datos de entrenamiento automáticamente al perfil del socio.',
    'Escalabilidad: la arquitectura cloud (Vercel + Koyeb + Neon) permite escalar a múltiples sedes del gimnasio sin rediseñar el sistema.',
]:
    add_bullet(doc, v)

doc.add_heading('5.4. Resumen de Impacto por Dimensión', 2)
impactos = [
    ('Control de acceso',   '60-90 seg por socio, errores frecuentes, pérdidas por membresías vencidas', '< 5 seg por socio, automático, cero pérdidas por acceso indebido'),
    ('Gestión de pagos',    'Registro en Excel, propenso a errores, sin historial claro',                'BD centralizada, historial completo, pago online en v2.0'),
    ('Reservas de clases',  'Por WhatsApp/llamada, sin control de aforo, cancelaciones sin registro',    'Online 24/7, aforo en tiempo real, penalizaciones automáticas'),
    ('Experiencia del socio','Sin autogestión, depende 100% de recepción',                              'QR personal, perfil web, app móvil en v3.0'),
    ('Control financiero',  'Dashboard inexistente, ingresos calculados manualmente',                    'Dashboard en tiempo real, reportes PDF en v2.0'),
    ('Seguimiento fitness',  'Cuadernos físicos, sin historial digital',                                 'Rutinas y progreso físico digitalizados, nutrición en v3.0'),
]
imp_t = doc.add_table(rows=1, cols=3)
imp_t.style = 'Table Grid'
header_row(imp_t, ['Dimensión', 'Situación Actual (Manual)', 'Con OLYMPUS CORE (Digital)'])
for row in impactos:
    r = imp_t.add_row()
    for c, t in zip(r.cells, row): c.text = t

doc.add_paragraph()


# ── GUARDAR ─────────────────────────────────────────────────────────────────────
output = (
    r'c:/Users/ACER/OneDrive/Desktop/PROGRAMACION/UNIVERSITY PROJECTS/'
    r'SIstemaGestorGimnasio/Documentacion/Documentos Actualizados/'
    r'GESTION_REQUISITOS_PROPUESTA_VALOR_actualizado.docx'
)
doc.save(output)
print(f'Documento guardado en: {output}')
