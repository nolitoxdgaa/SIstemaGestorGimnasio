from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table_row(table, cells, bold_first=False, bg=None):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].bold = True
        if bg:
            set_cell_bg(cell, bg)
    return row

doc = Document()

# === Estilo base ===
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─────────────────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('OLYMPUS CORE')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Sistema Gestor de Gimnasio')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x44, 0x4A, 0x57)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('DEFINICIÓN FUNCIONAL DEL MVP1')
r3.bold = True
r3.font.size = Pt(20)
r3.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()

for line in ['Versión: MVP1', 'Grupo: 5', 'Entrega: Julio 2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line).font.size = Pt(12)

doc.add_page_break()

# ─────────────────────────────────────────────────────────
# 1. PROPÓSITO DEL DOCUMENTO
# ─────────────────────────────────────────────────────────
add_heading(doc, '1. Propósito del documento', 1, '1E3A5F')
doc.add_paragraph(
    'Este documento define el alcance funcional completo del MVP1 de OLYMPUS CORE. '
    'Especifica qué hace el sistema, qué no hace (y por qué), bajo qué reglas de negocio opera, '
    'y cómo se mapea con los requisitos funcionales formalmente identificados en el proyecto.'
)

# ─────────────────────────────────────────────────────────
# 2. ALCANCE DEL MVP1
# ─────────────────────────────────────────────────────────
doc.add_heading('2. Alcance del MVP1', 1)
doc.add_paragraph(
    'El MVP1 implementa los módulos de mayor valor y menor riesgo según la priorización '
    'MoSCoW acordada en los documentos del curso. Cubre todos los Must Have y los '
    'Should Have más impactantes.'
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Módulo'
hdr[1].text = 'Descripción'
hdr[2].text = 'RF cubiertos'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, '1E3A5F')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

modules = [
    ('🔐 Autenticación y Seguridad', 'Login por roles, bloqueo de cuenta, JWT, cierre de sesión', 'RF-01, RF-10, RF-11'),
    ('👥 Gestión de Socios', 'Registro, edición, búsqueda, desactivación lógica de socios', 'RF-02, RF-12'),
    ('💳 Membresías y Planes', 'Activación de membresía, visualización de estado y vencimiento', 'RF-03, RF-04'),
    ('💵 Registro de Pagos', 'Registro de pagos con distintos métodos, historial, generación de membresía', 'RF-05'),
    ('📱 Acceso por Código QR', 'Generación de QR único por socio, validación en recepción, log de accesos', 'RF-06, RF-07'),
    ('📅 Clases Grupales', 'CRUD de clases, control de aforo, tipos: spinning, crossfit, yoga, zumba', 'RF-08'),
    ('📋 Reservas de Clases', 'Reserva con validaciones (aforo, membresía), cancelación con plazo límite', 'RF-09'),
    ('⚠️ Sistema de Strikes/Penalizaciones', 'Acumulación de inasistencias, bloqueo automático a los 3 strikes en 30 días', 'RF-13, RF-14'),
    ('📊 Dashboard del Administrador', 'Ingresos del día, membresías por vencer, ocupación de clases, socios activos', 'RF-15'),
    ('👤 Mi Perfil del Socio', 'Vista del socio con su QR, estado de membresía, reservas activas y strikes', 'RF-16'),
]
for row_data in modules:
    add_table_row(table, row_data)

# ─────────────────────────────────────────────────────────
# 3. MÓDULOS Y FUNCIONALIDADES DETALLADAS
# ─────────────────────────────────────────────────────────
doc.add_heading('3. Módulos y Funcionalidades Implementadas', 1)

module_details = [
    ('3.1 Módulo de Autenticación y Seguridad', [
        ('Inicio de sesión con email y contraseña', 'administrador, recepcionista, entrenador, socio'),
        ('Generación y validación de JWT', 'Todos los roles'),
        ('Bloqueo de cuenta tras 5 intentos fallidos (30 min)', 'Sistema automático'),
        ('Cierre de sesión seguro', 'Todos los roles'),
        ('Control de acceso por roles (RBAC)', 'Todos los módulos'),
    ]),
    ('3.2 Módulo de Gestión de Socios', [
        ('Registrar nuevo socio con DNI, email, teléfono', 'Administrador, Recepcionista'),
        ('Buscar socios por nombre o DNI', 'Administrador, Recepcionista'),
        ('Editar datos de un socio', 'Administrador, Recepcionista'),
        ('Desactivar socio (borrado lógico)', 'Administrador'),
        ('Visualizar estado de membresía y strikes en tarjeta', 'Administrador, Recepcionista'),
        ('Filtrar socios por estado (activo/inactivo)', 'Administrador, Recepcionista'),
        ('Paginación del listado de socios', 'Administrador, Recepcionista'),
    ]),
    ('3.3 Módulo de Membresías', [
        ('Visualizar membresías activas, por vencer (3d) y vencidas', 'Administrador, Recepcionista'),
        ('Filtro por categoría de estado de membresía', 'Administrador, Recepcionista'),
        ('Renovación de membresía mediante registro de pago', 'Administrador, Recepcionista'),
        ('Indicador de días restantes por vencimiento', 'Administrador, Recepcionista'),
        ('Estado de membresía: activa / vencida / bloqueada', 'Sistema automático'),
    ]),
    ('3.4 Módulo de Pagos', [
        ('Registrar pago por plan (Básico / Premium / VIP)', 'Administrador, Recepcionista'),
        ('Seleccionar método de pago: efectivo, Yape, Plin, tarjeta', 'Administrador, Recepcionista'),
        ('Activación automática de membresía tras pago exitoso', 'Sistema automático'),
        ('Historial de pagos de un socio', 'Administrador, Recepcionista, Socio'),
        ('Listado general de pagos filtrable por fecha y método', 'Administrador'),
    ]),
    ('3.5 Módulo de Acceso por Código QR', [
        ('Generación de código QR único con token JWT por socio', 'Administrador, Recepcionista'),
        ('Validación de QR en recepción: acceso permitido o denegado', 'Recepcionista'),
        ('Motivos de denegación: membresía vencida, penalización activa, QR inválido', 'Sistema automático'),
        ('Registro automático de log de acceso (PERMITIDO / DENEGADO)', 'Sistema automático'),
        ('Descarga e impresión del código QR', 'Administrador, Recepcionista'),
    ]),
    ('3.6 Módulo de Clases Grupales', [
        ('Crear clase con tipo, instructor, horario y aforo máximo', 'Administrador'),
        ('Editar datos de una clase existente', 'Administrador'),
        ('Cancelar una clase', 'Administrador'),
        ('Visualizar listado de clases con aforo disponible en tiempo real', 'Todos'),
        ('Filtrar clases por tipo y fecha', 'Todos'),
        ('Estado automático: disponible / llena / cancelada', 'Sistema automático'),
        ('Tipos de clase: Spinning, CrossFit, Yoga, Zumba', 'Administrador'),
    ]),
    ('3.7 Módulo de Reservas', [
        ('Reservar una clase con validación de membresía activa', 'Socio, Recepcionista'),
        ('Validación de cupo máximo al reservar', 'Sistema automático'),
        ('Bloqueo de reserva si el socio tiene penalización activa', 'Sistema automático'),
        ('Cancelar reserva (hasta 2 horas antes del inicio)', 'Socio, Recepcionista, Administrador'),
        ('Marcar asistencia o inasistencia a una clase', 'Recepcionista, Administrador'),
        ('Visualizar reservas del día en panel administrativo', 'Administrador, Recepcionista'),
        ('El socio puede ver sus propias reservas activas e historial', 'Socio'),
    ]),
    ('3.8 Módulo de Strikes y Penalizaciones', [
        ('Registro de inasistencia cuando el socio no se presenta a clase reservada', 'Recepcionista, Administrador'),
        ('Acumulación de strikes (inasistencias injustificadas en 30 días)', 'Sistema automático'),
        ('Bloqueo automático de membresía al llegar a 3 strikes', 'Sistema automático'),
        ('Bloqueo de nuevas reservas durante período de penalización', 'Sistema automático'),
        ('Visualización del contador de strikes en tarjeta de socio', 'Administrador, Recepcionista'),
        ('Justificación y eliminación de un strike (con registro)', 'Administrador'),
    ]),
    ('3.9 Módulo de Dashboard del Administrador', [
        ('Total de socios activos en el sistema', 'Administrador'),
        ('Ingresos del día (monto y cantidad de pagos)', 'Administrador'),
        ('Membresías próximas a vencer (en los próximos 3 días)', 'Administrador'),
        ('Ocupación de clases del día en tiempo real', 'Administrador'),
        ('Resumen visual de estado del gimnasio al iniciar sesión', 'Administrador'),
    ]),
    ('3.10 Módulo Mi Perfil del Socio', [
        ('Visualización del código QR personal del socio', 'Socio'),
        ('Estado actual de membresía (activa, vencida, bloqueada)', 'Socio'),
        ('Fecha de vencimiento de membresía', 'Socio'),
        ('Contador de strikes activos y pendientes', 'Socio'),
        ('Listado de reservas confirmadas del socio', 'Socio'),
    ]),
]

for mod_title, funcs in module_details:
    doc.add_heading(mod_title, 2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = 'Funcionalidad'
    hdr_cells[1].text = 'Rol que la ejecuta'
    for c in hdr_cells:
        c.paragraphs[0].runs[0].bold = True
        set_cell_bg(c, 'D0D8E4')
    for func, role in funcs:
        r = tbl.add_row()
        r.cells[0].text = func
        r.cells[1].text = role
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────
# 4. MÓDULOS EXCLUIDOS DEL MVP1
# ─────────────────────────────────────────────────────────
doc.add_heading('4. Módulos Excluidos del MVP1', 1)
doc.add_paragraph('Los siguientes módulos fueron definidos en el alcance completo del sistema pero quedan diferidos para versiones futuras:')

excl_table = doc.add_table(rows=1, cols=3)
excl_table.style = 'Table Grid'
for i, txt in enumerate(['Módulo excluido', 'Razón', 'Versión futura']):
    excl_table.rows[0].cells[i].text = txt
    excl_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(excl_table.rows[0].cells[i], '1E3A5F')
    excl_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

excluidos = [
    ('Rutinas de Entrenamiento', 'Alta complejidad; sin valor sin ficha médica previa', 'v1.1'),
    ('Progreso Físico', 'Depende del módulo de Rutinas', 'v1.1'),
    ('Ficha Médica Digital', 'Depende de Rutinas para tener sentido completo', 'v1.1'),
    ('Notificaciones Automáticas', 'Diferida por complejidad de integración (email/SMS)', 'v2.0'),
    ('Módulo Entrenador Completo', 'Sin rutinas/progreso, el rol entrenador queda vacío', 'v1.1'),
    ('Logs de Acceso (pantalla dedicada)', 'El endpoint existe pero sin UI dedicada', 'v1.1'),
    ('Reportes y exportación de datos', 'Requiere módulo de analytics adicional', 'v2.0'),
]
for row_data in excluidos:
    add_table_row(excl_table, row_data)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────
# 5. REGLAS DE NEGOCIO
# ─────────────────────────────────────────────────────────
doc.add_heading('5. Reglas de Negocio Implementadas', 1)

rn_table = doc.add_table(rows=1, cols=3)
rn_table.style = 'Table Grid'
for i, txt in enumerate(['ID', 'Regla de Negocio', 'Estado en MVP1']):
    rn_table.rows[0].cells[i].text = txt
    rn_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(rn_table.rows[0].cells[i], '1E3A5F')
    rn_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

reglas = [
    ('RN-01', 'Membresía vencida = acceso físico bloqueado por QR', '✅ Implementado'),
    ('RN-02', 'Cada clase tiene un aforo máximo; no se permiten reservas adicionales', '✅ Implementado'),
    ('RN-03', 'Cancelación de reserva solo hasta 2 horas antes del inicio de la clase', '✅ Implementado'),
    ('RN-04', '3 strikes en 30 días = bloqueo de reservas por 7 días', '✅ Implementado'),
    ('RN-05', 'Catálogo cerrado de precios (planes fijos gestionados por el admin)', '✅ Implementado'),
    ('RN-06', 'Ficha médica obligatoria antes de asignar una rutina de entrenamiento', '⏭️ Diferido v1.1'),
    ('RN-07', 'Un socio no puede tener dos reservas para la misma clase simultáneamente', '✅ Implementado'),
    ('RN-08', 'Dashboard obligatorio con métricas: ingresos, vencimientos, ocupación', '✅ Implementado'),
    ('RN-09', 'El QR tiene expiración (token JWT de tiempo limitado)', '✅ Implementado'),
    ('RN-10', 'Contraseña bloqueada por 30 minutos tras 5 intentos fallidos de login', '✅ Implementado'),
]
for row_data in reglas:
    r = rn_table.add_row()
    for i, (cell, text) in enumerate(zip(r.cells, row_data)):
        cell.text = text
        if '✅' in text:
            pass  # verde
        elif '⏭️' in text:
            set_cell_bg(cell, 'FFF3CD')

doc.add_paragraph()

# ─────────────────────────────────────────────────────────
# 6. MATRIZ DE TRAZABILIDAD RF
# ─────────────────────────────────────────────────────────
doc.add_heading('6. Matriz de Trazabilidad de Requisitos Funcionales', 1)
doc.add_paragraph(
    'La siguiente tabla asocia los 34 Requisitos Funcionales (RF) identificados en el proyecto '
    'con su nivel de prioridad MoSCoW, su presencia en el MVP1 y el módulo donde se implementa.'
)

mt = doc.add_table(rows=1, cols=5)
mt.style = 'Table Grid'
for i, txt in enumerate(['ID RF', 'Descripcion', 'MoSCoW', 'En MVP1', 'Modulo']):
    mt.rows[0].cells[i].text = txt
    mt.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(mt.rows[0].cells[i], '1E3A5F')
    mt.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

rfs = [
    # MUST HAVE
    ('RF-01', 'Autenticacion con email y contrasena por rol', 'Must Have', 'Si', 'Autenticacion'),
    ('RF-02', 'Registro y gestion de socios (CRUD completo)', 'Must Have', 'Si', 'Socios'),
    ('RF-03', 'Gestion de planes de membresia (nombre, precio, duracion)', 'Must Have', 'Si', 'Membresias / Planes'),
    ('RF-04', 'Visualizacion del estado de membresia (activa/vencida/bloqueada)', 'Must Have', 'Si', 'Membresias'),
    ('RF-05', 'Registro de pagos con generacion automatica de membresia', 'Must Have', 'Si', 'Pagos'),
    ('RF-06', 'Generacion de codigo QR unico por socio con token de acceso', 'Must Have', 'Si', 'Acceso QR'),
    ('RF-07', 'Validacion de acceso fisico mediante codigo QR', 'Must Have', 'Si', 'Acceso QR'),
    ('RF-08', 'Gestion de clases grupales (CRUD, aforo, horarios)', 'Must Have', 'Si', 'Clases'),
    ('RF-09', 'Reserva y cancelacion de clases grupales por el socio', 'Must Have', 'Si', 'Reservas'),
    ('RF-10', 'Bloqueo de cuenta por 5 intentos fallidos de autenticacion (30 min)', 'Must Have', 'Si', 'Autenticacion'),
    ('RF-11', 'Control de acceso basado en roles (RBAC) en cada endpoint', 'Must Have', 'Si', 'Autenticacion'),
    # SHOULD HAVE
    ('RF-12', 'Busqueda y filtrado de socios por nombre, DNI o estado', 'Should Have', 'Si', 'Socios'),
    ('RF-13', 'Registro de inasistencias y acumulacion de strikes de penalizacion', 'Should Have', 'Si', 'Penalizaciones'),
    ('RF-14', 'Bloqueo automatico por acumulacion de 3 strikes en 30 dias', 'Should Have', 'Si', 'Penalizaciones'),
    ('RF-15', 'Dashboard con metricas: ingresos, socios activos, vencimientos, ocupacion', 'Should Have', 'Si', 'Dashboard'),
    ('RF-16', 'Vista Mi Perfil del socio: QR, membresia, strikes y reservas activas', 'Should Have', 'Si', 'Mi Perfil'),
    ('RF-17', 'Justificacion y eliminacion de strikes por el administrador con registro', 'Should Have', 'Si', 'Penalizaciones'),
    ('RF-18', 'Historial de pagos accesible para admin, recepcionista y socio', 'Should Have', 'Si', 'Pagos'),
    ('RF-19', 'Log de accesos QR con resultado PERMITIDO/DENEGADO y motivo', 'Should Have', 'Parcial', 'Acceso QR'),
    # COULD HAVE
    ('RF-20', 'Ficha medica digital: lesiones, enfermedades y objetivo de entrenamiento', 'Could Have', 'No (v1.1)', 'Futuro'),
    ('RF-21', 'Asignacion de rutinas de entrenamiento personalizadas a socios', 'Could Have', 'No (v1.1)', 'Futuro'),
    ('RF-22', 'Registro de progreso fisico por ejercicio (peso, repeticiones, fecha)', 'Could Have', 'No (v1.1)', 'Futuro'),
    ('RF-23', 'Pantalla dedicada de logs de acceso con filtros avanzados', 'Could Have', 'No (v1.1)', 'Futuro'),
    ('RF-24', 'Generacion de comprobante de pago en PDF al momento del cobro', 'Could Have', 'No (v1.1)', 'Futuro'),
    ('RF-25', 'Indicador visual de ocupacion de clase para decision de reserva', 'Could Have', 'No (v1.1)', 'Futuro'),
    ('RF-26', 'Lista de espera para clases llenas con notificacion al liberarse cupo', 'Could Have', 'No (v1.1)', 'Futuro'),
    ('RF-27', 'Asignacion de entrenador responsable permanente a cada socio', 'Could Have', 'No (v1.1)', 'Futuro'),
    # WON'T HAVE
    ('RF-28', 'Notificaciones automaticas por email/SMS de vencimiento de membresia', "Won't Have", 'No (v2.0)', 'Futuro'),
    ('RF-29', 'Portal de pago online con pasarela de tarjeta credito/debito', "Won't Have", 'No (v2.0)', 'Futuro'),
    ('RF-30', 'Reportes exportables en Excel/PDF con filtros por rango de fechas', "Won't Have", 'No (v2.0)', 'Futuro'),
    ('RF-31', 'Aplicacion movil nativa (Android/iOS) para socios', "Won't Have", 'No (v3.0)', 'Futuro'),
    ('RF-32', 'Modulo de inventario de equipamiento del gimnasio', "Won't Have", 'No (v3.0)', 'Futuro'),
    ('RF-33', 'Integracion con torniquete electronico real via API/IoT', "Won't Have", 'No (v3.0)', 'Futuro'),
    ('RF-34', 'Codigo QR dinamico con renovacion automatica cada N minutos', "Won't Have", 'No (v3.0)', 'Futuro'),
]

for rf in rfs:
    r = mt.add_row()
    for i, (cell, text) in enumerate(zip(r.cells, rf)):
        cell.text = text

doc.add_paragraph()

# ─────────────────────────────────────────────────────────
# 7. TECNOLOGÍAS
# ─────────────────────────────────────────────────────────
doc.add_heading('7. Tecnologías Utilizadas', 1)

tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = 'Table Grid'
for i, txt in enumerate(['Capa', 'Tecnología']):
    tech_table.rows[0].cells[i].text = txt
    tech_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(tech_table.rows[0].cells[i], 'D0D8E4')

techs = [
    ('Frontend', 'React + Vite, Vanilla CSS, React Router DOM'),
    ('Backend', 'Node.js + Express.js'),
    ('Base de datos', 'PostgreSQL (Neon — cloud serverless)'),
    ('Autenticación', 'JWT (jsonwebtoken) + bcryptjs'),
    ('Generación QR', 'qrcode (server-side)'),
    ('Testing Backend', 'Jest + Supertest'),
    ('Despliegue Frontend', 'Vercel'),
    ('Despliegue Backend', 'Koyeb'),
    ('Control de versiones', 'Git + GitHub'),
]
for row_data in techs:
    add_table_row(tech_table, row_data)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────
# Guardar
# ─────────────────────────────────────────────────────────
output_path = r'c:/Users/ACER/OneDrive/Desktop/PROGRAMACION/UNIVERSITY PROJECTS/SIstemaGestorGimnasio/Documentacion/definicion_mvp1_actualizado_v2.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
