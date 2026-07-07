from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, cols, bg='1E3A5F', fg='FFFFFF'):
    row = table.rows[0]
    for cell, txt in zip(row.cells, cols):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        set_cell_bg(cell, bg)

def add_technique_block(doc, numero, nombre, subtitulo_proceso, pasos_proceso,
                         subtitulo_hallazgos, hallazgos,
                         subtitulo_resultados, resultados_rf, resultados_rn=None):
    """Agrega un bloque de tecnica de elicitacion con proceso y resultados."""
    doc.add_heading(f'{numero}. {nombre}', 2)

    # Proceso
    doc.add_heading(subtitulo_proceso, 3)
    for paso in pasos_proceso:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(paso)

    # Hallazgos / Resultados
    doc.add_heading(subtitulo_hallazgos, 3)
    for h in hallazgos:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(h)

    # RF/RN asociados
    doc.add_heading(subtitulo_resultados, 3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    header_row(tbl, ['Requisito', 'Descripcion breve', 'Prioridad MoSCoW'], bg='2E4A7F')
    for row_data in resultados_rf:
        r = tbl.add_row()
        for c, t in zip(r.cells, row_data):
            c.text = t

    if resultados_rn:
        doc.add_paragraph()
        doc.add_heading('Reglas de negocio identificadas:', 3)
        for rn in resultados_rn:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(rn)

    doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────────
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('OLYMPUS CORE')
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Sistema Gestor de Gimnasio')
r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x44,0x4A,0x57)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('DESCUBRIMIENTO Y ELICITACION DE REQUISITOS')
r3.bold = True; r3.font.size = Pt(20)
r3.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()
for line in ['Entregable: Iteracion 1 - Elicitacion', 'Grupo: 5', 'Julio 2026']:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run(line).font.size = Pt(12)

doc.add_page_break()

# ─── 1. INTRODUCCION ─────────────────────────────────────────────────────────
doc.add_heading('1. Introduccion', 1)
doc.add_paragraph(
    'El presente documento recoge el proceso completo de descubrimiento y elicitacion de '
    'requisitos para el sistema de gestion integral de gimnasios OLYMPUS CORE. '
    'A traves de la aplicacion de cuatro tecnicas formales de ingenieria de requisitos, '
    'el equipo de desarrollo identifico los actores involucrados, sus necesidades principales '
    'y los dolores operativos del negocio, formalizando un catalogo de 34 Requisitos Funcionales '
    '(RF) y 13 Requisitos No Funcionales (RNF) priorizados mediante la tecnica MoSCoW. '
    '\n\nEste documento sirve como base contractual entre el equipo de desarrollo y el cliente, '
    'garantizando que el sistema construido resuelva los problemas reales del negocio.'
)

# ─── 2. ACTORES ──────────────────────────────────────────────────────────────
doc.add_heading('2. Actores del Sistema', 1)
act_table = doc.add_table(rows=1, cols=4)
act_table.style = 'Table Grid'
header_row(act_table, ['Actor', 'Rol en el sistema', 'Necesidades principales', 'Identificado en'])
actors = [
    ('Administrador',
     'Maximo nivel de acceso. Gestiona la totalidad del sistema y toma decisiones de negocio.',
     'Visualizar metricas clave (ingresos, socios activos, membresías por vencer). Gestionar socios, planes, clases y pagos. Controlar accesos y penalizaciones.',
     'Entrevista + Observacion'),
    ('Recepcionista',
     'Atencion directa al cliente en el mostrador del gimnasio.',
     'Registrar socios rapidamente, cobrar membresías, validar el acceso por QR, gestionar reservas de clases y resolver consultas del socio en el momento.',
     'Entrevista + Observacion'),
    ('Entrenador',
     'Supervisa la actividad fisica de los socios y gestiona sus planes de entrenamiento.',
     'Acceder al listado de socios asignados, ver su ficha medica, gestionar clases y registrar el progreso fisico de los socios.',
     'Entrevista'),
    ('Socio',
     'Cliente del gimnasio. Usuario final del sistema.',
     'Acceder con su QR, ver el estado de su membresía, reservar y cancelar clases, ver su historial de pagos y conocer sus strikes activos.',
     'Brainstorming + Analisis'),
]
for a in actors:
    row = act_table.add_row()
    for c, t in zip(row.cells, a):
        c.text = t

doc.add_paragraph()

# ─── 3. TECNICAS DE ELICITACION ──────────────────────────────────────────────
doc.add_heading('3. Tecnicas de Elicitacion Utilizadas', 1)
doc.add_paragraph(
    'El equipo aplico cuatro tecnicas complementarias de elicitacion. Cada una se describe '
    'a continuacion con su proceso de aplicacion, los hallazgos obtenidos y los requisitos '
    'que se derivaron directamente de ella.'
)

# ── TECNICA 1: ENTREVISTAS ────────────────────────────────────────────────────
add_technique_block(
    doc,
    '3.1', 'Entrevistas con el Cliente',
    'Proceso de aplicacion:',
    [
        'Se realizaron dos sesiones de entrevista semiestructurada con el cliente objetivo: '
        'el dueno y administrador de un gimnasio de mediana escala con aproximadamente 200 socios activos.',
        'Primera sesion (duracion: 60 min): Se identifico la situacion actual del negocio, '
        'los procesos manuales existentes (registro en cuaderno, cobros sin comprobante formal, '
        'control de acceso visual sin herramienta) y las principales frustraciones del cliente.',
        'Segunda sesion (duracion: 45 min): Se validaron los requisitos preliminares levantados '
        'en la primera sesion, se priorizaron funcionalidades y se acordaron las reglas de negocio '
        'criticas (politica de strikes, bloqueo por membresía vencida, gestion de planes).',
        'Las entrevistas fueron registradas con notas estructuradas por modulo tematico: '
        'socios, pagos, acceso, clases y administracion.',
    ],
    'Hallazgos principales:',
    [
        'El cliente no tenia forma de saber cuantos socios tenian la membresía vencida hasta '
        'que intentaban entrar al local — lo que generaba conflictos diarios en recepcion.',
        'Las reservas de clases se hacian por WhatsApp, sin control de aforo ni confirmacion '
        'formal, ocasionando sobreaforo y quejas de socios.',
        'El cobro de membresías era completamente manual: una libreta con fecha de inicio y '
        'monto, sin historial digitalizado ni comprobante imprimible.',
        'El cliente menciono haberse topado con socios que prestaban su carnet a otras personas '
        'para que entraran al gimnasio sin pagar.',
        'Existia un sistema informal de "3 faltas y te quito las clases" que no se aplicaba '
        'de forma consistente por falta de registro.',
        'El dueno aseguro que el Dashboard era la funcionalidad mas critica: "con ver cuanto '
        'entre hoy, quien vence esta semana y si las clases estan llenas, me basta."',
    ],
    'Requisitos derivados de esta tecnica:',
    [
        ('RF-01', 'Login con roles (admin, recepcionista, entrenador, socio)', 'Must Have'),
        ('RF-02', 'Registro de socios con datos personales', 'Must Have'),
        ('RF-05', 'Registro de pagos y generacion automatica de membresía', 'Must Have'),
        ('RF-06', 'Generacion de QR unico por socio para evitar suplantacion', 'Must Have'),
        ('RF-07', 'Validacion de QR con resultado PERMITIDO/DENEGADO en recepcion', 'Must Have'),
        ('RF-13', 'Registro de inasistencias como strikes de penalizacion', 'Should Have'),
        ('RF-14', 'Bloqueo automatico por acumulacion de 3 strikes en 30 dias', 'Should Have'),
        ('RF-15', 'Dashboard con ingresos, socios activos, vencimientos y ocupacion de clases', 'Should Have'),
    ],
    [
        'RN-01: Membresía vencida = acceso fisico bloqueado automaticamente por QR.',
        'RN-04: 3 strikes de inasistencia injustificada en 30 dias activan un bloqueo de 7 dias.',
        'RN-09: El QR contiene un token con fecha de expiracion para evitar uso fraudulento.',
    ]
)

# ── TECNICA 2: OBSERVACION ───────────────────────────────────────────────────
add_technique_block(
    doc,
    '3.2', 'Observacion del Proceso Actual (As-Is)',
    'Proceso de aplicacion:',
    [
        'Se visito el gimnasio en dos momentos operativos distintos: hora pico (7:00 AM - 9:00 AM) '
        'y horario de tarde (6:00 PM - 8:00 PM), para observar el flujo real de trabajo en recepcion.',
        'Se documenta el flujo "as-is" (situacion actual) para cada proceso critico: '
        'ingreso al local, pago de membresía, inscripcion a clase y gestion de recepcion.',
        'Se tomaron notas de los cuellos de botella, errores frecuentes y quejas verbales '
        'de los socios durante la observacion.',
        'Se mapeo el tiempo promedio de cada operacion manual vs. el tiempo esperado con un sistema digital.',
    ],
    'Hallazgos del flujo as-is:',
    [
        'INGRESO: El recepcionista identificaba al socio visualmente o por nombre, buscaba en una '
        'libreta y verificaba manualmente si la fecha de vencimiento era vigente. Tiempo promedio: 45 segundos. '
        'Error observado: 2 socios con membresía vencida ingresaron sin ser detectados en la sesion de observacion.',
        'COBRO: El pago se anotaba en una libreta con nombre, fecha y monto. No habia separacion '
        'entre metodo de pago. Los comprobantes eran notas a mano que frecuentemente se perdian.',
        'CLASES: La lista de asistentes a cada clase era un papel impreso que se llenaba '
        'con lapiz. No habia control de aforo maximo; en una clase se contaron 24 personas '
        'en un espacio autorizado para 15.',
        'INASISTENCIAS: No existia registro formal de faltas. El entrenador recordaba de memoria '
        'quienes faltaban con frecuencia, sin aplicar ninguna consecuencia estandarizada.',
        'TIEMPO TOTAL del proceso manual de ingreso + verificacion: 60-90 segundos. '
        'Objetivo con sistema: menos de 5 segundos (escaneo de QR).',
    ],
    'Requisitos derivados de esta tecnica:',
    [
        ('RF-04', 'Visualizacion clara del estado de membresía (activa/vencida/bloqueada)', 'Must Have'),
        ('RF-07', 'Validacion de QR con tiempo de respuesta optimo en recepcion', 'Must Have'),
        ('RF-08', 'Control de aforo maximo por clase con estado automatico (disponible/llena)', 'Must Have'),
        ('RF-09', 'Reserva y cancelacion de clases con validacion de cupo en tiempo real', 'Must Have'),
        ('RF-12', 'Busqueda rapida de socios por nombre o DNI para recepcion', 'Should Have'),
        ('RF-19', 'Log de accesos con resultado y motivo para auditoria posterior', 'Should Have'),
        ('RF-24', 'Generacion de comprobante de pago en PDF', 'Could Have'),
    ],
    [
        'RN-02: Cada clase tiene un aforo maximo; el sistema no permite reservas adicionales cuando se alcanza.',
        'RN-07: Un socio no puede tener dos reservas activas para la misma clase.',
    ]
)

# ── TECNICA 3: ANALISIS DE SISTEMAS SIMILARES ────────────────────────────────
add_technique_block(
    doc,
    '3.3', 'Analisis de Sistemas Similares (Benchmarking)',
    'Proceso de aplicacion:',
    [
        'Se analizaron tres sistemas de gestion de gimnasios de uso comercial: '
        'GymMaster (Nueva Zelanda), PerfectGym (Polonia) y Mindbody (EE.UU.).',
        'El analisis se enfoco en identificar funcionalidades estandar del sector, '
        'patrones de UI/UX utilizados y flujos de usuario criticos (registro, pago, acceso, clases).',
        'Se identificaron funcionalidades presentes en todos los sistemas analizados '
        '(core obligatorio), funcionalidades presentes en la mayoria (importantes) '
        'y funcionalidades diferenciadas o de premium (opcionales).',
        'Se elaboro una tabla comparativa de caracteristicas para guiar la priorizacion MoSCoW del equipo.',
    ],
    'Hallazgos del benchmarking:',
    [
        'Los tres sistemas tienen como nucleo: autenticacion por roles, gestion de socios, '
        'cobro de membresías, acceso por codigo QR o codigo de barras, y dashboard de metricas. '
        'Esto valido los RF Must Have del proyecto.',
        'La funcionalidad de clases grupales con control de aforo y reserva online '
        'estaba presente en los tres sistemas como caracteristica diferencial respecto a '
        'la gestion manual. Esto elevo su prioridad a Must Have en nuestro catalogo.',
        'El sistema de strikes/penalizaciones no estaba presente en todos los sistemas, '
        'pero fue mencionado como pain point critico por el cliente en las entrevistas, '
        'por lo que se incluyo como Should Have.',
        'Funcionalidades como lista de espera para clases llenas, portal de pago online, '
        'aplicacion movil y notificaciones automaticas eran exclusivas de planes premium '
        'de pago — se documentaron como Could Have y Won\'t Have para versiones futuras.',
        'Los sistemas analizados usaban codigos QR dinamicos (con renovacion periodica) '
        'para mayor seguridad, lo cual inspiro el RNF de expiracion del token QR.',
    ],
    'Requisitos derivados de esta tecnica:',
    [
        ('RF-03', 'Gestion de planes de membresía con precios y duraciones configurables', 'Must Have'),
        ('RF-16', 'Vista Mi Perfil del socio con QR, estado de membresía y historial', 'Should Have'),
        ('RF-23', 'Pantalla dedicada de logs de acceso con filtros avanzados', 'Could Have'),
        ('RF-25', 'Indicador visual de ocupacion de clase para decision de reserva', 'Could Have'),
        ('RF-26', 'Lista de espera para clases con cupo lleno', 'Could Have'),
        ('RF-29', 'Portal de pago online (pasarela de tarjeta)', "Won't Have"),
        ('RF-31', 'Aplicacion movil nativa para socios', "Won't Have"),
    ],
    [
        'RN-05: Catalogo cerrado de planes (el precio lo define el administrador, no el socio).',
        'RN-09: Token QR con fecha de expiracion configurable para mayor seguridad de acceso.',
    ]
)

# ── TECNICA 4: BRAINSTORMING ─────────────────────────────────────────────────
add_technique_block(
    doc,
    '3.4', 'Sesiones de Brainstorming del Equipo',
    'Proceso de aplicacion:',
    [
        'Se realizaron dos sesiones de brainstorming estructurado con los 4 integrantes del equipo '
        'de desarrollo, siguiendo la metodologia de "Think, Pair, Share":',
        'Sesion 1 (60 min) — Generacion libre de ideas: cada integrante propuso funcionalidades '
        'sin restriccion de viabilidad. Se generaron 47 ideas en total.',
        'Sesion 2 (45 min) — Refinamiento y priorizacion: las ideas se agruparon por modulo '
        'tematico y se votaron usando puntos (cada integrante tenia 10 puntos para distribuir). '
        'Las ideas con mayor puntaje se mapearon a los RF formales.',
        'Se cruzo el resultado del brainstorming con los hallazgos de las entrevistas y '
        'la observacion para validar que los requisitos priorizados resolvian necesidades reales.',
        'Las ideas descartadas por viabilidad tecnica o tiempo se documentaron en la '
        "categoria Won't Have para consideracion futura.",
    ],
    'Ideas y decisiones de la sesion:',
    [
        'Idea mas votada: "El socio deberia poder ver su estado completo (membresía, QR, '
        'reservas, strikes) desde su propio perfil" — derivado en RF-16 (Mi Perfil) con prioridad Should Have.',
        'Idea de seguridad: "Que el acceso QR expire automaticamente para que no se pueda '
        'compartir por screenshot" — derivado en el RNF de expiracion de token y RF-34 QR dinamico.',
        'Debate sobre ficha medica: Se propuso incluirla en MVP1, pero se decidio diferir '
        'a v1.1 por dependencia con el modulo de rutinas y el mayor tiempo de desarrollo requerido. '
        'Documentada en RF-20 (Could Have).',
        'Propuesta de modulo de inventario: Descartada por estar fuera del dominio de negocio '
        'principal del proyecto. Documentada como RF-32 (Won\'t Have).',
        'Propuesta de notificaciones email/SMS: El equipo evaluo la complejidad de integracion '
        'con servicios externos (SendGrid, Twilio) y la costo asociado, decidiendose diferir '
        'para v2.0. Documentada en RF-28 (Won\'t Have).',
        'Se definio la arquitectura final: backend desacoplado (API REST con Node.js/Express) '
        '+ frontend SPA (React + Vite), con PostgreSQL en la nube (Neon).',
    ],
    'Requisitos derivados de esta tecnica:',
    [
        ('RF-10', 'Bloqueo de cuenta por 5 intentos fallidos de login (30 min)', 'Must Have'),
        ('RF-11', 'Control de acceso basado en roles RBAC en cada endpoint', 'Must Have'),
        ('RF-17', 'Justificacion y eliminacion de strikes por el administrador', 'Should Have'),
        ('RF-18', 'Historial de pagos accesible por administrador, recepcionista y socio', 'Should Have'),
        ('RF-27', 'Asignacion de entrenador responsable por socio', 'Could Have'),
        ('RF-30', 'Exportacion de reportes en Excel/PDF', "Won't Have"),
        ('RF-33', 'Integracion con torniquete electronico via API', "Won't Have"),
    ],
    [
        'RN-03: Cancelacion de reserva permitida solo hasta 2 horas antes del inicio de la clase.',
        'RN-08: El dashboard debe mostrar obligatoriamente: ingresos del dia, membresías por vencer y ocupacion de clases.',
        'RN-10: Contrasenas almacenadas con bcrypt (salt >= 10 rondas), nunca en texto plano.',
    ]
)

# ─── 4. REQUISITOS FUNCIONALES ───────────────────────────────────────────────
doc.add_heading('4. Catalogo Completo de Requisitos Funcionales (RF)', 1)
doc.add_paragraph(
    'Los siguientes 34 requisitos funcionales fueron identificados y formalizados a partir de '
    'las cuatro tecnicas de elicitacion aplicadas. Se presenta el catalogo completo con su '
    'identificador, descripcion, actor principal, tecnica de origen y prioridad MoSCoW.'
)

doc.add_heading('4.1 Must Have (Obligatorios) — 11 RF', 2)
must_table = doc.add_table(rows=1, cols=5)
must_table.style = 'Table Grid'
header_row(must_table, ['ID', 'Descripcion', 'Actor', 'Tecnica de origen', 'MVP1'])
must_rfs = [
    ('RF-01', 'Inicio de sesion con email/contrasena y autenticacion por roles', 'Todos', 'Entrevista', 'Si'),
    ('RF-02', 'Registro de socios con nombre, apellido, DNI, email y telefono', 'Admin / Recep.', 'Entrevista + Observacion', 'Si'),
    ('RF-03', 'Gestion de planes de membresía (nombre, precio, duracion)', 'Administrador', 'Benchmarking', 'Si'),
    ('RF-04', 'Visualizacion del estado de membresía (activa/vencida/bloqueada)', 'Admin / Recep.', 'Observacion', 'Si'),
    ('RF-05', 'Registro de pagos con metodo y generacion automatica de membresía', 'Admin / Recep.', 'Entrevista', 'Si'),
    ('RF-06', 'Generacion de codigo QR unico por socio con token de acceso', 'Sistema', 'Entrevista + Benchmarking', 'Si'),
    ('RF-07', 'Validacion de QR en recepcion con resultado PERMITIDO/DENEGADO', 'Recepcionista', 'Entrevista + Observacion', 'Si'),
    ('RF-08', 'CRUD de clases grupales con tipo, instructor, horario y aforo maximo', 'Administrador', 'Observacion', 'Si'),
    ('RF-09', 'Reserva y cancelacion de clases por parte del socio', 'Socio / Recep.', 'Observacion + Benchmarking', 'Si'),
    ('RF-10', 'Bloqueo de cuenta por 5 intentos de login fallidos (30 min)', 'Sistema', 'Brainstorming', 'Si'),
    ('RF-11', 'Control de acceso basado en roles (RBAC) para cada endpoint y modulo', 'Sistema', 'Brainstorming', 'Si'),
]
for rf in must_rfs:
    r = must_table.add_row()
    for c, t in zip(r.cells, rf): c.text = t

doc.add_paragraph()
doc.add_heading('4.2 Should Have (Importantes) — 8 RF', 2)
should_table = doc.add_table(rows=1, cols=5)
should_table.style = 'Table Grid'
header_row(should_table, ['ID', 'Descripcion', 'Actor', 'Tecnica de origen', 'MVP1'])
should_rfs = [
    ('RF-12', 'Busqueda y filtrado de socios por nombre, DNI o estado', 'Admin / Recep.', 'Observacion', 'Si'),
    ('RF-13', 'Registro de inasistencias a clases reservadas como strikes', 'Recep. / Sistema', 'Entrevista', 'Si'),
    ('RF-14', 'Bloqueo automatico de reservas por 3 strikes en 30 dias', 'Sistema', 'Entrevista', 'Si'),
    ('RF-15', 'Dashboard administrativo con ingresos, socios activos, vencimientos y ocupacion', 'Administrador', 'Entrevista', 'Si'),
    ('RF-16', 'Vista Mi Perfil del socio: QR, estado de membresía, strikes y reservas', 'Socio', 'Brainstorming + Benchmarking', 'Si'),
    ('RF-17', 'Justificacion y eliminacion de strikes por el administrador con registro', 'Administrador', 'Brainstorming', 'Si'),
    ('RF-18', 'Historial de pagos por socio accesible para admin, recep. y el socio', 'Todos', 'Brainstorming', 'Si'),
    ('RF-19', 'Log de accesos QR con resultado y motivo, filtrable por socio y fecha', 'Admin / Recep.', 'Observacion', 'Parcial'),
]
for rf in should_rfs:
    r = should_table.add_row()
    for c, t in zip(r.cells, rf): c.text = t

doc.add_paragraph()
doc.add_heading('4.3 Could Have (Deseables) — 8 RF', 2)
could_table = doc.add_table(rows=1, cols=5)
could_table.style = 'Table Grid'
header_row(could_table, ['ID', 'Descripcion', 'Actor', 'Tecnica de origen', 'MVP1'])
could_rfs = [
    ('RF-20', 'Ficha medica digital: lesiones, enfermedades cardiovasculares y objetivo', 'Entrenador / Socio', 'Entrevista + Brainstorming', 'No (v1.1)'),
    ('RF-21', 'Asignacion de rutinas de entrenamiento personalizadas a socios', 'Entrenador', 'Entrevista', 'No (v1.1)'),
    ('RF-22', 'Registro de progreso fisico por ejercicio (peso, reps, fecha)', 'Entrenador / Socio', 'Entrevista', 'No (v1.1)'),
    ('RF-23', 'Pantalla dedicada de logs de acceso con filtros avanzados', 'Admin / Recep.', 'Benchmarking', 'No (v1.1)'),
    ('RF-24', 'Generacion de comprobante de pago en PDF al momento del cobro', 'Admin / Recep.', 'Observacion', 'No (v1.1)'),
    ('RF-25', 'Indicador visual de ocupacion de clase para el socio antes de reservar', 'Socio', 'Benchmarking', 'No (v1.1)'),
    ('RF-26', 'Lista de espera para clases con cupo lleno; notificacion al liberarse lugar', 'Socio / Sistema', 'Benchmarking', 'No (v1.1)'),
    ('RF-27', 'Asignacion de entrenador responsable a cada socio de forma permanente', 'Administrador', 'Brainstorming', 'No (v1.1)'),
]
for rf in could_rfs:
    r = could_table.add_row()
    for c, t in zip(r.cells, rf): c.text = t

doc.add_paragraph()
doc.add_heading("4.4 Won't Have (Fuera de alcance actual) — 7 RF", 2)
wont_table = doc.add_table(rows=1, cols=4)
wont_table.style = 'Table Grid'
header_row(wont_table, ['ID', 'Descripcion', 'Razon de exclusion', 'Version futura'])
wont_rfs = [
    ('RF-28', 'Notificaciones automaticas por email/SMS de vencimiento de membresía', 'Requiere integracion con servicio externo (SendGrid/Twilio) y costo operativo', 'v2.0'),
    ('RF-29', 'Portal de autopago en linea (tarjeta credito/debito)', 'Requiere pasarela de pago y cumplimiento PCI-DSS', 'v2.0'),
    ('RF-30', 'Reportes exportables en Excel/PDF con filtros por rango de fechas', 'Requiere libreria de reportes; mayor tiempo de desarrollo', 'v2.0'),
    ('RF-31', 'Aplicacion movil nativa (Android/iOS) para socios', 'Requiere desarrollo movil separado; fuera del alcance academico', 'v3.0'),
    ('RF-32', 'Modulo de inventario de equipamiento del gimnasio', 'Fuera del dominio principal del negocio para este ciclo', 'v3.0'),
    ('RF-33', 'Integracion con torniquete electronico real via API/IoT', 'Requiere hardware fisico; fuera del alcance del proyecto', 'v3.0'),
    ('RF-34', 'Codigo QR dinamico con renovacion automatica cada N minutos', 'Complejidad de sincronizacion en tiempo real con hardware', 'v3.0'),
]
for rf in wont_rfs:
    r = wont_table.add_row()
    for c, t in zip(r.cells, rf): c.text = t

doc.add_paragraph()

# ─── 5. RESUMEN MOSOW ────────────────────────────────────────────────────────
doc.add_heading('5. Resumen de Priorizacion MoSCoW', 1)
summary = doc.add_table(rows=1, cols=5)
summary.style = 'Table Grid'
header_row(summary, ['Categoria', 'Cant. RF', 'Tecnica principal de origen', 'Implementado MVP1', 'Pendiente'])
s_data = [
    ('Must Have',   '11', 'Entrevista + Observacion', '11 / 11', '0'),
    ('Should Have', '8',  'Entrevista + Brainstorming', '7 / 8 (1 parcial)', '1'),
    ('Could Have',  '8',  'Benchmarking + Brainstorming', '0 / 8', '8 (v1.1)'),
    ("Won't Have",  '7',  'Brainstorming', '0 / 7', '7 (v2.0/v3.0)'),
    ('TOTAL',       '34', '—', '18+', '16'),
]
for row_data in s_data:
    r = summary.add_row()
    for c, t in zip(r.cells, row_data): c.text = t
    if row_data[0] == 'TOTAL':
        for c in r.cells:
            c.paragraphs[0].runs[0].bold = True
            set_cell_bg(c, 'D0D8E4')

doc.add_paragraph()

# ─── 6. REQUISITOS NO FUNCIONALES ───────────────────────────────────────────
doc.add_heading('6. Requisitos No Funcionales (RNF)', 1)
doc.add_paragraph(
    'Los siguientes 13 requisitos no funcionales garantizan la calidad, seguridad, '
    'rendimiento y mantenibilidad del sistema. Fueron identificados a partir de los '
    'hallazgos de las tecnicas de elicitacion, el analisis de sistemas similares y '
    'las mejores practicas de ingenieria de software aplicadas al dominio.'
)

rnf_sections = [
    ('6.1 Rendimiento', [
        ('RNF-01', 'El tiempo de respuesta de los endpoints de la API REST no debe superar 500 ms bajo carga normal (hasta 50 usuarios concurrentes)', 'Rendimiento', 'Benchmarking'),
        ('RNF-02', 'La validacion del codigo QR en el endpoint /acceso/validar debe completarse en menos de 300 ms para garantizar fluidez en la entrada al gimnasio', 'Rendimiento', 'Observacion'),
        ('RNF-03', 'El sistema debe soportar al menos 100 usuarios concurrentes sin degradacion perceptible del servicio durante las horas pico (7AM-9AM y 6PM-8PM)', 'Rendimiento', 'Observacion'),
    ]),
    ('6.2 Seguridad', [
        ('RNF-04', 'Las contrasenas deben almacenarse exclusivamente en formato hash usando bcrypt con un factor de costo (salt rounds) minimo de 10', 'Seguridad', 'Brainstorming'),
        ('RNF-05', 'Toda la comunicacion entre el frontend y el backend debe realizarse exclusivamente a traves de HTTPS con certificado TLS valido', 'Seguridad', 'Brainstorming + Benchmarking'),
        ('RNF-06', 'Los tokens JWT de autenticacion deben tener una duracion maxima de 24 horas y los tokens QR una duracion configurable (por defecto 12 horas)', 'Seguridad', 'Benchmarking'),
        ('RNF-07', 'La API debe validar y sanitizar todos los datos de entrada para prevenir inyeccion SQL y ataques XSS en cada endpoint', 'Seguridad', 'Brainstorming'),
    ]),
    ('6.3 Usabilidad y Accesibilidad', [
        ('RNF-08', 'La interfaz de usuario debe ser completamente funcional en los navegadores modernos mas utilizados: Chrome, Firefox, Edge y Safari (ultimas 2 versiones)', 'Usabilidad', 'Entrevista'),
        ('RNF-09', 'El sistema debe ser responsive y operable correctamente en pantallas a partir de 768 px de ancho (tablets y laptops)', 'Usabilidad', 'Benchmarking'),
        ('RNF-10', 'Las acciones criticas del flujo de recepcion (validar QR, registrar pago) deben completarse con un maximo de 3 clics o interacciones desde la pantalla principal', 'Usabilidad', 'Observacion'),
    ]),
    ('6.4 Disponibilidad y Mantenibilidad', [
        ('RNF-11', 'El sistema debe mantener una disponibilidad minima del 99% durante el horario de operacion del gimnasio (6:00 AM - 10:00 PM, de lunes a sabado)', 'Disponibilidad', 'Entrevista'),
        ('RNF-12', 'El codigo fuente debe estar versionado en Git con una rama por modulo de desarrollo y etiquetas de version para cada entrega formal', 'Mantenibilidad', 'Brainstorming'),
        ('RNF-13', 'La base de datos debe incluir indices en las columnas de busqueda y filtrado frecuente: dni, socio_id, estado, fecha_fin y fecha_hora de clases', 'Rendimiento / DB', 'Benchmarking + Observacion'),
    ]),
]

for section_title, rnfs in rnf_sections:
    doc.add_heading(section_title, 2)
    rnf_table = doc.add_table(rows=1, cols=4)
    rnf_table.style = 'Table Grid'
    header_row(rnf_table, ['ID', 'Descripcion', 'Categoria', 'Tecnica de origen'], bg='2E5090')
    for rnf in rnfs:
        r = rnf_table.add_row()
        for c, t in zip(r.cells, rnf): c.text = t
    doc.add_paragraph()

# ─── 7. VALIDACION ──────────────────────────────────────────────────────────
doc.add_heading('7. Validacion de Requisitos con el Cliente', 1)
doc.add_paragraph(
    'Tras la consolidacion del catalogo de requisitos, se realizo una sesion de validacion '
    'con el cliente para confirmar que los RF y RNF identificados representan correctamente '
    'sus necesidades. Los resultados de la validacion fueron:'
)
for v in [
    'El cliente confirmo que los 11 RF "Must Have" cubren completamente los problemas operativos criticos identificados en las entrevistas.',
    'Se acordo explicitamente que las funcionalidades de ficha medica, rutinas y progreso fisico (RF-20 a RF-22) son importantes pero pueden esperar a una segunda entrega.',
    'El cliente solicito expresamente que el sistema funcione desde cualquier computadora sin instalacion adicional — confirmando el RNF-08 de compatibilidad de navegadores.',
    'Se valido la politica de 3 strikes en 30 dias como regla de negocio oficial del gimnasio (RN-04).',
    'El cliente aprobó el catalogo de requisitos el dia 12 de mayo de 2026, dando inicio formal al proceso de desarrollo.',
]:
    doc.add_paragraph(v, style='List Bullet')

doc.add_paragraph()

# ─── Guardar ─────────────────────────────────────────────────────────────────
output_path = r'c:/Users/ACER/OneDrive/Desktop/PROGRAMACION/UNIVERSITY PROJECTS/SIstemaGestorGimnasio/Documentacion/DESCUBRIMIENTO_Y_ELICITACION_actualizado.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
