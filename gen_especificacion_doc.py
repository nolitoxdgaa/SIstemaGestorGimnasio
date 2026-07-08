from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Utilidades ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, cols, bg='1E3A5F', fg='FFFFFF'):
    row = table.rows[0]
    for cell, txt in zip(row.cells, cols):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        set_cell_bg(cell, bg)

def add_rf_block(doc, rf_id, nombre, actores, prioridad, modulo,
                 descripcion, precondiciones, flujo_normal,
                 flujos_alt, postcondiciones, reglas_negocio, endpoints=None):
    """Genera el bloque de especificacion formal de un RF."""
    doc.add_heading(f'{rf_id} — {nombre}', 3)

    # Tabla de cabecera
    t = doc.add_table(rows=6, cols=2)
    t.style = 'Table Grid'
    campos = [
        ('Identificador', rf_id),
        ('Nombre', nombre),
        ('Actores involucrados', actores),
        ('Prioridad MoSCoW', prioridad),
        ('Modulo del sistema', modulo),
        ('Descripcion', descripcion),
    ]
    for i, (key, val) in enumerate(campos):
        t.rows[i].cells[0].text = key
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(t.rows[i].cells[0], 'D0D8E4')
        t.rows[i].cells[1].text = val

    doc.add_paragraph()

    # Precondiciones
    p = doc.add_paragraph()
    p.add_run('Precondiciones:').bold = True
    for pre in precondiciones:
        doc.add_paragraph(pre, style='List Bullet')

    # Flujo normal
    p2 = doc.add_paragraph()
    p2.add_run('Flujo normal:').bold = True
    for i, paso in enumerate(flujo_normal, 1):
        doc.add_paragraph(f'{i}. {paso}')

    # Flujos alternativos / excepciones
    if flujos_alt:
        p3 = doc.add_paragraph()
        p3.add_run('Flujos alternativos / Excepciones:').bold = True
        for alt in flujos_alt:
            doc.add_paragraph(alt, style='List Bullet')

    # Postcondiciones
    p4 = doc.add_paragraph()
    p4.add_run('Postcondiciones:').bold = True
    for post in postcondiciones:
        doc.add_paragraph(post, style='List Bullet')

    # Reglas de negocio aplicadas
    if reglas_negocio:
        p5 = doc.add_paragraph()
        p5.add_run('Reglas de negocio aplicadas:').bold = True
        for rn in reglas_negocio:
            doc.add_paragraph(rn, style='List Bullet')

    # Endpoints REST
    if endpoints:
        p6 = doc.add_paragraph()
        p6.add_run('Endpoints REST asociados:').bold = True
        for ep in endpoints:
            doc.add_paragraph(ep, style='List Bullet')

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── PORTADA ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('OLYMPUS CORE')
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Sistema Gestor de Gimnasio')
r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x44, 0x4A, 0x57)

doc.add_paragraph()
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('ESPECIFICACION FORMAL DE REQUISITOS')
r3.bold = True; r3.font.size = Pt(20)
r3.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()
for line in ['Entregable: Iteracion 2 - Especificacion', 'Grupo: 5', 'Julio 2026']:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run(line).font.size = Pt(12)

doc.add_page_break()

# ── 1. INTRODUCCION ──────────────────────────────────────────────────────────
doc.add_heading('1. Introduccion', 1)
doc.add_paragraph(
    'Este documento presenta la Especificacion Formal de Requisitos del sistema OLYMPUS CORE. '
    'A diferencia de los documentos de elicitacion y priorizacion (entregables anteriores), '
    'el presente se centra exclusivamente en la descripcion tecnica y precisa del comportamiento '
    'del sistema para cada requisito funcional identificado.\n\n'
    'Cada requisito se especifica mediante una ficha estructurada que incluye: '
    'actores, precondiciones, flujo normal de ejecucion, flujos alternativos y de excepcion, '
    'postcondiciones, reglas de negocio aplicadas y los endpoints REST asociados a su implementacion. '
    'Adicionalmente, se incluye la especificacion de los 13 Requisitos No Funcionales con '
    'sus criterios de verificacion y aceptacion.'
)

# ── 2. MODELO DE DATOS ───────────────────────────────────────────────────────
doc.add_heading('2. Modelo de Datos del Sistema', 1)
doc.add_paragraph(
    'El sistema OLYMPUS CORE opera sobre las siguientes entidades de dominio principales. '
    'Cada entidad se especifica con sus atributos clave y relaciones.'
)

entidades = [
    ('usuarios', ['id (PK)', 'nombre', 'email (UNIQUE)', 'password_hash', 'rol (administrador|recepcionista|entrenador|socio)', 'creado_en']),
    ('socios', ['id (PK)', 'nombre', 'apellido', 'dni (UNIQUE, 8 dig.)', 'email (UNIQUE)', 'telefono', 'fecha_nacimiento', 'estado (activo|inactivo)', 'fecha_registro']),
    ('planes', ['id (PK)', 'nombre', 'descripcion', 'duracion_dias', 'precio', 'activo']),
    ('membresias', ['id (PK)', 'socio_id (FK)', 'plan_id (FK)', 'fecha_inicio', 'fecha_fin', 'estado (activa|vencida|bloqueada)', 'precio']),
    ('pagos', ['id (PK)', 'socio_id (FK)', 'plan_id (FK)', 'monto', 'metodo_pago (efectivo|yape|plin|tarjeta)', 'estado (completado|anulado)', 'creado_en']),
    ('clases_grupales', ['id (PK)', 'tipo (spinning|crossfit|yoga|zumba)', 'nombre', 'instructor', 'fecha_hora (TIMESTAMP)', 'duracion_minutos', 'aforo_maximo', 'estado (disponible|llena|cancelada)']),
    ('reservas', ['id (PK)', 'socio_id (FK)', 'clase_id (FK)', 'estado (confirmada|cancelada|asistio|inasistencia)', 'creado_en']),
    ('penalizaciones', ['id (PK)', 'socio_id (FK)', 'reserva_id (FK)', 'justificada (BOOL)', 'justificacion', 'creado_en']),
    ('penalizaciones_bloqueo', ['id (PK)', 'socio_id (FK)', 'bloqueado_hasta (TIMESTAMP)', 'creado_en']),
    ('logs_acceso', ['id (PK)', 'socio_id (FK)', 'resultado (PERMITIDO|DENEGADO)', 'motivo', 'registrado_en']),
]

for entidad, campos in entidades:
    doc.add_heading(entidad, 3)
    for campo in campos:
        doc.add_paragraph(campo, style='List Bullet')
    doc.add_paragraph()

# ── 3. ESPECIFICACION FORMAL POR REQUISITO FUNCIONAL ─────────────────────────
doc.add_heading('3. Especificacion Formal por Requisito Funcional', 1)
doc.add_paragraph(
    'A continuacion se presenta la especificacion formal de cada uno de los 19 Requisitos '
    'Funcionales implementados en el MVP1 (11 Must Have + 7 Should Have + 1 Should Have parcial). '
    'Los RF de categorias Could Have y Won\'t Have se documentan en la seccion 4.'
)

# ────────────────────────────────────────────────────────
doc.add_heading('3.1 Must Have — Requisitos Obligatorios (RF-01 a RF-11)', 2)

add_rf_block(
    doc, 'RF-01', 'Autenticacion por Roles',
    actores='Administrador, Recepcionista, Entrenador, Socio',
    prioridad='Must Have',
    modulo='Autenticacion',
    descripcion='El sistema permite que cualquier usuario registrado inicie sesion con su email y contrasena, recibiendo un token JWT con su identidad y rol para acceder a los recursos autorizados.',
    precondiciones=[
        'El usuario debe tener una cuenta registrada en la tabla usuarios.',
        'El servicio de base de datos debe estar disponible.',
    ],
    flujo_normal=[
        'El actor accede a la pantalla de login en la URL raiz del sistema.',
        'Ingresa su email y contrasena en el formulario.',
        'El frontend envia una peticion POST /auth/login con las credenciales.',
        'El backend verifica que el email exista en la tabla usuarios.',
        'El backend compara la contrasena ingresada con el hash almacenado usando bcrypt.',
        'Si las credenciales son correctas, el sistema genera un JWT firmado con duracion de 24 horas que incluye: id, nombre, email y rol del usuario.',
        'El frontend almacena el token y redirige al modulo correspondiente segun el rol (Dashboard para admin/recep, Clases para entrenador, Mi Perfil para socio).',
    ],
    flujos_alt=[
        'FA-01a: Credenciales incorrectas — El sistema responde con error 401. Se incrementa el contador de intentos fallidos.',
        'FA-01b: Cuenta bloqueada (5 intentos) — El sistema responde con error 429 e informa el tiempo de desbloqueo (30 min).',
        'FA-01c: Email no registrado — El sistema responde con error 401 generico (no revela si el email existe).',
    ],
    postcondiciones=[
        'El usuario tiene un token JWT valido para acceder a los endpoints protegidos.',
        'El contador de intentos fallidos se resetea a cero tras un login exitoso.',
    ],
    reglas_negocio=['RN-10: Contrasenas almacenadas con bcrypt, nunca en texto plano.'],
    endpoints=['POST /api/v1/auth/login', 'POST /api/v1/auth/logout']
)

add_rf_block(
    doc, 'RF-02', 'Gestion de Socios (CRUD)',
    actores='Administrador, Recepcionista',
    prioridad='Must Have',
    modulo='Socios',
    descripcion='El sistema permite crear, visualizar, actualizar y desactivar socios. Incluye validaciones de unicidad de DNI y email, y crea automaticamente una cuenta de usuario para el socio registrado.',
    precondiciones=[
        'El actor debe estar autenticado con rol administrador o recepcionista.',
        'Para crear: el DNI de 8 digitos y el email no deben existir previamente en el sistema.',
    ],
    flujo_normal=[
        'El actor accede al modulo Socios desde el menu lateral.',
        'Visualiza el listado paginado con su membresia y strikes.',
        'Para registrar: hace clic en "Nuevo Socio", completa el formulario (nombre, apellido, DNI, email, telefono, fecha de nacimiento) y confirma.',
        'El backend valida los campos requeridos, el formato del DNI (8 digitos), el formato de email y el telefono (9 digitos, empieza con 9).',
        'Si todo es valido, inserta el registro en socios y crea la cuenta correspondiente en usuarios con contrasena por defecto "Demo1234".',
        'El sistema genera automaticamente el codigo QR del nuevo socio.',
        'Para editar: el actor selecciona un socio y modifica los campos permitidos.',
        'Para desactivar: el actor presiona el boton de eliminar; el estado del socio cambia a "inactivo" (borrado logico).',
    ],
    flujos_alt=[
        'FA-02a: DNI duplicado — El sistema responde con error 409 "Ya existe un socio con ese DNI".',
        'FA-02b: Email duplicado — El sistema responde con error 409 "Ya existe un usuario con ese email".',
        'FA-02c: Campos requeridos faltantes — Error 400 con detalle de los campos invalidos.',
    ],
    postcondiciones=[
        'El nuevo socio queda registrado con estado "activo" y una cuenta de acceso creada.',
        'El socio desactivado permanece en la BD con estado "inactivo" (trazabilidad).',
    ],
    reglas_negocio=['RN: El borrado de socios es logico, nunca fisico.'],
    endpoints=['GET /api/v1/socios', 'GET /api/v1/socios/:id', 'POST /api/v1/socios', 'PUT /api/v1/socios/:id', 'DELETE /api/v1/socios/:id']
)

add_rf_block(
    doc, 'RF-03', 'Gestion de Planes de Membresia',
    actores='Administrador',
    prioridad='Must Have',
    modulo='Planes',
    descripcion='El sistema mantiene un catalogo de planes de membresia con nombre, descripcion, duracion en dias y precio. El administrador puede consultar los planes disponibles para asignarlos durante el registro de pagos.',
    precondiciones=['El actor debe estar autenticado con rol administrador.'],
    flujo_normal=[
        'El administrador accede al listado de planes disponibles (GET /planes).',
        'Los planes activos se muestran con su nombre, descripcion, duracion y precio.',
        'El administrador puede seleccionar un plan al registrar un pago para un socio.',
    ],
    flujos_alt=['FA-03a: No hay planes activos — El sistema muestra una lista vacia y el administrador debe crear planes desde la base de datos.'],
    postcondiciones=['Los planes quedan disponibles para ser asignados en el registro de pagos.'],
    reglas_negocio=['RN-05: Catalogo cerrado de precios; el precio lo define el administrador, no el socio.'],
    endpoints=['GET /api/v1/planes']
)

add_rf_block(
    doc, 'RF-04', 'Visualizacion del Estado de Membresia',
    actores='Administrador, Recepcionista, Socio',
    prioridad='Must Have',
    modulo='Membresias',
    descripcion='El sistema muestra en tiempo real el estado de la membresia de cada socio (activa, vencida o bloqueada) junto con la fecha de vencimiento y los dias restantes. Se incluye alertas visuales para membresías proximas a vencer (menos de 3 dias).',
    precondiciones=['El actor debe estar autenticado.'],
    flujo_normal=[
        'El actor consulta el listado de membresías o el perfil de un socio especifico.',
        'El sistema realiza una consulta JOIN entre socios, membresías y planes para obtener la membresia mas reciente de cada socio.',
        'Se calcula automaticamente si la membresia esta activa (fecha_fin >= HOY), vencida (fecha_fin < HOY) o bloqueada (estado = bloqueada por strikes).',
        'Se muestran los dias restantes y se aplica un color de alerta (amarillo) si vence en 3 dias o menos.',
        'El modulo de Membresías permite filtrar por estado: activas, por vencer y vencidas.',
    ],
    flujos_alt=['FA-04a: Socio sin membresia — El sistema muestra "Sin membresia" en la tarjeta del socio.'],
    postcondiciones=['El actor obtiene una vision clara del estado actual de cada membresia.'],
    reglas_negocio=['RN-01: Membresia vencida o bloqueada impide el acceso fisico al gimnasio.'],
    endpoints=['GET /api/v1/membresias', 'GET /api/v1/membresias/:id', 'GET /api/v1/socios/:id']
)

add_rf_block(
    doc, 'RF-05', 'Registro de Pagos y Activacion de Membresia',
    actores='Administrador, Recepcionista',
    prioridad='Must Have',
    modulo='Pagos',
    descripcion='El sistema registra el pago de una membresia y activa automaticamente la membresia del socio por la duracion correspondiente al plan seleccionado. Soporta cuatro metodos de pago: efectivo, Yape, Plin y tarjeta.',
    precondiciones=[
        'El actor debe estar autenticado con rol administrador o recepcionista.',
        'El socio debe estar registrado y activo en el sistema.',
        'El plan seleccionado debe existir y estar activo.',
    ],
    flujo_normal=[
        'El actor accede al modulo Pagos o al perfil del socio.',
        'Selecciona el socio y el plan al que desea suscribirlo.',
        'Elige el metodo de pago (efectivo, Yape, Plin o tarjeta) e ingresa el monto.',
        'El sistema inserta el registro en la tabla pagos.',
        'El sistema crea o actualiza la membresia del socio: calcula fecha_inicio (hoy) y fecha_fin (hoy + duracion_dias del plan) y establece estado = "activa".',
        'Se muestra confirmacion del pago y el nuevo estado de la membresia.',
    ],
    flujos_alt=[
        'FA-05a: Monto no coincide con el precio del plan — El sistema muestra advertencia pero permite continuar (el admin puede aplicar descuentos manuales).',
        'FA-05b: Socio inactivo — El sistema no permite registrar pagos para socios desactivados.',
    ],
    postcondiciones=[
        'El pago queda registrado en el historial con fecha y metodo.',
        'La membresia del socio pasa a estado "activa" con la nueva fecha de vencimiento.',
        'El socio puede ahora validar su acceso por QR.',
    ],
    reglas_negocio=['RN: El registro de pago es el unico mecanismo para activar una membresia.'],
    endpoints=['POST /api/v1/pagos', 'GET /api/v1/pagos', 'GET /api/v1/socios/:id/pagos']
)

add_rf_block(
    doc, 'RF-06', 'Generacion de Codigo QR por Socio',
    actores='Sistema (automatico), Administrador, Recepcionista',
    prioridad='Must Have',
    modulo='Acceso QR',
    descripcion='El sistema genera un codigo QR unico por socio que contiene un token JWT firmado con el identificador del socio. El QR se genera en formato base64 embebible en la interfaz y tiene una duracion configurable.',
    precondiciones=['El socio debe estar registrado en el sistema.'],
    flujo_normal=[
        'El actor solicita el QR del socio (GET /socios/:id/qr) o el socio accede a Mi Perfil.',
        'El backend genera un JWT de corto plazo (12 horas por defecto) con el payload { socioId: id }.',
        'Se genera la imagen del QR en formato data:image/png;base64 usando la libreria qrcode.',
        'El QR se muestra en pantalla y puede descargarse o imprimirse.',
        'El QR generado automaticamente al crear un socio nuevo se devuelve en la respuesta del POST /socios.',
    ],
    flujos_alt=['FA-06a: Socio no encontrado — Error 404.'],
    postcondiciones=['El socio dispone de un QR valido que puede presentar en recepcion para acceder al gimnasio.'],
    reglas_negocio=['RN-09: El token del QR tiene fecha de expiracion para evitar su uso fraudulento mediante capturas de pantalla.'],
    endpoints=['GET /api/v1/socios/:id/qr']
)

add_rf_block(
    doc, 'RF-07', 'Validacion de Acceso por QR',
    actores='Recepcionista (o sistema de torniquete)',
    prioridad='Must Have',
    modulo='Acceso QR',
    descripcion='El sistema valida el codigo QR presentado por el socio en recepcion y determina si el acceso debe ser PERMITIDO o DENEGADO, registrando el resultado y el motivo en el log de accesos.',
    precondiciones=['El actor debe estar autenticado con rol recepcionista o administrador.'],
    flujo_normal=[
        'El recepcionista abre el modulo Acceso QR en la pantalla de recepcion.',
        'El socio presenta su QR (desde su perfil en el movil o en papel).',
        'El recepcionista escanea o ingresa el token y envia la peticion POST /acceso/validar.',
        'El backend valida la firma y expiracion del JWT del QR.',
        'Verifica que el socio tenga membresia activa (fecha_fin >= HOY y estado != "vencida").',
        'Verifica que el socio no tenga un bloqueo activo por strikes (tabla penalizaciones_bloqueo).',
        'Si todas las validaciones pasan: resultado = "PERMITIDO". Se muestra pantalla verde.',
        'El sistema registra el acceso en logs_acceso con resultado y timestamp.',
    ],
    flujos_alt=[
        'FA-07a: Token QR expirado — DENEGADO, motivo: QR_EXPIRADO.',
        'FA-07b: Membresia vencida — DENEGADO, motivo: MEMBRESIA_VENCIDA. Se muestra la fecha de vencimiento.',
        'FA-07c: Penalizacion activa (3 strikes) — DENEGADO, motivo: PENALIZACION_ACTIVA. Se muestra la fecha de desbloqueo.',
        'FA-07d: Socio inactivo — DENEGADO, motivo: SOCIO_INACTIVO.',
    ],
    postcondiciones=[
        'El acceso queda registrado en logs_acceso independientemente del resultado.',
        'El recepcionista tiene confirmacion visual inmediata (verde = permitido, rojo = denegado).',
    ],
    reglas_negocio=['RN-01: Membresia vencida bloquea el acceso.', 'RN-04: 3 strikes activos bloquean el acceso.'],
    endpoints=['POST /api/v1/acceso/validar', 'GET /api/v1/acceso/logs']
)

add_rf_block(
    doc, 'RF-08', 'Gestion de Clases Grupales',
    actores='Administrador',
    prioridad='Must Have',
    modulo='Clases',
    descripcion='El sistema permite crear, editar y cancelar clases grupales. Cada clase tiene un tipo predefinido, instructor, horario, duracion y aforo maximo. El estado de la clase (disponible/llena/cancelada) se actualiza automaticamente segun las reservas.',
    precondiciones=['El actor debe estar autenticado con rol administrador.'],
    flujo_normal=[
        'El administrador accede al modulo Clases.',
        'Para crear: completa el formulario con tipo, nombre, instructor, fecha-hora, duracion y aforo maximo.',
        'El sistema inserta la clase con estado inicial "disponible".',
        'Para editar: selecciona la clase y modifica los campos permitidos.',
        'Para cancelar: el estado de la clase cambia a "cancelada" y se notifica a los socios con reservas confirmadas (en v2.0).',
        'El listado de clases permite filtrar por tipo, fecha y disponibilidad.',
        'Los socios y recepcionistas pueden ver el listado con el aforo disponible en tiempo real.',
    ],
    flujos_alt=[
        'FA-08a: Fecha-hora en el pasado — El sistema muestra advertencia.',
        'FA-08b: Aforo maximo menor que las reservas existentes — Error de validacion.',
    ],
    postcondiciones=['La clase queda registrada y visible para todos los usuarios autenticados.'],
    reglas_negocio=['RN-02: Cuando el aforo disponible llega a 0, el estado de la clase cambia automaticamente a "llena" y no se permiten nuevas reservas.'],
    endpoints=['GET /api/v1/clases', 'GET /api/v1/clases/:id', 'POST /api/v1/clases', 'PUT /api/v1/clases/:id', 'DELETE /api/v1/clases/:id']
)

add_rf_block(
    doc, 'RF-09', 'Reserva y Cancelacion de Clases',
    actores='Socio, Recepcionista',
    prioridad='Must Have',
    modulo='Reservas',
    descripcion='El sistema permite a los socios reservar cupos en clases grupales y cancelar sus reservas. Aplica multiples validaciones: membresia activa, cupo disponible, sin penalizaciones activas y sin horarios duplicados.',
    precondiciones=[
        'El actor debe estar autenticado.',
        'La clase debe existir y tener estado "disponible".',
    ],
    flujo_normal=[
        'El socio navega al modulo Clases y selecciona una clase disponible.',
        'Hace clic en "Reservar" y confirma la accion.',
        'El backend valida: (a) el socio tiene membresia activa, (b) el aforo no esta lleno, (c) el socio no tiene penalizacion activa, (d) el socio no tiene ya una reserva en esa clase.',
        'Si todas las validaciones pasan, se crea la reserva con estado "confirmada".',
        'El aforo disponible de la clase se decrementa en 1.',
        'Para cancelar: el socio selecciona la reserva y hace clic en cancelar.',
        'El backend verifica que la clase comience en mas de 2 horas desde el momento actual.',
        'Si la validacion pasa, la reserva cambia a estado "cancelada" y el aforo se incrementa en 1.',
    ],
    flujos_alt=[
        'FA-09a: Membresia vencida — Error 400, el socio debe renovar antes de reservar.',
        'FA-09b: Clase sin cupos — Error 409 CLASE_SIN_CUPOS.',
        'FA-09c: Penalizacion activa — Error 400 PENALIZACION_ACTIVA con fecha de desbloqueo.',
        'FA-09d: Reserva duplicada — Error 409 RESERVA_DUPLICADA.',
        'FA-09e: Cancelacion fuera de plazo — Error 400 CANCELACION_FUERA_DE_PLAZO (menos de 2h antes).',
    ],
    postcondiciones=[
        'La reserva queda registrada con estado "confirmada" y visible en el historial del socio.',
        'El aforo disponible de la clase se actualiza en tiempo real.',
    ],
    reglas_negocio=[
        'RN-02: No se permiten reservas cuando el aforo llega a cero.',
        'RN-03: Cancelacion solo hasta 2 horas antes del inicio de la clase.',
        'RN-07: Un socio no puede tener dos reservas confirmadas para la misma clase.',
    ],
    endpoints=['POST /api/v1/reservas', 'GET /api/v1/reservas', 'GET /api/v1/socios/:id/reservas', 'DELETE /api/v1/reservas/:id/cancelar']
)

add_rf_block(
    doc, 'RF-10', 'Bloqueo por Intentos Fallidos de Login',
    actores='Sistema (automatico)',
    prioridad='Must Have',
    modulo='Autenticacion',
    descripcion='El sistema bloquea temporalmente una cuenta de usuario durante 30 minutos cuando se detectan 5 intentos de inicio de sesion fallidos consecutivos, para proteger contra ataques de fuerza bruta.',
    precondiciones=['El usuario debe intentar autenticarse con credenciales incorrectas.'],
    flujo_normal=[
        'El sistema lleva un contador de intentos fallidos por email en memoria o cache.',
        'Cada intento fallido incrementa el contador.',
        'Al llegar al intento 5, el sistema establece un bloqueo temporal de 30 minutos.',
        'Durante el bloqueo, cualquier intento de login (incluso con credenciales correctas) es rechazado con error 429.',
        'Pasados los 30 minutos, el contador se resetea y el usuario puede intentar nuevamente.',
    ],
    flujos_alt=['FA-10a: Login exitoso antes del quinto intento — El contador se resetea a cero.'],
    postcondiciones=['La cuenta permanece bloqueada por 30 minutos desde el quinto intento fallido.'],
    reglas_negocio=['RN-10: Seguridad contra ataques de diccionario y fuerza bruta.'],
    endpoints=['POST /api/v1/auth/login']
)

add_rf_block(
    doc, 'RF-11', 'Control de Acceso por Roles (RBAC)',
    actores='Sistema (automatico)',
    prioridad='Must Have',
    modulo='Autenticacion',
    descripcion='El sistema aplica control de acceso basado en roles en cada endpoint protegido. Cada operacion solo puede ser ejecutada por los roles autorizados, y cualquier intento de acceso no autorizado es rechazado con error 403.',
    precondiciones=['El usuario debe estar autenticado con un JWT valido.'],
    flujo_normal=[
        'Cada peticion al backend incluye el header Authorization: Bearer <token>.',
        'El middleware auth.middleware.js verifica la firma y vigencia del JWT.',
        'El middleware roles.middleware.js verifica que el rol del usuario esta en la lista de roles permitidos para ese endpoint.',
        'Si el rol es valido, la peticion continua al controlador correspondiente.',
        'Si el JWT es invalido o expirado: error 401.',
        'Si el rol no esta autorizado: error 403.',
    ],
    flujos_alt=[
        'FA-11a: Token ausente — Error 401 UNAUTHORIZED.',
        'FA-11b: Token expirado — Error 401 TOKEN_EXPIRED.',
        'FA-11c: Rol insuficiente — Error 403 FORBIDDEN.',
    ],
    postcondiciones=['Solo los roles autorizados pueden ejecutar cada operacion del sistema.'],
    reglas_negocio=['RN: El principio de minimo privilegio se aplica a todos los endpoints de la API.'],
    endpoints=['Todos los endpoints protegidos de la API /api/v1/*']
)

# ────────────────────────────────────────────────────────
doc.add_heading('3.2 Should Have — Requisitos Importantes (RF-12 a RF-19)', 2)

add_rf_block(
    doc, 'RF-12', 'Busqueda y Filtrado de Socios',
    actores='Administrador, Recepcionista',
    prioridad='Should Have',
    modulo='Socios',
    descripcion='El sistema permite buscar socios por nombre, apellido o DNI, y filtrarlos por estado del socio (activo/inactivo) o estado de membresia (activa/vencida/bloqueada). Los resultados se presentan paginados.',
    precondiciones=['El actor debe estar autenticado con rol administrador o recepcionista.'],
    flujo_normal=[
        'El actor escribe en la barra de busqueda o selecciona un filtro en el listado de socios.',
        'El frontend envia los parametros como query params: ?busqueda=nombre&estado=activo&membresiaEstado=activa.',
        'El backend construye la consulta SQL dinamicamente segun los filtros activos.',
        'Los resultados se devuelven paginados (20 por pagina por defecto) con el total de coincidencias.',
    ],
    flujos_alt=['FA-12a: Sin coincidencias — El sistema muestra un estado vacio "No se encontraron socios".'],
    postcondiciones=['El actor visualiza unicamente los socios que coinciden con los criterios de busqueda.'],
    reglas_negocio=[],
    endpoints=['GET /api/v1/socios?busqueda=&estado=&membresiaEstado=&pagina=&limite=']
)

add_rf_block(
    doc, 'RF-13', 'Registro de Inasistencias y Strikes',
    actores='Recepcionista, Administrador',
    prioridad='Should Have',
    modulo='Penalizaciones',
    descripcion='El sistema permite registrar la inasistencia de un socio a una clase reservada. Cada inasistencia injustificada genera un strike. El contador de strikes es visible en la tarjeta del socio y se cuenta en una ventana de 30 dias.',
    precondiciones=[
        'El socio debe tener una reserva con estado "confirmada" para la clase.',
        'La clase debe haber pasado su horario de inicio.',
    ],
    flujo_normal=[
        'El recepcionista o administrador accede a la lista de reservas de una clase pasada.',
        'Selecciona al socio que no se presento y marca su asistencia como "inasistencia".',
        'El sistema cambia el estado de la reserva a "inasistencia".',
        'Se inserta un registro en la tabla penalizaciones asociado a esa reserva, con justificada = false.',
        'El contador de strikes del socio (penalizaciones no justificadas en los ultimos 30 dias) se incrementa.',
        'Si el conteo llega a 3: se activa automaticamente el flujo de RF-14.',
    ],
    flujos_alt=['FA-13a: El socio ya tenia la reserva cancelada — No se puede marcar inasistencia.'],
    postcondiciones=['El strike queda registrado y el contador actualizado en el perfil del socio.'],
    reglas_negocio=['RN-04: Los strikes se contabilizan en una ventana movil de 30 dias. Un strike antiguo deja de contar automaticamente.'],
    endpoints=['PUT /api/v1/reservas/:id (estado: inasistencia)', 'GET /api/v1/socios/:id/strikes']
)

add_rf_block(
    doc, 'RF-14', 'Bloqueo Automatico por Acumulacion de Strikes',
    actores='Sistema (automatico)',
    prioridad='Should Have',
    modulo='Penalizaciones',
    descripcion='Cuando un socio acumula 3 strikes de inasistencia injustificada en los ultimos 30 dias, el sistema bloquea automaticamente su capacidad de hacer nuevas reservas y su acceso fisico al gimnasio por 7 dias.',
    precondiciones=['El socio debe tener exactamente 3 penalizaciones no justificadas en los ultimos 30 dias.'],
    flujo_normal=[
        'Al registrarse el tercer strike del socio (RF-13), el sistema verifica el conteo total.',
        'Al detectar el tercer strike activo, se inserta un registro en penalizaciones_bloqueo con bloqueado_hasta = AHORA + 7 dias.',
        'El estado de la membresia activa del socio cambia a "bloqueada".',
        'A partir de ese momento: el socio no puede hacer nuevas reservas (RF-09 lo verifica) y su QR es denegado en recepcion (RF-07 lo verifica).',
        'Pasados los 7 dias, el bloqueo expira automaticamente y la membresia vuelve a estado "activa".',
    ],
    flujos_alt=['FA-14a: El administrador justifica un strike antes de llegar al tercer bloqueo — El contador baja y el bloqueo no se activa.'],
    postcondiciones=['El socio queda bloqueado durante 7 dias. El bloqueo consta en penalizaciones_bloqueo con su fecha de expiracion.'],
    reglas_negocio=['RN-04: 3 strikes en 30 dias = bloqueo de 7 dias.'],
    endpoints=['GET /api/v1/socios/:id/strikes', 'POST interno al registrar inasistencia']
)

add_rf_block(
    doc, 'RF-15', 'Dashboard Administrativo',
    actores='Administrador',
    prioridad='Should Have',
    modulo='Dashboard',
    descripcion='El sistema presenta al administrador un panel con las metricas operativas clave del dia: ingresos totales, cantidad de pagos, socios activos, membresías proximas a vencer y ocupacion de clases del dia.',
    precondiciones=['El actor debe estar autenticado con rol administrador.'],
    flujo_normal=[
        'El administrador inicia sesion y es redirigido al Dashboard automaticamente.',
        'El frontend realiza una peticion GET /dashboard/resumen.',
        'El backend ejecuta las consultas necesarias para calcular: ingresos del dia (SUM de pagos de HOY), socios activos (COUNT con membresia activa), membresías por vencer en los proximos 3 dias, y ocupacion de clases programadas para HOY.',
        'Los datos se presentan en tarjetas de metricas con iconos y valores destacados.',
    ],
    flujos_alt=['FA-15a: Sin datos del dia (primer dia del mes) — Las metricas muestran 0 con mensaje informativo.'],
    postcondiciones=['El administrador tiene una vision rapida del estado operativo del gimnasio al iniciar su jornada.'],
    reglas_negocio=['RN-08: El Dashboard debe mostrar obligatoriamente: ingresos del dia, membresías por vencer y ocupacion de clases.'],
    endpoints=['GET /api/v1/dashboard/resumen']
)

add_rf_block(
    doc, 'RF-16', 'Vista Mi Perfil del Socio',
    actores='Socio',
    prioridad='Should Have',
    modulo='Mi Perfil',
    descripcion='El sistema provee al socio autenticado una vista personalizada con su informacion completa: codigo QR de acceso, estado y fecha de vencimiento de su membresia, contador de strikes activos, y lista de reservas confirmadas.',
    precondiciones=['El actor debe estar autenticado con rol socio.'],
    flujo_normal=[
        'El socio inicia sesion y es redirigido a Mi Perfil.',
        'El sistema carga: datos del socio, estado de su membresia activa, QR generado, contador de strikes de los ultimos 30 dias, y reservas confirmadas futuras.',
        'El socio puede descargar su QR en imagen.',
        'El socio puede ver el historial de sus pagos y reservas pasadas.',
    ],
    flujos_alt=['FA-16a: Membresia vencida — El perfil muestra alerta de renovacion y el QR indica estado vencido.'],
    postcondiciones=['El socio puede acceder a toda su informacion relevante sin necesidad de contactar a recepcion.'],
    reglas_negocio=[],
    endpoints=['GET /api/v1/socios/:id', 'GET /api/v1/socios/:id/qr', 'GET /api/v1/socios/:id/reservas', 'GET /api/v1/socios/:id/pagos']
)

add_rf_block(
    doc, 'RF-17', 'Justificacion y Eliminacion de Strikes',
    actores='Administrador',
    prioridad='Should Have',
    modulo='Penalizaciones',
    descripcion='El sistema permite al administrador justificar y eliminar un strike de un socio, registrando la razon de la justificacion. Si el socio tenia 3 strikes y el administrador elimina uno, el bloqueo se levanta automaticamente.',
    precondiciones=[
        'El actor debe estar autenticado con rol administrador.',
        'El strike a eliminar debe existir y estar marcado como justificada = false.',
    ],
    flujo_normal=[
        'El administrador accede al perfil del socio y visualiza sus strikes activos.',
        'Selecciona el strike a justificar e ingresa la razon (ej: "Baja medica con certificado").',
        'El sistema actualiza el strike: justificada = true, justificacion = texto ingresado.',
        'El contador de strikes activos del socio se recalcula.',
        'Si el socio tenia bloqueo activo y ahora tiene menos de 3 strikes, el bloqueo se levanta.',
    ],
    flujos_alt=['FA-17a: Strike no encontrado — Error 404.'],
    postcondiciones=['El strike queda justificado con trazabilidad de la razon. El contador del socio se actualiza.'],
    reglas_negocio=['RN: Las justificaciones quedan registradas para auditoria futura.'],
    endpoints=['DELETE /api/v1/socios/:id/strikes/:strikeId']
)

add_rf_block(
    doc, 'RF-18', 'Historial de Pagos',
    actores='Administrador, Recepcionista, Socio',
    prioridad='Should Have',
    modulo='Pagos',
    descripcion='El sistema permite consultar el historial completo de pagos, tanto a nivel global (admin) como por socio especifico. Incluye filtros por fecha, metodo de pago y socio.',
    precondiciones=['El actor debe estar autenticado.'],
    flujo_normal=[
        'El administrador accede al modulo Pagos para ver todos los registros.',
        'Puede filtrar por rango de fechas, metodo de pago o socio especifico.',
        'El socio accede a Mi Perfil para ver unicamente sus propios pagos.',
        'Cada pago muestra: fecha, plan, monto, metodo y estado.',
    ],
    flujos_alt=['FA-18a: Sin pagos registrados — Lista vacia con mensaje informativo.'],
    postcondiciones=['El actor obtiene el historial completo de transacciones segun su nivel de acceso.'],
    reglas_negocio=[],
    endpoints=['GET /api/v1/pagos', 'GET /api/v1/socios/:id/pagos']
)

add_rf_block(
    doc, 'RF-19', 'Log de Accesos QR',
    actores='Administrador, Recepcionista',
    prioridad='Should Have',
    modulo='Acceso QR',
    descripcion='El sistema registra cada intento de acceso por QR (permitido o denegado) en una tabla de logs. El endpoint de consulta esta disponible pero la pantalla dedicada con filtros avanzados se implementara en v1.1.',
    precondiciones=['El actor debe estar autenticado con rol administrador o recepcionista.'],
    flujo_normal=[
        'Cada validacion de QR (RF-07) genera automaticamente un registro en logs_acceso.',
        'El registro incluye: socio_id, resultado (PERMITIDO/DENEGADO), motivo y timestamp.',
        'El administrador puede consultar los logs via GET /acceso/logs con filtros de socio, fecha y resultado.',
    ],
    flujos_alt=[],
    postcondiciones=['Todos los accesos quedan trazados para auditoria.'],
    reglas_negocio=[],
    endpoints=['GET /api/v1/acceso/logs']
)

# ── 4. RF DIFERIDOS ──────────────────────────────────────────────────────────
doc.add_heading('4. Requisitos Funcionales Diferidos', 1)
doc.add_paragraph(
    'Los siguientes RF fueron identificados durante la elicitacion pero quedan fuera del '
    'alcance del MVP1. Se documentan brevemente con su justificacion de exclusion.'
)

diferidos = doc.add_table(rows=1, cols=4)
diferidos.style = 'Table Grid'
header_row(diferidos, ['ID', 'Nombre', 'Categoria', 'Version'])
dif_data = [
    ('RF-20', 'Ficha medica digital (lesiones, objetivo, enfermedades)', 'Could Have', 'v1.1'),
    ('RF-21', 'Asignacion de rutinas de entrenamiento personalizadas', 'Could Have', 'v1.1'),
    ('RF-22', 'Registro de progreso fisico por ejercicio', 'Could Have', 'v1.1'),
    ('RF-23', 'Pantalla dedicada de logs de acceso con filtros avanzados', 'Could Have', 'v1.1'),
    ('RF-24', 'Generacion de comprobante de pago en PDF', 'Could Have', 'v1.1'),
    ('RF-25', 'Indicador visual de ocupacion de clase antes de reservar', 'Could Have', 'v1.1'),
    ('RF-26', 'Lista de espera para clases llenas con notificacion', 'Could Have', 'v1.1'),
    ('RF-27', 'Asignacion de entrenador responsable a cada socio', 'Could Have', 'v1.1'),
    ('RF-28', 'Notificaciones automaticas por email/SMS de vencimiento', "Won't Have", 'v2.0'),
    ('RF-29', 'Portal de pago online (pasarela tarjeta)', "Won't Have", 'v2.0'),
    ('RF-30', 'Exportacion de reportes en Excel/PDF', "Won't Have", 'v2.0'),
    ('RF-31', 'Aplicacion movil nativa (Android/iOS)', "Won't Have", 'v3.0'),
    ('RF-32', 'Modulo de inventario de equipamiento', "Won't Have", 'v3.0'),
    ('RF-33', 'Integracion con torniquete electronico real', "Won't Have", 'v3.0'),
    ('RF-34', 'Codigo QR dinamico con renovacion periodica automatica', "Won't Have", 'v3.0'),
]
for d in dif_data:
    r = diferidos.add_row()
    for c, t in zip(r.cells, d): c.text = t

doc.add_paragraph()

# ── 5. ESPECIFICACION FORMAL DE RNF ──────────────────────────────────────────
doc.add_heading('5. Especificacion Formal de Requisitos No Funcionales', 1)
doc.add_paragraph(
    'Los 13 Requisitos No Funcionales se especifican con sus criterios de verificacion '
    'y aceptacion, que permitiran validar su cumplimiento durante las pruebas del sistema.'
)

rnfs_formales = [
    ('RNF-01', 'Tiempo de respuesta de la API', 'Rendimiento',
     'Los endpoints de la API REST deben responder en menos de 500 ms bajo carga normal.',
     'Hasta 50 usuarios concurrentes realizando peticiones simultaneas.',
     'Medir con herramienta de carga (Apache JMeter o k6). El percentil 95 de tiempos de respuesta no debe superar 500 ms.',
     'Optimizacion de consultas SQL con indices en columnas clave. Uso de connection pooling en PostgreSQL.'),
    ('RNF-02', 'Tiempo de validacion de QR', 'Rendimiento',
     'La validacion del codigo QR debe completarse en menos de 300 ms.',
     'Solicitudes normales de validacion desde recepcion.',
     'Medir el tiempo de respuesta del endpoint POST /acceso/validar con cronometro en el frontend. Media de 20 mediciones < 300 ms.',
     'La logica de validacion del QR es simple (verificar JWT + consulta de membresia indexada).'),
    ('RNF-03', 'Capacidad de usuarios concurrentes', 'Rendimiento',
     'El sistema debe soportar 100 usuarios concurrentes sin degradacion perceptible.',
     'Horario pico del gimnasio (7AM-9AM y 6PM-8PM).',
     'Prueba de carga con 100 usuarios virtuales concurrentes. Tasa de error < 1%. Tiempo de respuesta medio < 800 ms.',
     'Uso de la nube para backend (Koyeb) y base de datos (Neon) con escalado automatico.'),
    ('RNF-04', 'Almacenamiento seguro de contrasenas', 'Seguridad',
     'Las contrasenas deben almacenarse en formato bcrypt con salt >= 10 rondas. Jamas en texto plano.',
     'Durante el registro de cualquier usuario o cambio de contrasena.',
     'Inspeccionar la columna password_hash en la BD: debe iniciar con $2b$10$ (formato bcrypt). No debe existir ninguna columna "password" en texto plano.',
     'Uso de la libreria bcryptjs con factor de costo 10 en todas las operaciones de hash.'),
    ('RNF-05', 'Comunicacion HTTPS obligatoria', 'Seguridad',
     'Toda comunicacion entre frontend y backend debe ser exclusivamente por HTTPS.',
     'En produccion (Vercel + Koyeb).',
     'Verificar que las URLs de produccion usen https://. Verificar que el backend rechace peticiones HTTP (redireccion o error).',
     'Los proveedores de despliegue (Vercel y Koyeb) proveen HTTPS automatico con Let\'s Encrypt.'),
    ('RNF-06', 'Expiracion de tokens JWT', 'Seguridad',
     'Tokens JWT de autenticacion: duracion maxima 24 horas. Tokens QR: duracion configurable (por defecto 12 horas).',
     'En cada generacion de token.',
     'Decodificar el JWT y verificar que el campo "exp" corresponde a 24 horas desde la emision. Para QR, verificar 12 horas.',
     'Configurado en las variables de entorno JWT_EXPIRES_IN y QR_EXPIRES_IN.'),
    ('RNF-07', 'Prevencion de inyeccion SQL y XSS', 'Seguridad',
     'La API debe sanitizar todos los datos de entrada.',
     'En cada peticion entrante con datos del usuario.',
     'Pruebas de penetracion basicas: intentar enviar payloads SQL (OR 1=1) y XSS (<script>alert(1)</script>). El sistema debe rechazarlos o ignorarlos sin errores internos.',
     'Uso de consultas parametrizadas (prepared statements) con la libreria pg de Node.js. Nunca concatenacion de strings en SQL.'),
    ('RNF-08', 'Compatibilidad de navegadores', 'Usabilidad',
     'La interfaz debe funcionar correctamente en Chrome, Firefox, Edge y Safari (ultimas 2 versiones).',
     'Siempre.',
     'Ejecutar el flujo completo (login, registrar socio, validar QR) en cada uno de los 4 navegadores. Sin errores de consola criticos.',
     'Uso de React y CSS estandar sin caracteristicas experimentales de navegador.'),
    ('RNF-09', 'Diseno Responsive', 'Usabilidad',
     'La interfaz debe ser operable en pantallas >= 768 px de ancho.',
     'Siempre.',
     'Abrir la aplicacion en una tablet (768px) o redimensionar el navegador. Todos los elementos deben ser visibles y operables sin scroll horizontal.',
     'Implementado con media queries en el CSS global del proyecto.'),
    ('RNF-10', 'Eficiencia de flujos criticos (max 3 clics)', 'Usabilidad',
     'Las acciones criticas de recepcion deben completarse con maximo 3 interacciones.',
     'En los flujos: validar QR y registrar pago.',
     'Contar los clics requeridos para: (1) abrir Acceso QR, (2) ingresar token, (3) ver resultado. El total no debe superar 3 acciones principales.',
     'La pantalla de Acceso QR es accesible directamente desde el menu lateral en 1 clic.'),
    ('RNF-11', 'Disponibilidad del servicio', 'Disponibilidad',
     'El sistema debe mantener disponibilidad >= 99% en horario de operacion del gimnasio.',
     'De lunes a sabado, entre las 6:00 AM y las 10:00 PM.',
     'Monitorear el uptime durante un mes usando herramienta de monitoreo (UptimeRobot). El resultado debe ser >= 99% en el periodo evaluado.',
     'Despliegue en plataformas cloud con alta disponibilidad (Vercel para frontend, Koyeb para backend, Neon para BD).'),
    ('RNF-12', 'Control de versiones con Git', 'Mantenibilidad',
     'El codigo debe estar versionado en Git con ramas por modulo y etiquetas de version.',
     'Durante todo el desarrollo.',
     'Verificar en GitHub: existencia de ramas por modulo (feature/socios, feature/clases, etc.), historial de commits descriptivos y etiquetas v1.0 para cada entrega.',
     'Politica de desarrollo acordada por el equipo: commits atomicos, pull requests con revision antes de merge a main.'),
    ('RNF-13', 'Indices en base de datos', 'Rendimiento',
     'La BD debe tener indices en columnas de busqueda frecuente: dni, socio_id, estado, fecha_fin, fecha_hora.',
     'Desde el primer despliegue de produccion.',
     'Ejecutar EXPLAIN ANALYZE en las consultas criticas del sistema (buscar socio por DNI, obtener membresia activa). El plan de ejecucion debe mostrar "Index Scan" en lugar de "Seq Scan".',
     'Indices definidos en el script de migracion de la base de datos.'),
]

for rnf in rnfs_formales:
    rnf_id, nombre, cat, descripcion, contexto, criterio, implementacion = rnf
    doc.add_heading(f'{rnf_id} — {nombre}', 3)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = 'Table Grid'
    filas = [
        ('Identificador', rnf_id),
        ('Categoria', cat),
        ('Descripcion / Requerimiento', descripcion),
        ('Contexto de aplicacion', contexto),
        ('Criterio de verificacion y aceptacion', criterio),
        ('Estrategia de implementacion', implementacion),
    ]
    for i, (key, val) in enumerate(filas):
        tbl.rows[i].cells[0].text = key
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[i].cells[0], 'D0D8E4')
        tbl.rows[i].cells[1].text = val
    doc.add_paragraph()

# ── 6. GLOSARIO ──────────────────────────────────────────────────────────────
doc.add_heading('6. Glosario de Terminos', 1)
glosario = [
    ('JWT', 'JSON Web Token. Estandar de token de autenticacion firmado digitalmente que contiene la identidad y rol del usuario.'),
    ('QR', 'Quick Response. Codigo de barras bidimensional que el sistema usa como mecanismo de acceso fisico al gimnasio.'),
    ('Strike', 'Penalizacion registrada por inasistencia injustificada a una clase reservada. Tres strikes en 30 dias activan un bloqueo.'),
    ('Membresia', 'Periodo de suscripcion activa de un socio a un plan del gimnasio, delimitado por una fecha de inicio y una fecha de fin.'),
    ('Aforo', 'Capacidad maxima de personas permitidas en una clase grupal especifica.'),
    ('Borrado logico', 'Tecnica de eliminacion que no borra fisicamente el registro de la BD, sino que cambia su estado a "inactivo" para preservar la trazabilidad.'),
    ('RBAC', 'Role-Based Access Control. Sistema de control de acceso donde los permisos se asignan segun el rol del usuario.'),
    ('MVP', 'Minimum Viable Product. Version funcional del sistema con las caracteristicas minimas necesarias para demostrar valor.'),
    ('MoSCoW', 'Tecnica de priorizacion: Must Have, Should Have, Could Have, Won\'t Have.'),
    ('bcrypt', 'Algoritmo de hashing de contrasenas que incorpora un factor de costo para resistir ataques de fuerza bruta.'),
]

g_table = doc.add_table(rows=1, cols=2)
g_table.style = 'Table Grid'
header_row(g_table, ['Termino', 'Definicion'])
for term, defin in glosario:
    r = g_table.add_row()
    r.cells[0].text = term
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[1].text = defin

doc.add_paragraph()

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = r'c:/Users/ACER/OneDrive/Desktop/PROGRAMACION/UNIVERSITY PROJECTS/SIstemaGestorGimnasio/Documentacion/ESPECIFICACION_FORMAL_actualizado.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
