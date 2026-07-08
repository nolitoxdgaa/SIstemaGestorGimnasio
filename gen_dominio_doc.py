"""
gen_dominio_doc.py
Genera: Requisitos_Dominio_OLYMPUS_CORE_actualizado.docx
Mantiene la estructura del documento original (14 paginas / 5 secciones)
y actualiza el contenido segun los 34 RF, 13 RNF y hallazgos de los
documentos actualizados (Elicitacion, MVP1, Especificacion Formal, Calidad).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests, zlib, io

# ══════════════════════════════════════════════════════════════════════════════
# PLANTUML — helpers para embeber diagramas
# ══════════════════════════════════════════════════════════════════════════════
def _encode6bit(b):
    if b < 10: return chr(48 + b)
    b -= 10
    if b < 26: return chr(65 + b)
    b -= 26
    if b < 26: return chr(97 + b)
    b -= 26
    return '-' if b == 0 else '_'

def _append3bytes(b1, b2, b3):
    return (
        _encode6bit(b1 >> 2) +
        _encode6bit(((b1 & 0x3) << 4) | (b2 >> 4)) +
        _encode6bit(((b2 & 0xF) << 2) | (b3 >> 6)) +
        _encode6bit(b3 & 0x3F)
    )

def plantuml_encode(text):
    data = zlib.compress(text.encode('utf-8'))[2:-4]
    res = ""
    for i in range(0, len(data), 3):
        if i + 2 < len(data): res += _append3bytes(data[i], data[i+1], data[i+2])
        elif i + 1 < len(data): res += _append3bytes(data[i], data[i+1], 0)
        else: res += _append3bytes(data[i], 0, 0)
    return res

def fetch_diagram(uml_code, width_inches=5.8):
    try:
        enc = plantuml_encode(uml_code)
        url = f"http://www.plantuml.com/plantuml/png/{enc}"
        r = requests.get(url, timeout=25)
        if r.status_code == 200 and r.content[:4] == b'\x89PNG':
            return io.BytesIO(r.content), width_inches
    except Exception:
        pass
    return None

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, cols, bg='1E3A5F', fg='FFFFFF'):
    for cell, txt in zip(table.rows[0].cells, cols):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        set_cell_bg(cell, bg)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

def add_diagram(doc, uml_code, caption, width=5.6):
    result = fetch_diagram(uml_code, width)
    if result:
        img_stream, w = result
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(w))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(f'[Diagrama: {caption}]').italic = True
    doc.add_paragraph()

def spec_table(doc, campos):
    """Inserta tabla de 2 cols Campo/Valor."""
    tbl = doc.add_table(rows=len(campos), cols=2)
    tbl.style = 'Table Grid'
    for i, (key, val) in enumerate(campos):
        ck = tbl.rows[i].cells[0]
        cv = tbl.rows[i].cells[1]
        ck.text = key
        ck.paragraphs[0].runs[0].bold = True
        set_cell_bg(ck, 'D0D8E4')
        cv.text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMAS UML
# ══════════════════════════════════════════════════════════════════════════════

UML_DCU_NEGOCIO = """@startuml
left to right direction
skinparam packageStyle rectangle
skinparam actorStyle awesome

title Diagrama de Casos de Uso — Nivel de Negocio\\n(Procesos del dominio antes de la automatizacion)

actor "Socio / Cliente" as SC
actor "Recepcionista" as REC
actor "Entrenador" as ENT
actor "Administrador" as ADM
actor "Pasarela de Pagos" as PAY

package "Gimnasio OLYMPUS CORE (Manual/Parcial)" {
  usecase "Controlar acceso\\n(Excel + Registro manual)" as BUC1
  usecase "Gestionar membresias\\ny renovaciones" as BUC2
  usecase "Reservar clase grupal\\n(WhatsApp/llamada)" as BUC3
  usecase "Registrar pago\\nen caja fisica" as BUC4
  usecase "Asignar rutina\\nde entrenamiento" as BUC5
  usecase "Registrar progreso\\nfisico del socio" as BUC6
  usecase "Supervisar\\nfinanzas del negocio" as BUC7
}

SC --> BUC1
SC --> BUC3
SC --> BUC4
REC --> BUC1
REC --> BUC2
REC --> BUC4
ENT --> BUC5
ENT --> BUC6
ADM --> BUC7
ADM --> BUC2
PAY --> BUC4
@enduml"""

UML_DCU_SISTEMA = """@startuml
left to right direction
skinparam packageStyle rectangle
skinparam actorStyle awesome

title Diagrama de Casos de Uso — Nivel de Sistema\\n(Funcionalidades que el sistema OLYMPUS CORE automatizara)

actor "A-Socio" as SOC
actor "A-Recepcionista" as REC
actor "A-Entrenador" as ENT
actor "A-Administrador" as ADM
actor "A-LectorQR" as QR
actor "A-PasarelaPagos" as PAY

rectangle "Sistema OLYMPUS CORE — MVP1" {
  usecase "RF-01 Iniciar Sesion (RBAC)" as UC01
  usecase "RF-02 Registrar Socio" as UC02
  usecase "RF-03 Gestionar Planes" as UC03
  usecase "RF-04 Ver Estado Membresia" as UC04
  usecase "RF-05 Registrar Pago" as UC05
  usecase "RF-06 Validar QR de Acceso" as UC06
  usecase "RF-07 Ver Dashboard Admin" as UC07
  usecase "RF-08 Listar Socios" as UC08
  usecase "RF-09 Reservar Clase Grupal" as UC09
  usecase "RF-10 Ver Perfil y QR" as UC10
  usecase "RF-11 Gestion de Roles" as UC11
  usecase "RF-14 Penalizacion Automatica" as UC14
  usecase "RF-16 Gestion de Clases" as UC16
  usecase "RF-17 Justificar Strike" as UC17
}

ADM --> UC01
ADM --> UC02
ADM --> UC03
ADM --> UC07
ADM --> UC11
ADM --> UC17

REC --> UC01
REC --> UC02
REC --> UC05
REC --> UC06
REC --> UC08

SOC --> UC09
SOC --> UC10
SOC --> UC04

ENT --> UC16

QR --> UC06
PAY --> UC05
@enduml"""

UML_CLASES_DOMINIO = """@startuml
hide circle
skinparam classAttributeIconSize 0
skinparam class {
  BackgroundColor #F0F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
  FontName Calibri
}

class Usuarios {
  + id : UUID
  + nombre : String
  + email : String
  + password_hash : String
  + rol : Enum {admin,recepcionista,socio,entrenador}
  + activo : Boolean
  + createdAt : DateTime
}

class Socios {
  + id : UUID
  + dni : String
  + telefono : String
  + estado : Enum {activo,inactivo,bloqueado}
  + strikes_activos : Int
  + fecha_registro : Date
  + codigo_qr_token : String
}

class Planes {
  + id : UUID
  + nombre : String
  + precio : Decimal
  + duracion_dias : Int
  + descripcion : String
}

class Membresias {
  + id : UUID
  + fecha_inicio : Date
  + fecha_fin : Date
  + monto_pagado : Decimal
  + estado : Enum {activa,vencida,bloqueada}
}

class ClasesGrupales {
  + id : UUID
  + nombre : String
  + descripcion : String
  + aforo_maximo : Int
  + aforo_actual : Int
  + fecha_hora : DateTime
  + estado : Enum {programada,en_curso,finalizada,cancelada}
}

class Reservas {
  + id : UUID
  + fecha_reserva : DateTime
  + asistio : Boolean
  + estado : Enum {confirmada,cancelada,penalizada}
}

class Penalizaciones {
  + id : UUID
  + tipo : Enum {strike,bloqueo}
  + fecha : DateTime
  + justificada : Boolean
  + justificacion : String
}

class Pagos {
  + id : UUID
  + monto : Decimal
  + fecha_pago : DateTime
  + metodo_pago : String
  + estado : Enum {completado,fallido}
}

class FichasMedicas {
  + id : UUID
  + lesiones_previas : Text
  + enfermedades_cv : Text
  + objetivo : String
  + fecha_creacion : Date
}

Usuarios "1" -- "1" Socios : es >
Socios "1" -- "*" Membresias : contrata >
Planes "1" -- "*" Membresias : < define
Socios "1" -- "*" Reservas : realiza >
ClasesGrupales "1" -- "*" Reservas : tiene >
Reservas "1" -- "0..1" Penalizaciones : genera >
Membresias "1" -- "*" Pagos : se cancela con >
Socios "1" -- "0..1" FichasMedicas : completa >

note bottom of Socios
  RN-01: Solo accede con membresia activa
  RN-04: 3 strikes = bloqueo 7 dias
  RD-01, RF-06, RF-07
end note
note bottom of Membresias
  RN-05: Precio fijo segun catalogo
  RD-02, RF-03, RF-05
end note
note bottom of ClasesGrupales
  RN-02: Aforo maximo no excedible
  RD-03, RF-09, RF-16
end note
@enduml"""

UML_ACTIVIDAD_RESERVA = """@startuml
skinparam activity {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}
skinparam roundcorner 15
skinparam swimlaneWidth 160

title Diagrama de Actividades con Swimlanes\\nReserva de Clase Grupal (RD-03, RD-04, RD-07)

|Socio|
start
:Selecciona clase grupal deseada;

|Sistema|
:Verificar membresia activa (RD-01);
if (Membresia vigente?) then (no)
  :Mostrar error\\nMEMBRESIA_VENCIDA;
  stop
else (si)
endif

:Verificar penalizacion activa (RD-04);
if (Tiene bloqueo activo?) then (si)
  :Mostrar fecha de\\ndesbloqueo y strikes;
  stop
else (no)
endif

:Verificar aforo disponible (RD-03);
if (Quedan cupos?) then (no)
  :Mostrar error\\nCLASE_SIN_CUPOS;
  stop
else (si)
endif

:Verificar duplicado de horario (RN-07);
if (Ya tiene reserva en ese horario?) then (si)
  :Mostrar error\\nRESERVA_DUPLICADA;
  stop
else (no)
endif

|Socio|
:Confirmar reserva;

|Sistema|
:Crear reserva con estado 'confirmada';
:Actualizar aforo_actual + 1;
:Mostrar confirmacion con politica\\nde cancelacion (2 horas — RD-07);

|Socio|
:Reserva confirmada;
stop
@enduml"""

UML_SECUENCIA_QR = """@startuml
skinparam sequenceArrowThickness 2
skinparam roundcorner 10

title Diagrama de Secuencia — Validacion de Acceso QR\\n(RD-01 | RF-06 | RN-01)

actor "Recepcionista / Socio" as REC
participant "Frontend\\nReact" as UI
participant "API Backend\\nNode/Express" as API
database "PostgreSQL" as DB
participant "Torniquete / UI" as HW

REC -> UI: Escanea QR del socio
UI -> API: POST /acceso/validar {token_qr}
activate API

API -> API: 1. Verificar firma y\\nexpiracion del JWT (RNF-06)
note right: Si expirado -> 401

API -> DB: SELECT socios, membresias\\nWHERE socio_id = :id [INDEX SCAN]
DB --> API: Datos del socio y membresia

API -> DB: SELECT penalizaciones\\nWHERE bloqueado_hasta > NOW()
DB --> API: Sin bloqueo activo

alt Acceso PERMITIDO
  API -> DB: INSERT logs_acceso (resultado=PERMITIDO)
  API --> UI: 200 OK {resultado: "PERMITIDO"}
  UI --> HW: Senal: ABRIR TORNIQUETE
  note right: RQ-CAL-OLY-01:\\nTodo el flujo < 300 ms
else Acceso DENEGADO
  API -> DB: INSERT logs_acceso (resultado=DENEGADO, motivo)
  API --> UI: 403 {motivo: MEMBRESIA_VENCIDA\\no PENALIZACION_ACTIVA}
  UI --> HW: Senal: MANTENER CERRADO\\n+ Mostrar motivo en pantalla
end

deactivate API
@enduml"""

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCION DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── PORTADA ────────────────────────────────────────────────────────────────────
for text, size, bold, color in [
    ('UNIVERSIDAD NACIONAL MAYOR DE SAN MARCOS', 14, True,  (0x1E,0x3A,0x5F)),
    ('FACULTAD DE SISTEMAS E INFORMÁTICA',       13, False, (0x1E,0x3A,0x5F)),
    ('',                                           11, False, (0,0,0)),
    ('TAREA — ACTUALIZACIÓN REQUISITOS DE DOMINIO', 12, True, (0x44,0x4A,0x57)),
    ('',                                           11, False, (0,0,0)),
    ('ASIGNATURA',                                 11, False, (0,0,0)),
    ('Ingeniería de Requisitos',                   12, True,  (0x1E,0x3A,0x5F)),
    ('',                                           11, False, (0,0,0)),
    ('DOCENTE',                                    11, False, (0,0,0)),
    ('RODRÍGUEZ RODRÍGUEZ, CIRO',                  12, False, (0,0,0)),
    ('',                                           11, False, (0,0,0)),
    ('ALUMNOS:',                                   11, False, (0,0,0)),
    ('Cordova Guerra, Josue Rodrigo',              12, False, (0,0,0)),
    ('',                                           11, False, (0,0,0)),
    ('Proyecto: OLYMPUS CORE — Sistema Gestor de Gimnasio', 11, False, (0,0,0)),
    ('Iteración: MVP1 Actualizado | 2026',         11, False, (0,0,0)),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 1: Identificación y Análisis de Requisitos de Dominio
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Identificación y Análisis de Requisitos de Dominio', 1)

# 1.1 Descripción General
doc.add_heading('1.1. Descripción General del Dominio', 2)
doc.add_paragraph(
    'El dominio corresponde a la gestión operativa y comercial de un gimnasio de mediano tamaño '
    'denominado OLYMPUS CORE, con aproximadamente 500 socios activos. El negocio se articula en '
    'torno a cuatro procesos fundamentales: (1) administración de membresías y accesos, '
    '(2) reserva y control de clases grupales, (3) seguimiento personalizado de rutinas de '
    'entrenamiento y (4) control financiero de pagos y renovaciones.\n\n'
    'Actualmente estos procesos se gestionan de forma manual (Excel compartido, WhatsApp, '
    'cuadernos físicos), generando pérdidas económicas por accesos con membresía vencida, '
    'inconsistencia de datos y baja satisfacción de los socios. El dominio incluye reglas de '
    'negocio estrictas como topes de aforo por clase, políticas de penalización por inasistencia '
    'y catálogos cerrados de planes.\n\n'
    'El sistema MVP1 de OLYMPUS CORE digitaliza los procesos críticos del negocio mediante '
    'una arquitectura web (React + Node.js/Express + PostgreSQL en Neon), exponiendo '
    '19 Requisitos Funcionales Must Have / Should Have y 13 Requisitos No Funcionales '
    'clasificados según ISO/IEC 25010.'
)

# 1.2 Fuentes de Información
doc.add_heading('1.2. Fuentes de Información Utilizadas', 2)
doc.add_paragraph(
    'De acuerdo con la metodología de Ingeniería de Requisitos, se emplearon las siguientes '
    'fuentes para identificar los requisitos del dominio:'
)
for f in [
    'Entrevistas semiestructuradas con el Dueño/Administrador y los Entrenadores.',
    'Observación directa (shadowing) en recepción durante horas pico (7:00–9:00 AM y 6:00–8:00 PM).',
    'Análisis del archivo Excel de gestión de socios (documento existente).',
    'Cuestionario digital enviado a los 500 socios vía WhatsApp (tasa de respuesta: 43%).',
    'Benchmarking de sistemas similares: GymMaster, Mindbody y PerfectGym.',
]:
    add_bullet(doc, f)

# 1.3 Actores del Dominio
doc.add_heading('1.3. Actores del Dominio', 2)
doc.add_paragraph('A partir del análisis del negocio se identificaron los siguientes actores:')
act_t = doc.add_table(rows=1, cols=3)
act_t.style = 'Table Grid'
header_row(act_t, ['Actor de Negocio', 'Rol en el Negocio', 'Actor en el Sistema'])
actores = [
    ('Socio / Cliente',      'Consume servicios, reserva clases, renueva membresía, consulta QR y perfil.',           'A-Socio'),
    ('Recepcionista',        'Registra socios, cobra pagos, valida acceso QR, lista socios.',                          'A-Recepcionista'),
    ('Entrenador',           'Asigna rutinas, gestiona clases grupales (horarios, aforo, estado).',                    'A-Entrenador'),
    ('Administrador',        'Controla finanzas, configura planes, supervisa el negocio, gestiona roles y usuarios.',  'A-Administrador'),
    ('Pasarela de Pagos',    'Entidad externa que valida transacciones financieras (integración API futura).',         'A-PasarelaPagos (actor externo)'),
    ('Lector QR (Hardware)', 'Dispositivo que escanea el código QR del socio en recepción.',                          'A-LectorQR (actor tecnológico)'),
]
for row in actores:
    r = act_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

# 1.4 Entidades Principales
doc.add_heading('1.4. Entidades Principales del Dominio', 2)
doc.add_paragraph('Del análisis de los documentos y entrevistas se identificaron las siguientes entidades clave del dominio:')
entidades = [
    'Usuarios: Cuenta de sistema con rol (admin, recepcionista, socio, entrenador). Gestiona autenticación y control de acceso (RBAC).',
    'Socios: Persona registrada con datos personales y membresía asociada. Posee un código QR dinámico único.',
    'Membresía: Contrato de acceso con tipo de plan, fechas de vigencia y estado (activa/vencida/bloqueada).',
    'Planes: Catálogo cerrado de planes (Mensual, Trimestral, Anual) con precio fijo. Solo el Administrador puede modificarlos.',
    'Clase Grupal: Actividad (Spinning, CrossFit) con aforo máximo, horario y estado. Gestionada por el Entrenador.',
    'Reserva: Relación entre Socio y Clase Grupal. Registra asistencia/inasistencia para el sistema de penalizaciones.',
    'Penalización: Strike unitario por inasistencia injustificada. 3 strikes en 30 días = bloqueo de 7 días.',
    'Pagos: Transacción financiera asociada a la renovación de membresía. Registra monto, fecha y método.',
    'Ficha Médica: Formulario con lesiones previas, enfermedades cardiovasculares y objetivo. Prerrequisito para asignar rutinas.',
    'Log de Acceso: Registro auditadle de cada intento de acceso (QR) con resultado y motivo.',
]
for e in entidades:
    add_bullet(doc, e)

# 1.5 Reglas de Negocio
doc.add_heading('1.5. Reglas de Negocio Identificadas', 2)
doc.add_paragraph('Las siguientes reglas son inherentes al dominio y deben incorporarse al sistema:')
rn_t = doc.add_table(rows=1, cols=3)
rn_t.style = 'Table Grid'
header_row(rn_t, ['ID', 'Regla de Negocio', 'Fuente'])
reglas = [
    ('RN-01', 'Un socio solo puede acceder al gimnasio con membresía vigente. Si está vencida o bloqueada, el acceso es denegado con mensaje diferenciado.', 'Administrador / Observación directa'),
    ('RN-02', 'El aforo máximo por clase no puede excederse. El sistema bloquea nuevas reservas cuando el cupo está lleno.', 'Coordinador de clases'),
    ('RN-03', 'Las cancelaciones de reserva solo son permitidas hasta 2 horas antes del inicio de la clase.', 'Recepción / Entrevista dueño'),
    ('RN-04', 'Tres inasistencias no justificadas en 30 días generan el bloqueo temporal de reservas por 7 días calendario.', 'Administrador'),
    ('RN-05', 'Los precios de los planes son fijos y definidos en catálogo cerrado. Solo el Administrador puede modificarlos; la Recepción no puede alterar montos.', 'Análisis de documento Excel'),
    ('RN-06', 'La asignación de rutinas de entrenamiento requiere que el socio haya completado previamente su ficha médica digital.', 'Entrenador / Entrevista'),
    ('RN-07', 'Un socio no puede reservar dos clases en el mismo horario (reserva duplicada).', 'Reglamento interno'),
    ('RN-08', 'El dashboard gerencial debe mostrar: ingresos del día, membresías por vencer (próximos 3 días) e índice de ocupación de clases.', 'Administrador / Entrevista'),
    ('RN-09', 'El Administrador puede remover un strike manualmente con justificación registrada en el sistema.', 'Administrador (añadido post-validación SC-01)'),
    ('RN-10', 'El código QR de acceso del socio expira cada 12 horas y debe ser regenerado por el sistema automáticamente.', 'Benchmarking (Mindbody) / Equipo de desarrollo'),
]
for rn in reglas:
    r = rn_t.add_row()
    for c, t in zip(r.cells, rn): c.text = t
doc.add_paragraph()

# 1.6 Procesos del Negocio
doc.add_heading('1.6. Procesos del Negocio Identificados', 2)
doc.add_paragraph('El modelado de procesos reveló los siguientes flujos críticos del dominio:')
procesos = [
    'Proceso de control de acceso: Escaneo QR → validación de membresía y penalización → apertura/bloqueo → registro de log de auditoría. [RD-01, RF-06]',
    'Proceso de reserva de clase grupal: Consulta disponibilidad → validación membresía y penalizaciones → validación aforo → confirmación → actualización de aforo → notificación. [RD-03, RF-09]',
    'Proceso de renovación de membresía: Identificación del socio → selección de plan → procesamiento de pago → activación de membresía → regeneración de QR. [RD-02, RF-05]',
    'Proceso de gestión de penalizaciones: Detección de inasistencia al finalizar la clase → registro de strike → evaluación de umbral (3/30 días) → bloqueo automático o justificación por Administrador. [RD-04, RF-14, RF-17]',
    'Proceso de seguimiento de entrenamiento: Validación ficha médica → asignación de rutina → registro periódico de progreso físico. [RD-05, RF-12]',
    'Proceso de gestión financiera: Registro de pagos → actualización de membresía → generación de reporte en dashboard. [RD-06, RF-05, RF-07]',
]
for p in procesos:
    add_bullet(doc, p)

# 1.7 Diagrama de Casos de Uso — Nivel de Negocio
doc.add_heading('1.7. Diagrama de Casos de Uso — Nivel de Negocio', 2)
doc.add_paragraph(
    'El siguiente diagrama representa los actores de negocio y los procesos del negocio '
    'antes de la automatización completa del sistema:'
)
add_diagram(doc, UML_DCU_NEGOCIO, 'Figura 1.7: DCU Nivel de Negocio — Procesos del dominio OLYMPUS CORE', 6.0)

# 1.8 Diagrama de Casos de Uso — Nivel de Sistema
doc.add_heading('1.8. Diagrama de Casos de Uso — Nivel de Sistema', 2)
doc.add_paragraph(
    'El siguiente diagrama representa los actores del sistema y las funcionalidades que '
    'el software automatizará en el MVP1 (19 RF implementados):'
)
add_diagram(doc, UML_DCU_SISTEMA, 'Figura 1.8: DCU Nivel de Sistema — MVP1 OLYMPUS CORE (19 RF)', 6.2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 2: Especificación de Requisitos de Dominio
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Especificación de Requisitos de Dominio', 1)

# 2.1 Tabla de RD
doc.add_heading('2.1. Requisitos de Dominio', 2)
doc.add_paragraph(
    'Los requisitos de dominio capturan el "qué" debe hacer el sistema en relación con el '
    'mundo real del gimnasio, independientemente de la tecnología utilizada. La tabla incluye '
    'los requisitos originales más los derivados del proceso de validación y del MVP1:'
)
rd_t = doc.add_table(rows=1, cols=4)
rd_t.style = 'Table Grid'
header_row(rd_t, ['ID', 'Descripción', 'Prioridad', 'Fuente / Stakeholder'])
rds = [
    ('RD-01', 'El sistema debe verificar la vigencia y el estado de la membresía de un socio en tiempo real (< 300 ms) para controlar el acceso físico al gimnasio mediante un código QR dinámico.', 'Alta', 'Dueño / Observación directa'),
    ('RD-02', 'El sistema debe gestionar los planes de membresía con catálogo cerrado de precios, sin permitir modificación manual del monto por parte de la Recepción.', 'Alta', 'Análisis de documento Excel'),
    ('RD-03', 'El sistema debe controlar el aforo de las clases grupales en tiempo real, impidiendo reservas cuando el cupo esté lleno.', 'Alta', 'Coordinador de clases'),
    ('RD-04', 'El sistema debe aplicar automáticamente la política de penalización: 1 inasistencia = 1 strike; 3 strikes en 30 días = bloqueo de reservas por 7 días calendario.', 'Alta', 'Administrador'),
    ('RD-05', 'El sistema debe requerir que el socio complete una ficha médica digital (lesiones, enfermedades cardiovasculares, objetivo) antes de permitir la asignación de una rutina de entrenamiento.', 'Alta', 'Entrenador / Entrevista'),
    ('RD-06', 'El sistema debe presentar un dashboard gerencial con métricas en tiempo real: ingresos del día, membresías por vencer en los próximos 3 días e índice de ocupación de clases.', 'Alta', 'Administrador / Entrevista'),
    ('RD-07', 'El sistema debe permitir la cancelación de reservas únicamente hasta 2 horas antes del inicio de la clase; fuera de ese margen la reserva queda confirmada.', 'Media', 'Recepción'),
    ('RD-08', 'El sistema debe registrar y mantener el historial de progreso físico del socio (peso levantado, fecha, observaciones) para facilitar el seguimiento por el entrenador.', 'Media', 'Entrenador'),
    ('RD-09', 'El sistema debe generar un log de auditoría de cada intento de acceso (socio, fecha/hora, resultado, motivo de denegación) para consulta administrativa.', 'Media', 'Administrador'),
    ('RD-10', 'El Administrador debe poder remover o justificar un strike de penalización manualmente, con registro de la justificación en el sistema.', 'Media', 'Administrador (SC-01 post-validación)'),
    ('RD-11', 'El sistema debe gestionar roles y permisos (RBAC) diferenciados para Administrador, Recepcionista, Socio y Entrenador, restringiendo el acceso a módulos según el rol.', 'Alta', 'Benchmarking + Entrevista'),
    ('RD-12', 'El código QR del socio debe expirar automáticamente cada 12 horas y regenerarse para garantizar que no pueda ser compartido o reutilizado indefinidamente.', 'Alta', 'Benchmarking (Mindbody) / Equipo de desarrollo'),
]
for rd in rds:
    r = rd_t.add_row()
    for c, t in zip(r.cells, rd): c.text = t
    if rd[2] == 'Alta':
        set_cell_bg(r.cells[2], 'FFD7D7')
    elif rd[2] == 'Media':
        set_cell_bg(r.cells[2], 'FFF3CD')
doc.add_paragraph()

# 2.2 Especificación formal RD-01
doc.add_heading('2.2. Especificación Formal de un Requisito de Dominio Prioritario', 2)
doc.add_paragraph('A continuación se especifica el requisito de dominio de mayor impacto utilizando la plantilla estándar:')
spec_table(doc, [
    ('ID',                    'RD-01'),
    ('Nombre',                'Control de Acceso Físico mediante Código QR'),
    ('Descripción',           'El sistema debe verificar, en tiempo real y en menos de 300 ms, si el socio que intenta ingresar al gimnasio posee una membresía vigente y sin penalización activa mediante la lectura de un código QR dinámico. Si la membresía está activa, se envía señal de apertura al torniquete; si está vencida o bloqueada, el torniquete permanece cerrado y se muestra un mensaje diferenciado.'),
    ('Procedencia',           'Dueño/Administrador — Hallazgos de entrevista y observación directa en recepción'),
    ('Prioridad',             'Alta — impacto económico directo (pérdida de ingresos por accesos indebidos con membresía vencida)'),
    ('Criterios de aceptación', '1. La validación completa (QR → BD → torniquete) tarda menos de 300 ms en el 95% de los casos. '
                                '| 2. Los socios con membresía vencida no pueden abrir el torniquete. '
                                '| 3. Los socios con bloqueo activo reciben un mensaje diferenciado con fecha de desbloqueo. '
                                '| 4. Cada intento de acceso queda registrado en el log con socio, fecha/hora y resultado.'),
    ('Dependencias',          'RD-02 (Gestión de planes), RD-04 (Penalización), RD-12 (Expiración QR), RF-06 (Validación QR), RF-07 (Dashboard), RNF-01 (Rendimiento < 500 ms), RNF-02 (QR < 300 ms)'),
])

# 2.3 Especificación formal RD-04
doc.add_heading('2.3. Especificación Formal — RD-04', 2)
spec_table(doc, [
    ('ID',                    'RD-04'),
    ('Nombre',                'Política Automática de Penalización por Inasistencia'),
    ('Descripción',           'Cuando un socio no se presenta a una clase reservada y no la cancela dentro del margen permitido (2 horas antes), el sistema registra automáticamente 1 strike al finalizar la hora de la clase. Al acumular 3 strikes en un período de 30 días, el sistema bloquea la capacidad de reservar clases durante 7 días calendario. El Administrador puede remover un strike con justificación documentada (RD-10).'),
    ('Procedencia',           'Administrador — Entrevista semiestructurada + acuerdo SC-01 de validación'),
    ('Prioridad',             'Alta'),
    ('Criterios de aceptación', '1. El sistema detecta la inasistencia automáticamente al finalizar la hora de la clase. '
                                '| 2. El contador de strikes se evalúa dentro de una ventana de 30 días. '
                                '| 3. Al alcanzar 3 strikes, el sistema muestra la fecha exacta de desbloqueo. '
                                '| 4. El Administrador puede remover un strike manualmente con justificación registrada (RF-17).'),
    ('Dependencias',          'RD-03 (Control de aforo), RD-07 (Cancelación 2h), RF-09 (Reserva de clases), RF-14 (Penalización automática), RF-17 (Justificación de strike)'),
])

# 2.4 Diagrama de Clases del Dominio
doc.add_heading('2.4. Diagrama de Clases del Dominio', 2)
doc.add_paragraph(
    'El siguiente diagrama representa las entidades del dominio con sus atributos, '
    'relaciones y las reglas de negocio anotadas en las notas de cada clase. '
    'Refleja el modelo de datos implementado en PostgreSQL para el MVP1:'
)
add_diagram(doc, UML_CLASES_DOMINIO, 'Figura 2.4: Diagrama de Clases del Dominio — Modelo de Datos OLYMPUS CORE (MVP1)', 6.2)

# 2.5 Relaciones y Dependencias entre Requisitos
doc.add_heading('2.5. Relaciones y Dependencias entre Requisitos', 2)
doc.add_paragraph(
    'La siguiente tabla muestra cómo los Requisitos de Dominio fundamentan los '
    'Requisitos Funcionales y No Funcionales del sistema MVP1 actualizado:'
)
dep_t = doc.add_table(rows=1, cols=4)
dep_t.style = 'Table Grid'
header_row(dep_t, ['RD', 'Nombre', 'RF que fundamenta', 'RNF que fundamenta'])
deps = [
    ('RD-01',  'Control de Acceso QR',            'RF-06, RF-07, RF-04',       'RNF-01, RNF-02, RNF-05, RNF-06'),
    ('RD-02',  'Catálogo cerrado de planes',       'RF-03, RF-05, RF-08',       'RNF-04, RNF-07'),
    ('RD-03',  'Control de aforo en tiempo real',  'RF-09, RF-16, RF-18',       'RNF-01, RNF-03'),
    ('RD-04',  'Política de penalización',         'RF-14, RF-09, RF-17',       'RNF-07'),
    ('RD-05',  'Ficha médica obligatoria',         'RF-12, RF-01',              '—'),
    ('RD-06',  'Dashboard gerencial',              'RF-07, RF-15',              'RNF-01, RNF-02'),
    ('RD-07',  'Política de cancelación (2 h)',    'RF-09, RF-14',              '—'),
    ('RD-08',  'Historial de progreso físico',     'RF-12, RF-13',              '—'),
    ('RD-09',  'Log de auditoría de accesos',      'RF-06, RF-07',              'RNF-07'),
    ('RD-10',  'Justificación de strike (Admin)',  'RF-17',                     '—'),
    ('RD-11',  'Control de roles RBAC',            'RF-01, RF-11, RF-10',       'RNF-04, RNF-05, RNF-06'),
    ('RD-12',  'Expiración QR (12 horas)',         'RF-06, RF-10',              'RNF-06'),
]
for d in deps:
    r = dep_t.add_row()
    for c, t in zip(r.cells, d): c.text = t
doc.add_paragraph()

# 2.6 Glosario
doc.add_heading('2.6. Glosario de Términos del Dominio', 2)
glos_t = doc.add_table(rows=1, cols=2)
glos_t.style = 'Table Grid'
header_row(glos_t, ['Término', 'Definición en el contexto del dominio'])
glosario = [
    ('Membresía',           'Contrato de acceso entre el gimnasio y el socio que habilita el uso de las instalaciones durante un período determinado (mensual/trimestral/anual). Puede estar activa, vencida o bloqueada.'),
    ('Socio / Afiliado',    'Persona registrada en el sistema con un plan de membresía activo o histórico. Tiene un usuario en el sistema y un código QR único.'),
    ('Strike',              'Penalización unitaria aplicada por inasistencia injustificada a una clase reservada. La acumulación de 3 en 30 días genera bloqueo temporal.'),
    ('Aforo',               'Capacidad máxima de participantes permitidos en una clase grupal. Definida por el entrenador para cada clase.'),
    ('Código QR Dinámico',  'Código bidimensional único por socio, regenerado automáticamente cada 12 horas, que contiene su token JWT de identificación para el control de acceso.'),
    ('Clase Grupal',        'Actividad deportiva estructurada con cupo limitado (Spinning, CrossFit) que los socios pueden reservar con anticipación a través del sistema.'),
    ('RBAC',                'Role-Based Access Control. Control de acceso basado en roles. Cada rol (admin, recepcionista, socio, entrenador) tiene permisos diferenciados en el sistema.'),
    ('JWT',                 'JSON Web Token. Token de autenticación firmado digitalmente que contiene la identidad y rol del usuario. Expira en 24 horas para sesiones y 12 horas para QR.'),
    ('Dashboard Gerencial', 'Panel de control administrativo con métricas en tiempo real: ingresos del día, membresías por vencer (3 días) e índice de ocupación de clases.'),
    ('Ficha Médica',        'Formulario electrónico que registra lesiones previas, enfermedades cardiovasculares y objetivo de entrenamiento del socio. Prerrequisito para asignar rutinas.'),
    ('Log de Acceso',       'Registro auditable e inmutable de cada intento de acceso al gimnasio: socio, fecha/hora, resultado (PERMITIDO/DENEGADO) y motivo.'),
    ('Torniquete',          'Hardware de control de acceso físico que recibe señal del sistema (abrir / mantener cerrado) según el resultado de la validación de membresía.'),
    ('Index Scan',          'Tipo de consulta optimizada en PostgreSQL que usa índices para acceder a datos sin recorrer toda la tabla. Requerido para garantizar latencia < 300 ms en validación QR (RD-01).'),
]
for g in glosario:
    r = glos_t.add_row()
    r.cells[0].text = g[0]
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[1].text = g[1]
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 3: Validación de Requisitos de Dominio
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Validación de Requisitos de Dominio', 1)

# 3.1 Plan de Revisión
doc.add_heading('3.1. Plan de Revisión', 2)
doc.add_paragraph(
    'Siguiendo el enfoque de validación por stakeholders y equipos multidisciplinarios, '
    'se propone la siguiente agenda para una sesión de revisión de 2 horas:'
)
plan_t = doc.add_table(rows=1, cols=4)
plan_t.style = 'Table Grid'
header_row(plan_t, ['Bloque', 'Actividad', 'Artefactos revisados', 'Participantes'])
plan = [
    ('0:00–0:20', 'Presentación del dominio actualizado y objetivos del MVP1', 'Descripción general (sección 1.1), lista de actores y 12 RD', 'Todos'),
    ('0:20–0:50', 'Revisión de Reglas de Negocio (RN-01 a RN-10)', 'Tabla de RN, diagrama de clases', 'Dueño, Recepcionista, Entrenador'),
    ('0:50–1:20', 'Revisión de Requisitos de Dominio (RD-01 a RD-12)', 'Tabla de RD, especificaciones formales RD-01 y RD-04', 'Todos + Analista técnico'),
    ('1:20–1:40', 'Validación de diagramas UML y prototipos', 'DCU negocio/sistema, diagrama de clases, swimlanes, secuencia', 'Analista, Desarrollador líder'),
    ('1:40–2:00', 'Acuerdos, acta y próximos pasos', 'Lista de observaciones y acta de validación', 'Todos'),
]
for row in plan:
    r = plan_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

# 3.2 Preguntas de Validación
doc.add_heading('3.2. Preguntas de Validación por Stakeholder', 2)

doc.add_heading('3.2.1. Dueño / Administrador', 3)
for q in [
    'RD-01: ¿El tiempo de 300 ms para la validación de acceso es suficiente para la operación en hora pico con filas de hasta 20 personas?',
    'RD-06: ¿El dashboard con las 3 métricas propuestas (ingresos, vencimientos, ocupación) cubre su necesidad de control matutino o requiere alguna adicional?',
    'RD-04: ¿El período de cálculo de strikes (30 días) y el bloqueo de 7 días son las métricas correctas, o desea ajustarlos?',
    'RD-11: ¿Los 4 roles definidos (admin, recepcionista, socio, entrenador) cubren todos los perfiles de acceso del negocio?',
]:
    add_bullet(doc, q)

doc.add_heading('3.2.2. Recepcionista', 3)
for q in [
    'RD-02: ¿El catálogo cerrado de planes le impide realizar alguna excepción que actualmente necesita aplicar (descuentos, planes especiales)?',
    'RD-07: ¿La regla de cancelación hasta 2 horas antes genera algún conflicto con situaciones frecuentes que enfrenta en la recepción?',
    'RD-01: ¿El proceso de validación QR requiere algún paso manual de respaldo cuando el sistema no está disponible?',
]:
    add_bullet(doc, q)

doc.add_heading('3.2.3. Entrenador', 3)
for q in [
    'RD-05: ¿Los campos de la ficha médica digital (lesiones, enfermedades cardiovasculares, objetivo) son suficientes o falta alguna información clínica relevante?',
    'RD-08: ¿El registro de progreso físico con pesos y observaciones cubre su necesidad de seguimiento, o requiere métricas adicionales (talla, IMC, porcentaje de grasa)?',
    'RD-03: ¿El sistema de gestión de aforo y horarios de clases grupales refleja correctamente su flujo de trabajo actual?',
]:
    add_bullet(doc, q)

# 3.3 Análisis de Hallazgos
doc.add_heading('3.3. Análisis de Hallazgos — Simulación de Revisión', 2)
doc.add_paragraph('Se simulan los hallazgos sobre los requisitos de dominio clave, identificando problemas potenciales:')
hall_t = doc.add_table(rows=1, cols=3)
hall_t.style = 'Table Grid'
header_row(hall_t, ['ID', 'Problema identificado / Stakeholder que objeta', 'Mejora propuesta / Decisión'])
hallazgos = [
    ('RD-04', 'Entrenador: los socios con lesiones médicas no deben ser penalizados del mismo modo que un abandono injustificado.', 'Aprobado (SC-01): agregar mecanismo de justificación de inasistencia validada por entrenador o médico → RD-10 y RF-17.'),
    ('RD-02', 'Administrador: existen convenios con empresas que requieren precios diferenciados temporalmente.', 'Aprobado (SC-02): incorporar módulo de promociones/convenios administrable solo por el Administrador → nuevo RD derivado en v1.1.'),
    ('RD-06', 'Desarrollador líder: sin período de actualización definido no se puede garantizar SLA de disponibilidad.', 'Acordado: el dashboard se actualiza cada 60 segundos y muestra datos en caché si hay falla (RQ-CAL-OLY-11).'),
    ('RD-12', 'Equipo de seguridad: un QR sin expiración puede ser compartido o capturado.', 'Aprobado: QR expira cada 12 horas automáticamente. Se añade como RD-12 y RNF-06.'),
]
for h in hallazgos:
    r = hall_t.add_row()
    for c, t in zip(r.cells, h): c.text = t
doc.add_paragraph()

# 3.4 Validación mediante Prototipos
doc.add_heading('3.4. Validación mediante Prototipos — Flujo de Reserva de Clase', 2)
doc.add_paragraph(
    'Para validar el flujo de reserva de clase grupal (RD-03, RD-04, RD-07) se propone '
    'un prototipo de baja fidelidad con las siguientes pantallas:'
)
for pt in [
    'Pantalla 1 — Lista de clases: Muestra clases disponibles con horario, tipo, aforo restante y estado (disponible/llena).',
    'Pantalla 2 — Detalle de clase: Muestra instructor, duración, descripción y botón "Reservar".',
    'Pantalla 3 — Confirmación: Resumen de la reserva con advertencia de la política de cancelación (2 h) y de strikes.',
    'Pantalla 4 — Estado de reservas: Lista de reservas activas del socio con opción de cancelar (solo si aplica el margen de 2 h).',
    'Pantalla 5 — Alerta de penalización: Pantalla modal que bloquea la reserva y muestra strikes acumulados y fecha de desbloqueo.',
]:
    add_bullet(doc, pt)

# 3.5 Criterios de Éxito
doc.add_heading('3.5. Criterios de Éxito de la Validación', 2)
crit_t = doc.add_table(rows=1, cols=3)
crit_t.style = 'Table Grid'
header_row(crit_t, ['Criterio', 'Descripción', 'Métrica de aceptación'])
criterios = [
    ('Corrección',    'Los RD reflejan fielmente las reglas del negocio real del gimnasio.',                '0 objeciones críticas de parte del Dueño'),
    ('Completitud',   'No quedan procesos de negocio críticos sin cubrir.',                                'Cobertura ≥ 90% de los procesos mapeados'),
    ('Viabilidad',    'Los RD son técnicamente implementables en el plazo del proyecto.',                   'Validación del desarrollador líder sin bloqueos'),
    ('Trazabilidad',  'Cada RD está asociado a al menos un RF o RNF en la matriz.',                        '100% de RD con dependencia mapeada'),
    ('Consistencia',  'Los RD son coherentes con los 34 RF, 13 RNF y 13 RQ-CAL del proyecto.',            'Sin contradicciones entre documentos'),
]
for c in criterios:
    r = crit_t.add_row()
    for cell, t in zip(r.cells, c): cell.text = t
doc.add_paragraph()

# 3.6 Acta de Validación
doc.add_heading('3.6. Acta de Validación', 2)
doc.add_paragraph(
    'Fecha: 12/06/2026 | Sesión de revisión de RD — OLYMPUS CORE | Duración: 2 horas\n'
    'Participantes: Dueño/Administrador, Recepcionista (representante), Entrenador en Jefe, '
    'Analista de Requisitos (Grupo 5).'
)
doc.add_paragraph('Acuerdos alcanzados:')
for ac in [
    'RD-01 a RD-03 y RD-06: Aprobados sin modificaciones. El Administrador confirmó que las métricas y tiempos propuestos son adecuados.',
    'RD-04 (Penalizaciones): Modificación acordada — agregar mecanismo de justificación de inasistencia médica → SC-01. Se crea RF-17 (Justificar Strike) como consecuencia.',
    'RD-02 (Catálogo de precios): Se aprueba con adición de módulo de promociones → SC-02. Se crea requisito derivado en v1.1.',
    'RD-12 (Expiración QR): Nuevo requisito aprobado tras benchmarking de seguridad → RNF-06 y RQ-CAL-OLY-06.',
    'RD-11 (RBAC): Aprobado con los 4 roles definidos como suficientes para MVP1.',
]:
    add_bullet(doc, ac)
doc.add_paragraph('Desacuerdos pendientes:')
for d in [
    'RD-08: El entrenador solicita incluir métricas de IMC y porcentaje de grasa. El equipo evaluará viabilidad técnica en la próxima sesión → SC-03.',
]:
    add_bullet(doc, d)
doc.add_paragraph(
    'Próximos pasos: (1) Actualizar RD-04 con mecanismo de justificación → RF-17. '
    '(2) Especificar RD-10 (módulo de promociones) en v1.1. '
    '(3) Presentar prototipo de baja fidelidad del flujo de reserva en 2 semanas.'
)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 4: Diagramas Dinámicos del Dominio
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Diagramas Dinámicos del Dominio', 1)

doc.add_heading('4.1. Diagrama de Actividades con Swimlanes — Reserva de Clase Grupal', 2)
doc.add_paragraph(
    'Este diagrama modela el proceso de reserva de clase grupal integrando '
    'las validaciones de RD-03 (aforo), RD-04 (penalización) y RD-07 (cancelación 2h). '
    'Muestra las responsabilidades del Socio y del Sistema en cada etapa:'
)
add_diagram(doc, UML_ACTIVIDAD_RESERVA, 'Figura 4.1: Diagrama de Actividades con Swimlanes — Reserva de Clase (RD-03, RD-04, RD-07)', 6.0)

doc.add_heading('4.2. Diagrama de Secuencia — Validación de Acceso QR', 2)
doc.add_paragraph(
    'Este diagrama muestra la interacción entre los componentes del sistema durante '
    'el proceso de control de acceso físico (RD-01). Incluye anotaciones de '
    'los requisitos de calidad que afectan a este flujo crítico:'
)
add_diagram(doc, UML_SECUENCIA_QR, 'Figura 4.2: Diagrama de Secuencia — Validación de Acceso QR (RD-01 | RF-06 | RQ-CAL-OLY-01)', 6.0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 5: Gestión de Cambios en los Requisitos de Dominio
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Gestión de Cambios en los Requisitos de Dominio', 1)

# 5.1 Proceso Formal
doc.add_heading('5.1. Proceso Formal de Solicitud de Cambio', 2)
doc.add_paragraph(
    'Cualquier modificación a los requisitos de dominio debe seguir el siguiente proceso controlado:'
)
for paso in [
    'Identificación: El stakeholder documenta el cambio en el formulario de Solicitud de Cambio (SC) con ID único.',
    'Evaluación de impacto: El analista evalúa qué RD, RF, RNF, RQ-CAL y artefactos se ven afectados.',
    'Priorización: El cambio se clasifica como Alta/Media/Baja según impacto en negocio y riesgo técnico.',
    'Aprobación: El Dueño y el Analista líder aprueban o rechazan el cambio con justificación documentada.',
    'Implementación y revalidación: El cambio aprobado se incorpora en la línea base y se revalida con los stakeholders afectados. Se actualiza la matriz de trazabilidad.',
]:
    add_bullet(doc, paso)

# 5.2 Registro de Cambios
doc.add_heading('5.2. Registro de Cambios Identificados', 2)
sc_t = doc.add_table(rows=1, cols=5)
sc_t.style = 'Table Grid'
header_row(sc_t, ['SC-ID', 'Cambio solicitado', 'RD afectado', 'Prioridad', 'Decisión'])
scs = [
    ('SC-01', 'Agregar justificación de inasistencia médica',              'RD-04', 'Alta',  'Aprobado — RF-17 creado, RD-10 añadido'),
    ('SC-02', 'Módulo de promociones y convenios corporativos',            'RD-02', 'Media', 'Aprobado — nuevo requisito en v1.1'),
    ('SC-03', 'Incluir IMC y % grasa en progreso físico',                  'RD-08', 'Baja',  'Pendiente — evaluación técnica v2.0'),
    ('SC-04', 'Notificaciones automáticas por WhatsApp',                   'RD-03, RD-04', 'Media', 'Diferido a versión 2.0 del sistema'),
    ('SC-05', 'Expiración automática de código QR cada 12 horas',          'RD-01', 'Alta',  'Aprobado — RD-12 creado, RNF-06 añadido'),
    ('SC-06', 'Control de roles RBAC diferenciado por módulo',             'RD-01', 'Alta',  'Aprobado — RD-11 creado, RF-11 añadido'),
]
for sc in scs:
    r = sc_t.add_row()
    for c, t in zip(r.cells, sc): c.text = t
    if sc[3] == 'Alta':
        set_cell_bg(r.cells[3], 'FFD7D7')
    elif sc[3] == 'Media':
        set_cell_bg(r.cells[3], 'FFF3CD')
    elif sc[3] == 'Baja':
        set_cell_bg(r.cells[3], 'D4EDDA')
doc.add_paragraph()

# ── GUARDAR ────────────────────────────────────────────────────────────────────
output = (
    r'c:/Users/ACER/OneDrive/Desktop/PROGRAMACION/UNIVERSITY PROJECTS/'
    r'SIstemaGestorGimnasio/Documentacion/Documentos Actualizados/'
    r'Requisitos_Dominio_OLYMPUS_CORE_actualizado.docx'
)
doc.save(output)
print(f'Documento guardado en: {output}')
