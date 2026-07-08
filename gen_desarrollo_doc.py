"""
gen_desarrollo_doc.py
Genera: REQUISITOS_DESARROLLO_actualizado.docx
Mantiene la estructura del documento original (5 secciones + conclusiones).
- Tecnologia de produccion realista (sin Vercel/Neon).
- Alineado con 34 RF, 13 RNF, 13 RQ-CAL y los RD actualizados.
- Incluye diagramas UML embebidos desde la API de PlantUML.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests, zlib, io

# ══════════════════════════════════════════════════════════════════════════════
# PLANTUML helpers
# ══════════════════════════════════════════════════════════════════════════════
def _encode6bit(b):
    if b < 10: return chr(48 + b)
    b -= 10
    if b < 26: return chr(65 + b)
    b -= 26
    if b < 26: return chr(97 + b)
    b -= 26
    return '-' if b == 0 else '_'

def _append3bytes(b1, b2, b3):
    return (
        _encode6bit(b1 >> 2) +
        _encode6bit(((b1 & 0x3) << 4) | (b2 >> 4)) +
        _encode6bit(((b2 & 0xF) << 2) | (b3 >> 6)) +
        _encode6bit(b3 & 0x3F)
    )

def plantuml_encode(text):
    data = zlib.compress(text.encode('utf-8'))[2:-4]
    res = ""
    for i in range(0, len(data), 3):
        if i + 2 < len(data): res += _append3bytes(data[i], data[i+1], data[i+2])
        elif i + 1 < len(data): res += _append3bytes(data[i], data[i+1], 0)
        else: res += _append3bytes(data[i], 0, 0)
    return res

def fetch_diagram(uml_code, width_inches=5.8):
    try:
        enc = plantuml_encode(uml_code)
        r = requests.get(f"http://www.plantuml.com/plantuml/png/{enc}", timeout=25)
        if r.status_code == 200 and r.content[:4] == b'\x89PNG':
            return io.BytesIO(r.content), width_inches
    except Exception:
        pass
    return None

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, cols, bg='1E3A5F', fg='FFFFFF'):
    for cell, txt in zip(table.rows[0].cells, cols):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        set_cell_bg(cell, bg)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

def add_diagram(doc, uml_code, caption, width=5.8):
    result = fetch_diagram(uml_code, width)
    if result:
        img_stream, w = result
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_stream, width=Inches(w))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph(f'[Diagrama: {caption}]').italic = True
    doc.add_paragraph()

def two_col_table(doc, rows_data, key_bg='D0D8E4'):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    for i, (key, val) in enumerate(rows_data):
        ck, cv = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        ck.text = key
        ck.paragraphs[0].runs[0].bold = True
        set_cell_bg(ck, key_bg)
        cv.text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMAS UML
# ══════════════════════════════════════════════════════════════════════════════

# Arquitectura de produccion (Componentes)
UML_ARQUITECTURA = """@startuml
skinparam componentStyle uml2
skinparam component {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}

title Arquitectura de Produccion — OLYMPUS CORE\\n(Stack tecnologico realista)

node "Servidor de Aplicacion\\n(Linux VPS / Ubuntu 22.04)" {
  component "Nginx\\n(Reverse Proxy + SSL/TLS)" as NGINX
  component "PM2\\n(Process Manager)" as PM2
  component "Backend\\nNode.js 20 LTS + Express" as BE {
    [Auth Controller (JWT)]
    [Socios Controller]
    [Reservas Controller (Transacciones)]
    [QR Controller]
    [Dashboard Controller]
  }
  component "Frontend\\nReact 18 + Vite\\n(Build estatico servido por Nginx)" as FE
}

node "Servidor de Base de Datos\\n(PostgreSQL 16 — mismo VPS o VPS dedicado)" {
  database "PostgreSQL 16" as DB {
    [socios]
    [membresias]
    [clases_grupales]
    [reservas]
    [pagos]
    [logs_acceso]
  }
}

node "Red Local del Gimnasio" {
  [Lector QR / Tablet\\nRecepcion] as QR
  [PC Administrador] as PC
  [Torniquete (Hardware)] as TORN
}

node "GitHub" {
  [Repositorio Git\\n(CI/CD con GitHub Actions)] as GIT
}

PC --> NGINX : HTTPS (LAN + Internet)
QR --> NGINX : HTTPS (LAN)
NGINX --> FE : Sirve archivos estaticos
NGINX --> BE : Proxy reverso /api/*
BE --> DB : TCP puerto 5432\\n(Pool de conexiones - pg)
BE --> TORN : Señal HTTP/Serial (apertura)
GIT --> BE : Deploy automatico\\nvía SSH + PM2 reload
@enduml"""

# CI/CD Pipeline (Actividad)
UML_CICD = """@startuml
skinparam activity {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}
skinparam roundcorner 15

title Pipeline CI/CD — GitHub Actions\\nOLYMPUS CORE

start
:Developer hace git push\\na rama feature/;

fork
  :Trigger: GitHub Actions Workflow;
fork again
  :Code Review por compañero\\n(Pair Programming - XP);
end fork

:Stage 1 — Lint y formato\\n(ESLint + Prettier);
if (Lint OK?) then (no)
  :Notificar error al developer;
  stop
else (si)
endif

:Stage 2 — Tests unitarios\\n(Jest + Supertest);
if (Tests 100% OK?) then (no)
  :Notificar fallo de tests;
  stop
else (si)
endif

:Stage 3 — Tests de integracion\\n(endpoints con BD de prueba);
if (Integracion OK?) then (no)
  :Notificar fallo;
  stop
else (si)
endif

:PR aprobado — merge a main;

:Stage 4 — Build de produccion;
:Stage 5 — Deploy via SSH al VPS;
note right: PM2 reload (zero-downtime)\\nnginx -s reload

:Stage 6 — Health check del endpoint\\nGET /api/health -> 200 OK;
if (Health check OK?) then (no)
  :Rollback automatico\\nal commit anterior;
  stop
else (si)
endif

:Deploy exitoso;
:Notificar al equipo via email;
stop
@enduml"""

# Concurrencia en reservas (Secuencia)
UML_CONCURRENCIA = """@startuml
skinparam sequenceArrowThickness 2
skinparam roundcorner 10

title TEC-001: Control de Concurrencia en Reservas\\nTransacciones PostgreSQL con SELECT FOR UPDATE

participant "Socio A\\n(Request 1)" as A
participant "Socio B\\n(Request 2)" as B
participant "API Backend\\n(Node.js)" as API
database "PostgreSQL 16\\n(MVCC + Row-Level Lock)" as DB

== Peticiones concurrentes al mismo endpoint ==

A -> API: POST /api/reservas {clase_id: 42}
B -> API: POST /api/reservas {clase_id: 42}

== Request A adquiere el lock primero ==

API -> DB: BEGIN TRANSACTION
API -> DB: SELECT aforo_actual, aforo_maximo\\nFROM clases_grupales\\nWHERE id = 42\\nFOR UPDATE NOWAIT
note right of DB: Row-Level Lock adquirido\\npor Transaction A

DB --> API: aforo_actual=19, aforo_maximo=20
API -> DB: INSERT INTO reservas (socio_id=A, clase_id=42)
API -> DB: UPDATE clases_grupales\\nSET aforo_actual = 20\\nWHERE id = 42
API -> DB: COMMIT
DB --> API: 201 CREATED

== Request B intenta adquirir el mismo lock ==

API -> DB: BEGIN TRANSACTION (B)
API -> DB: SELECT ... FOR UPDATE NOWAIT\\n(clase_id = 42)
DB --> API: ERROR: could not obtain lock — aforo_maximo alcanzado
API -> DB: ROLLBACK
DB --> API: 409 CLASE_SIN_CUPOS

API --> A: HTTP 201 — Reserva creada exitosamente
API --> B: HTTP 409 — Clase sin cupos disponibles

note bottom
  RNF-03: 100 usuarios concurrentes sin sobreventa
  TEC-001: Exactamente 20 HTTP 201 y N-20 HTTP 409
  RQ-CAL-OLY-03: Latencia P95 < 800 ms bajo 100 VUs
end note
@enduml"""

# Validacion QR — Estado
UML_ESTADO_QR = """@startuml
skinparam state {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}
skinparam roundcorner 15

title TEC-002: Estados del Token QR durante la Validacion de Acceso\\n(RF-06, RNF-02, RNF-06)

[*] --> GeneradoFresco : Socio inicia sesion\\nJWT firmado (exp: +12h)

GeneradoFresco : Token QR valido
GeneradoFresco : Firmado con HS256
GeneradoFresco : Payload: socio_id, iat, exp

GeneradoFresco --> EnValidacion : Recepcionista\\nescanea QR

EnValidacion : POST /api/acceso/validar {token}
EnValidacion : Verificando firma JWT
EnValidacion : Consultando BD [Index Scan]

EnValidacion --> Permitido : Membresia activa\\nSin penalizacion\\n(< 300 ms)
EnValidacion --> DenegadoVencido : Membresia vencida
EnValidacion --> DenegadoBloqueado : Strike activo (3/30 dias)
EnValidacion --> DenegadoTokenInvalido : Firma JWT invalida\\no token expirado

Permitido : INSERT logs_acceso (PERMITIDO)
Permitido : Señal apertura torniquete
Permitido --> GeneradoFresco : Listo para\\nproximo acceso

DenegadoVencido : INSERT logs_acceso (DENEGADO — VENCIDA)
DenegadoBloqueado : INSERT logs_acceso (DENEGADO — BLOQUEADO)
DenegadoTokenInvalido : INSERT logs_acceso (DENEGADO — TOKEN_INVALIDO)

DenegadoVencido --> [*]
DenegadoBloqueado --> [*]
DenegadoTokenInvalido --> [*]

note right of GeneradoFresco
  RNF-06: Token QR expira cada 12 horas
  RD-12: Regenerado automaticamente
end note
@enduml"""

# Stack tecnologico (Clases)
UML_STACK = """@startuml
hide circle
skinparam classAttributeIconSize 0
skinparam class {
  BackgroundColor #F0F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}

title Stack Tecnologico de Produccion — OLYMPUS CORE\\n(Capas del sistema)

class "Capa de Presentacion" as FRONT {
  framework = "React 18 + Vite 5"
  routing = "React Router DOM 6"
  estilos = "CSS Modules / Vanilla CSS"
  http_client = "Axios"
  estado = "React Context + useState"
  build = "npm run build -> dist/"
  servidor = "Nginx (archivos estaticos)"
}

class "Capa de API y Logica" as BACK {
  runtime = "Node.js 20 LTS"
  framework = "Express 4.x"
  autenticacion = "jsonwebtoken (JWT HS256)"
  cifrado = "bcryptjs (rounds: 12)"
  validacion = "Joi / express-validator"
  process_manager = "PM2 (cluster mode)"
  logs = "Winston (INFO/WARN/ERROR)"
}

class "Capa de Datos" as DATA {
  sgbd = "PostgreSQL 16"
  driver = "node-postgres (pg)"
  conexion = "Pool (max: 20 conexiones)"
  migraciones = "node-pg-migrate"
  indices = "dni, socio_id, estado, fecha_fin"
  transacciones = "SELECT FOR UPDATE NOWAIT"
}

class "Infraestructura" as INFRA {
  so = "Ubuntu 22.04 LTS"
  servidor = "VPS (2 vCPU, 4GB RAM, 50GB SSD)"
  proxy = "Nginx 1.24 (reverse proxy + SSL)"
  ssl = "Let's Encrypt (certbot)"
  ci_cd = "GitHub Actions"
  repositorio = "GitHub (Git Flow)"
  monitoring = "UptimeRobot + logs PM2"
}

FRONT --> BACK : HTTPS REST API (JSON)
BACK --> DATA : TCP 5432 (Pool)
INFRA --> BACK : Gestiona proceso (PM2)
INFRA --> FRONT : Sirve archivos (Nginx)
INFRA --> DATA : Backup diario (cron)
@enduml"""

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCION DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── PORTADA ────────────────────────────────────────────────────────────────────
for text, size, bold, color in [
    ('UNIVERSIDAD NACIONAL MAYOR DE SAN MARCOS', 14, True,  (0x1E,0x3A,0x5F)),
    ('FACULTAD DE SISTEMAS E INFORMÁTICA',       13, False, (0x1E,0x3A,0x5F)),
    ('', 11, False, (0,0,0)),
    ('REQUISITOS DE DESARROLLO', 18, True, (0x1E,0x3A,0x5F)),
    ('', 11, False, (0,0,0)),
    ('ASIGNATURA', 11, False, (0,0,0)),
    ('Ingeniería de Requisitos', 13, True, (0x1E,0x3A,0x5F)),
    ('', 11, False, (0,0,0)),
    ('DOCENTE', 11, False, (0,0,0)),
    ('Rodríguez Rodríguez, Ciro', 12, False, (0,0,0)),
    ('', 11, False, (0,0,0)),
    ('INTEGRANTES GRUPO 5:', 11, True, (0,0,0)),
    ('Alva Chacon, Jose Benjamin          24200045', 11, False, (0,0,0)),
    ('Cordova Guerra, Josue Rodrigo       24200155', 11, False, (0,0,0)),
    ('Sandoval Dominguez, Erick Marco     24200172', 11, False, (0,0,0)),
    ('Melendez Bustamante, Alvaro Mathias 24200166', 11, False, (0,0,0)),
    ('', 11, False, (0,0,0)),
    ('Proyecto: OLYMPUS CORE — Sistema Gestor de Gimnasio', 11, False, (0,0,0)),
    ('2026', 12, False, (0,0,0)),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introducción', 1)
doc.add_paragraph(
    'Durante la fase de identificación y análisis de requisitos de desarrollo del proyecto '
    'OLYMPUS CORE, el equipo evaluó las restricciones técnicas, metodológicas y operativas '
    'necesarias para garantizar la construcción de una plataforma robusta, segura, escalable '
    'y alineada a las necesidades del gimnasio.\n\n'
    'El análisis permitió identificar factores críticos relacionados con concurrencia en '
    'reservas de clases, validación de acceso mediante código QR dinámico, rendimiento del '
    'sistema, seguridad de la información, compatibilidad tecnológica y despliegue en '
    'infraestructura de producción realista. Asimismo, se definieron estrategias de mitigación '
    'de riesgos y mecanismos de validación técnica orientados a garantizar la calidad del '
    'producto final.\n\n'
    'El desarrollo del proyecto se basa en metodologías ágiles (Extreme Programming) y '
    'herramientas estándar de la industria, priorizando la construcción incremental del MVP1, '
    'la trazabilidad de cambios y la validación continua de requisitos. El stack tecnológico '
    'seleccionado corresponde a tecnologías de uso real en producción: Node.js, PostgreSQL, '
    'React y Nginx sobre un VPS Linux, asegurando que la plataforma pueda operar de forma '
    'autónoma y segura en el entorno del gimnasio.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. IDENTIFICACIÓN Y ANÁLISIS DE REQUISITOS DE DESARROLLO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Identificación y Análisis de Requisitos de Desarrollo', 1)

doc.add_heading('2.1. Mapeo de Restricciones del Proyecto', 2)
doc.add_paragraph(
    'Durante esta etapa se identificaron las restricciones técnicas, temporales, '
    'metodológicas y operativas que condicionan el desarrollo del sistema:'
)
rest_t = doc.add_table(rows=1, cols=3)
rest_t.style = 'Table Grid'
header_row(rest_t, ['Tipo de restricción', 'Descripción de la restricción', 'Justificación / Fuente'])
restricciones = [
    ('Tecnológica',   'El sistema requiere hardware específico en recepción: lectores de códigos QR (smartphones o lectores dedicados) y torniquete físico con interfaz HTTP o serial.',
                      'Análisis de requisitos del gimnasio — RD-01, RF-06'),
    ('Tecnológica',   'El SGBD será estrictamente PostgreSQL 16 para asegurar integridad relacional, soporte ACID, transacciones con SELECT FOR UPDATE y manejo eficiente de concurrencia.',
                      'Decisión técnica del equipo — RNF-07, TEC-001'),
    ('Tecnológica',   'El servidor de producción será un VPS Linux (Ubuntu 22.04 LTS) gestionado con Nginx como reverse proxy y PM2 como process manager para Node.js. Se excluyen plataformas PaaS gratuitas (Heroku free tier, etc.) por limitaciones de tiempo de inactividad y recursos.',
                      'Requisito de disponibilidad — RQ-CAL-OLY-11 (99% uptime)'),
    ('Tecnológica',   'La autenticación se implementará con JWT (HS256) y bcrypt (factor 12). Las sesiones duran 24 horas y los tokens QR expiran en 12 horas. No se almacenan contraseñas en texto plano.',
                      'RNF-04, RNF-05, RNF-06 — Seguridad'),
    ('Temporal',      'El tiempo máximo para el desarrollo y entrega del MVP1 está fijado en 5 meses. Las iteraciones semanales de XP garantizan entregables funcionales al final de cada sprint.',
                      'Límite académico del proyecto'),
    ('Proceso',       'El desarrollo se gestionará mediante Extreme Programming (XP): Pair Programming, TDD, integración continua con GitHub Actions, Planning Game semanal y cliente in-situ.',
                      'Metodología de desarrollo seleccionada'),
    ('Seguridad',     'Todas las comunicaciones deben ir sobre HTTPS (TLS 1.2 mínimo) con certificado SSL válido (Let\'s Encrypt). El tráfico HTTP debe ser redirigido automáticamente a HTTPS.',
                      'RNF-05, RQ-CAL-OLY-05'),
    ('Rendimiento',   'La validación de acceso QR debe completarse en menos de 300 ms (P95). La respuesta general de la API debe ser < 500 ms (P95). Se requieren índices en columnas críticas de la BD.',
                      'RNF-01, RNF-02 — RQ-CAL-OLY-01, RQ-CAL-OLY-02'),
    ('Concurrencia',  'El sistema debe soportar mínimo 100 usuarios concurrentes sin sobreventa de cupos en reservas. Se implementa SELECT FOR UPDATE NOWAIT en PostgreSQL.',
                      'Shadowing horas pico — RNF-03, RQ-CAL-OLY-03'),
    ('Escalabilidad', 'La arquitectura modular (separación frontend/backend/BD) debe permitir escalar de forma independiente cada capa sin rediseñar el sistema completo.',
                      'Stakeholders — RD-03, versión 2.0 y 3.0'),
    ('Compatibilidad','La interfaz web debe ser funcional en Chrome, Firefox, Edge y Safari (últimas 2 versiones). Compatible con pantallas >= 768px (tablets de recepción).',
                      'RNF-08, RNF-09 — RQ-CAL-OLY-08, RQ-CAL-OLY-09'),
    ('Integración',   'El sistema debe permitir futuras integraciones con APIs externas (pasarela de pagos — SC-002, WhatsApp — SC-001) mediante diseño de endpoints REST versionados.',
                      'Dueño del gimnasio — Gestión de Cambios v2.0'),
    ('Legal',         'El registro de datos personales (DNI, teléfono, historial médico) debe cumplir con la Ley N.° 29733 de Protección de Datos Personales del Perú. Se requiere consentimiento explícito del socio.',
                      'Requisito legal del dominio — RD-05'),
]
for row in restricciones:
    r = rest_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

doc.add_paragraph('La siguiente tabla clasifica los elementos del proyecto según su naturaleza:')
clas_t = doc.add_table(rows=1, cols=3)
clas_t.style = 'Table Grid'
header_row(clas_t, ['Elemento', 'Clasificación', 'Justificación'])
clasificaciones = [
    ('Desarrollo iterativo en sprints de 1 semana',    'Requisito de Proyecto / Proceso', 'Define la frecuencia de hitos de validación y la metodología de trabajo. No visible para el usuario final.'),
    ('PostgreSQL con transacciones ACID',              'Requisito de Proyecto (Restricción técnica)', 'Condiciona el motor de BD y la estrategia de concurrencia. Impacta en la arquitectura de toda la capa de datos.'),
    ('VPS Linux + Nginx + PM2',                        'Requisito de Proyecto (Restricción de infraestructura)', 'Define el entorno de despliegue de producción. Limita las decisiones de hosting del equipo.'),
    ('Integración obligatoria con torniquete físico',  'Requisito de Proyecto (Restricción técnica)', 'Obliga al equipo a adaptarse al hardware preexistente del gimnasio (protocolo HTTP/serial).'),
    ('Control de versiones con Git y GitHub',          'Requisito de Proyecto / Proceso', 'Herramienta exclusiva del equipo para gestión del código fuente y CI/CD.'),
    ('Validación de código QR en menos de 300 ms',     'Requisito del Producto (RNF)', 'Comportamiento del sistema de cara al usuario (recepcionista y socio). Verificable con métricas.'),
    ('Registro de Ficha Médica bajo Ley N.° 29733',    'Requisito de Dominio / Restricción Legal', 'Normativa peruana que el producto debe cumplir para operar legalmente. Impacta en formularios y consentimientos.'),
    ('JWT con expiración de 24h / QR con 12h',         'Requisito de Producto (RNF-06)', 'Comportamiento de seguridad verificable técnicamente con herramientas de decodificación JWT.'),
]
for row in clasificaciones:
    r = clas_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

doc.add_paragraph('El siguiente diagrama muestra el stack tecnológico de producción elegido y la relación entre sus capas:')
add_diagram(doc, UML_STACK, 'Figura 2.1: Diagrama de Clases — Stack Tecnológico de Producción OLYMPUS CORE', 6.2)

add_diagram(doc, UML_ARQUITECTURA, 'Figura 2.1b: Diagrama de Componentes — Arquitectura de Producción (VPS + Nginx + PostgreSQL)', 6.2)

# 2.2 Análisis de Riesgos
doc.add_heading('2.2. Análisis de Riesgos', 2)
risk_t = doc.add_table(rows=1, cols=5)
risk_t.style = 'Table Grid'
header_row(risk_t, ['Restricción', 'Riesgo Identificado', 'Probabilidad', 'Impacto', 'Mitigación'])
riesgos = [
    ('Límite de tiempo (5 meses)',
     'No completar el alcance total de los módulos operativos del MVP1 en el tiempo establecido.',
     'Alta', 'Alto',
     'Usar Planning Game de XP para priorizar el MVP1. Ciclos semanales para asegurar valor temprano. Matriz MoSCoW congelada.'),
    ('Metodología XP (TDD y Pair Programming)',
     'Resistencia del equipo a programar en pares o curva de aprendizaje lenta al escribir pruebas antes que el código.',
     'Media', 'Medio',
     'Rotación constante de parejas para compartir conocimiento. Mentoría técnica y adopción gradual del TDD en los componentes más críticos primero.'),
    ('Base de datos PostgreSQL',
     'Inconsistencias en el esquema de BD debido a cambios rápidos y refactorización propia del ciclo ágil.',
     'Baja', 'Alto',
     'Control de versiones de la BD con node-pg-migrate integrado al pipeline de CI/CD. Cada cambio al esquema requiere un archivo de migración aprobado.'),
    ('Integración QR con hardware del gimnasio',
     'Incompatibilidad del protocolo del torniquete físico actual con el sistema web (HTTP vs. serial RS-232).',
     'Media', 'Alto',
     'Desarrollar un spike (prototipo arquitectónico) en la semana 1 para validar la comunicación con el torniquete. Definir protocolo (HTTP o adaptador serial) antes de la implementación.'),
    ('Seguridad (JWT + bcrypt)',
     'Robo de credenciales, tokens JWT comprometidos o accesos indebidos por mal manejo de secretos.',
     'Media', 'Crítico',
     'Variables de entorno gestionadas con dotenv + archivos .env.production en el servidor (nunca en el repositorio). JWT con expiración estricta. Rotación de secretos en caso de compromiso.'),
    ('Escalabilidad bajo carga',
     'Lentitud o caída del sistema ante pico de usuarios (hora pico 7AM y 6PM con 100+ socios concurrentes).',
     'Media', 'Alto',
     'Arquitectura modular con PM2 en modo cluster (usa todos los núcleos del VPS). Índices en PostgreSQL en columnas críticas. Pruebas de carga con k6 antes del go-live.'),
    ('APIs externas (SC-001, SC-002)',
     'Caída del servicio de WhatsApp Business API o pasarela de pagos durante horas de operación.',
     'Baja', 'Medio',
     'Patrón retry con exponential backoff (máximo 3 intentos). Canal de respaldo (correo electrónico). Las fallas de APIs externas no deben bloquear el flujo principal del MVP1.'),
    ('Infraestructura VPS',
     'Caída del servidor físico o agotamiento de recursos (CPU/RAM/disco) en momentos críticos.',
     'Baja', 'Alto',
     'Monitoreo con UptimeRobot (ping cada 5 min). Alertas por email al equipo. Backups diarios automáticos de la BD con pg_dump + cron. Plan de recuperación documentado.'),
]
IMP_BG = {'Alto': 'FFD7D7', 'Medio': 'FFF3CD', 'Crítico': 'F8D7DA', 'Bajo': 'D4EDDA'}
for row in riesgos:
    r = risk_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
    set_cell_bg(r.cells[3], IMP_BG.get(row[3], 'FFFFFF'))
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ESPECIFICACIÓN DE REQUISITOS DE DESARROLLO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Especificación de Requisitos de Desarrollo', 1)
doc.add_paragraph(
    'Después de la identificación y análisis de restricciones, el equipo documentó detalladamente '
    'los requisitos técnicos necesarios para construir una plataforma robusta, escalable y '
    'alineada a los objetivos del gimnasio.'
)

# 3.1 Lista de Requisitos de Desarrollo
doc.add_heading('3.1. Requisitos de Desarrollo', 2)
rds_list = [
    ('RD-DEV-01', 'Frontend',        'El frontend será desarrollado con React 18 y Vite 5. La UI será construida con CSS Modules/Vanilla CSS. La comunicación con el backend se realizará con Axios sobre HTTPS. El build de producción genera archivos estáticos servidos por Nginx.'),
    ('RD-DEV-02', 'Backend',         'El backend será implementado con Node.js 20 LTS y Express 4.x utilizando arquitectura basada en APIs REST versionadas (/api/v1/). Se gestionará con PM2 en modo cluster para aprovechar todos los núcleos del VPS.'),
    ('RD-DEV-03', 'Base de Datos',   'PostgreSQL 16 será el SGBD. Se usará el driver node-postgres (pg) con pool de conexiones (máximo 20). Las migraciones se gestionarán con node-pg-migrate. Se crearán índices en columnas de búsqueda frecuente (dni, socio_id, estado, fecha_fin).'),
    ('RD-DEV-04', 'Seguridad',       'El sistema implementará autenticación con JWT (HS256, exp: 24h para sesiones, 12h para QR) y cifrado bcrypt (factor: 12 rondas). Todas las rutas protegidas requerirán el header Authorization: Bearer <token>. Control de acceso RBAC por middleware.'),
    ('RD-DEV-05', 'Concurrencia',    'El módulo de reservas implementará transacciones PostgreSQL con SELECT FOR UPDATE NOWAIT para prevenir condiciones de carrera (Race Conditions) en la actualización del aforo. El sistema debe soportar 100 usuarios concurrentes sin sobreventa de cupos.'),
    ('RD-DEV-06', 'QR y Acceso',     'El módulo de validación de acceso generará tokens QR con JWT firmado (exp: 12h). El endpoint POST /api/acceso/validar validará la firma, membresía y penalizaciones en < 300 ms (P95) y registrará cada intento en la tabla logs_acceso.'),
    ('RD-DEV-07', 'Infraestructura', 'El servidor de producción será un VPS Linux (Ubuntu 22.04). Nginx actuará como reverse proxy con SSL/TLS (Let\'s Encrypt). El despliegue se automatizará con GitHub Actions vía SSH. No se utilizarán plataformas PaaS gratuitas con limitaciones de inactividad.'),
    ('RD-DEV-08', 'Control de versiones', 'El equipo utilizará GitHub con Git Flow: ramas main (producción), develop (integración) y feature/* (funcionalidades). Cada PR requiere revisión de al menos 1 compañero antes del merge.'),
    ('RD-DEV-09', 'Metodología',     'El desarrollo seguirá XP: Pair Programming rotativo, TDD para módulos críticos (reservas, QR, autenticación), sprints de 1 semana con Planning Game, cliente in-situ (o representante) para feedback continuo.'),
    ('RD-DEV-10', 'CI/CD',          'El pipeline de GitHub Actions ejecutará: lint (ESLint), tests unitarios (Jest + Supertest), tests de integración, build de producción y deploy automático al VPS vía SSH con PM2 reload (zero-downtime). El pipeline incluye health check post-deploy con rollback automático.'),
    ('RD-DEV-11', 'Pruebas',        'Se implementarán pruebas unitarias (Jest) para controladores y servicios, pruebas de integración (Supertest) para endpoints REST, y pruebas de carga con k6 para validar RNF-01, RNF-02 y RNF-03 antes del go-live.'),
    ('RD-DEV-12', 'Rendimiento',    'El sistema debe mantener tiempos de respuesta: POST /acceso/validar < 300 ms (P95), endpoints generales < 500 ms (P95), bajo carga de 100 VUs concurrentes. Se implementarán índices en PostgreSQL y caché a nivel de aplicación donde aplique.'),
    ('RD-DEV-13', 'Logs y Monitoreo', 'Se implementará logging estructurado con Winston (niveles INFO/WARN/ERROR) con rotación diaria de archivos. El monitoreo de disponibilidad (uptime) se realizará con UptimeRobot. Los logs de acceso QR se almacenarán en la tabla logs_acceso de PostgreSQL.'),
    ('RD-DEV-14', 'Backups',         'Se configurará un job cron diario (3:00 AM) para ejecutar pg_dump de la BD de producción. Los backups se almacenarán en una ruta segura del VPS con retención de 30 días. En caso de falla crítica, se podrá restaurar en < 4 horas.'),
]
rd_t = doc.add_table(rows=1, cols=4)
rd_t.style = 'Table Grid'
header_row(rd_t, ['ID', 'Categoría', 'Descripción', ''])
# Rebuild as 3-col
rd_t2 = doc.add_table(rows=1, cols=3)
rd_t2.style = 'Table Grid'
header_row(rd_t2, ['ID', 'Categoría', 'Descripción'])
for row in rds_list:
    r = rd_t2.add_row()
    for c, t in zip(r.cells, row): c.text = t
rd_t._tbl.getparent().remove(rd_t._tbl)
doc.add_paragraph()

# 3.2 TEC-001: Reservas con concurrencia
doc.add_heading('3.2. Requisito Técnico TEC-001', 2)
doc.add_paragraph('ID: TEC-001', style='Normal').runs[0].bold = True
two_col_table(doc, [
    ('ID',          'TEC-001'),
    ('Título',      'Implementar módulo de reservas de clase grupal con control de concurrencia'),
    ('Tipo',        'Tarea técnica (backend + concurrencia + BD)'),
    ('RF asociados','RF-09 (Reservar clase grupal), RF-14 (Penalización), RF-16 (Gestión de clases), RF-18 (Control de aforo)'),
    ('RNF asociados','RNF-01 (Latencia < 500ms), RNF-03 (100 VUs concurrentes)'),
    ('RQ-CAL asociados','RQ-CAL-OLY-02 (Latencia API), RQ-CAL-OLY-03 (Capacidad)'),
])

doc.add_paragraph(
    'Se necesita la implementación de un módulo backend que permita a los socios reservar '
    'clases grupales en tiempo real, validando automáticamente la disponibilidad de cupos, '
    'membresía activa y restricciones por penalización. El módulo debe soportar múltiples '
    'reservas simultáneas sin pérdida ni sobrescritura de información, implementando '
    'transacciones PostgreSQL con bloqueos a nivel de fila (SELECT FOR UPDATE NOWAIT) '
    'para prevenir condiciones de carrera (Race Conditions).'
)

doc.add_paragraph('Criterios de aceptación (funcionales):')
for c in [
    'Dado un usuario con membresía activa, sin penalización activa y cupos disponibles, el sistema registra la reserva correctamente (HTTP 201).',
    'Después de cada reserva exitosa, el aforo_actual de la clase se actualiza automáticamente en la BD.',
    'Si la clase no tiene cupos disponibles, el sistema devuelve HTTP 409 con mensaje CLASE_SIN_CUPOS.',
    'El sistema bloquea la reserva si el socio tiene membresía vencida (HTTP 400 MEMBRESIA_VENCIDA).',
    'El sistema bloquea la reserva si el socio tiene penalización activa (HTTP 400 PENALIZACION_ACTIVA).',
    'El sistema rechaza reservas duplicadas en el mismo horario (HTTP 409 RESERVA_DUPLICADA).',
    'Se registran todas las reservas en la BD con fecha, hora y usuario asociado.',
    'El sistema soporta reservas simultáneas sin duplicidad ni inconsistencia (verificado con prueba de carga).',
]:
    add_bullet(doc, c)

doc.add_paragraph('Requisitos no funcionales:')
for c in [
    'Rendimiento: El proceso de reserva debe completarse en < 500 ms (P95) bajo carga de 100 VUs.',
    'Concurrencia: El sistema debe soportar al menos 100 usuarios simultáneos sin sobreventa de cupos.',
    'Disponibilidad: El módulo debe mantenerse operativo el 99% del tiempo (RQ-CAL-OLY-11).',
    'Integridad: No debe existir sobrescritura ni pérdida de reservas. Transacciones ACID garantizadas.',
    'Compatibilidad: Los endpoints REST deben ser consumibles desde el cliente web y (en v3.0) desde la app móvil.',
]:
    add_bullet(doc, c)

doc.add_paragraph('Dependencias técnicas:')
for d in [
    'PostgreSQL 16 configurado con soporte transaccional y row-level locking.',
    'Módulo de autenticación JWT operativo (middleware auth).',
    'Módulo de penalizaciones operativo (RF-14 implementado).',
    'Horarios y aforos previamente registrados en la BD.',
    'Índice en clases_grupales(id) para garantizar que el SELECT FOR UPDATE use Index Scan.',
]:
    add_bullet(doc, d)

doc.add_paragraph('El siguiente diagrama ilustra el control de concurrencia en el módulo de reservas:')
add_diagram(doc, UML_CONCURRENCIA,
            'Figura 3.2: Diagrama de Secuencia — Control de Concurrencia en Reservas con SELECT FOR UPDATE (TEC-001)', 6.0)

doc.add_paragraph('Definition of Ready (DoR):')
dor1_t = doc.add_table(rows=1, cols=3)
dor1_t.style = 'Table Grid'
header_row(dor1_t, ['Criterio', 'Estado requerido', 'Responsable'])
for row in [
    ('Modelo de datos validado',   'Tablas clases_grupales y reservas implementadas con índices', 'Backend Team'),
    ('APIs REST definidas',        'Endpoints documentados (POST /api/reservas, etc.)',            'Backend Team'),
    ('Reglas de negocio definidas','Penalizaciones, aforos y cancelaciones documentadas',          'QA Team'),
    ('Diagramas UML disponibles',  'Casos de uso y secuencias completados',                        'Equipo de Análisis'),
    ('Casos de prueba definidos',  'Escenarios concurrentes y de error documentados',              'QA Team'),
    ('Entorno backend operativo',  'Servidor VPS, PostgreSQL y API base funcionando',              'DevOps / Backend'),
]:
    r = dor1_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

doc.add_paragraph('Definition of Done (DoD):')
dod1_t = doc.add_table(rows=1, cols=2)
dod1_t.style = 'Table Grid'
header_row(dod1_t, ['Criterio', 'Verificación'])
for row in [
    ('Reservas operativas',        'Registro correcto de reservas con todos los campos requeridos'),
    ('Validación de aforo',        'aforo_actual se actualiza automáticamente tras cada reserva exitosa'),
    ('Manejo concurrente',         'Prueba de carga k6: 0 reservas duplicadas con 100 VUs simultáneos'),
    ('Pruebas Jest/Supertest OK',  'Todos los casos de prueba (éxito, cupo lleno, membresía vencida) pasan'),
    ('Logs implementados',         'Tabla reservas con trazabilidad completa (socio_id, clase_id, timestamp)'),
    ('Seguridad validada',         'Endpoint protegido por middleware JWT, solo socios autenticados pueden reservar'),
    ('Documentación actualizada',  'Endpoints documentados con ejemplos de request/response'),
    ('Revisión técnica',           'Code review aprobado por al menos 1 integrante del equipo'),
]:
    r = dod1_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

# 3.3 TEC-002: Validación QR
doc.add_heading('3.3. Requisito Técnico TEC-002', 2)
two_col_table(doc, [
    ('ID',          'TEC-002'),
    ('Título',      'Implementar módulo de validación de acceso mediante código QR dinámico'),
    ('Tipo',        'Tarea técnica (backend + integración QR + hardware)'),
    ('RF asociados','RF-06 (Validar acceso QR), RF-07 (Dashboard con logs), RF-10 (Ver perfil y QR), RD-01, RD-12'),
    ('RNF asociados','RNF-02 (QR < 300ms), RNF-05 (HTTPS), RNF-06 (JWT exp.)'),
    ('RQ-CAL asociados','RQ-CAL-OLY-01 (Latencia QR), RQ-CAL-OLY-06 (Expiración token)'),
])

doc.add_paragraph(
    'Implementación de un servicio backend que reciba tokens QR generados desde la aplicación '
    'del socio (JWT firmado con exp de 12 horas), consulte el estado de la membresía en tiempo '
    'real en PostgreSQL y permita o bloquee automáticamente el acceso al gimnasio mediante '
    'integración con el sistema de torniquete físico. El módulo debe registrar cada intento '
    'de acceso en logs estructurados para fines de auditoría y trazabilidad, y responder '
    'en < 300 ms (P95) para no generar colas en recepción.'
)

doc.add_paragraph('Criterios de aceptación (funcionales):')
for c in [
    'Dado un token QR válido (firma JWT correcta, no expirado) y membresía activa sin penalización: el sistema devuelve HTTP 200 {resultado: "PERMITIDO"} y señaliza apertura del torniquete.',
    'Dado un token QR expirado o con firma inválida: el sistema devuelve HTTP 401 {motivo: "TOKEN_INVALIDO"} sin consultar la BD.',
    'Dado un QR válido pero membresía vencida: el sistema devuelve HTTP 403 {motivo: "MEMBRESIA_VENCIDA"}.',
    'Dado un QR válido pero penalización de bloqueo activa: el sistema devuelve HTTP 403 {motivo: "PENALIZACION_ACTIVA", desbloqueado_en: "<fecha>"}.',
    'El sistema registra socio_id, fecha/hora, resultado y motivo en la tabla logs_acceso para cada intento.',
    'Dado un error de conexión a BD o timeout (> 5s): el sistema reintenta hasta 2 veces antes de devolver HTTP 503 {error: "SERVICE_UNAVAILABLE"}.',
    'El sistema registra logs estructurados (INFO, WARN, ERROR) para cada intento de acceso usando Winston.',
]:
    add_bullet(doc, c)

doc.add_paragraph('Requisitos no funcionales:')
for c in [
    'Seguridad: tokens JWT firmados con HS256. Secreto almacenado en variable de entorno (nunca en código fuente). bcrypt para contraseñas (RNF-04).',
    'Rendimiento: el tiempo promedio de validación completa (JWT → BD → respuesta) debe ser < 300 ms (P95) — RQ-CAL-OLY-01.',
    'Disponibilidad: el módulo debe mantenerse operativo el 99% del tiempo en horario operativo del gimnasio — RQ-CAL-OLY-11.',
    'Trazabilidad: todos los accesos (permitidos y denegados) deben almacenarse en logs_acceso con campos: socio_id, timestamp, resultado, motivo — RD-09.',
    'Escalabilidad: el sistema debe soportar más de 100 validaciones concurrentes sin degradación — RNF-03.',
]:
    add_bullet(doc, c)

doc.add_paragraph('Dependencias técnicas:')
for d in [
    'PostgreSQL 16 operativo con índices en socios(id), membresias(socio_id, estado, fecha_fin).',
    'Integración validada con lector QR y torniquete físico del gimnasio (protocolo HTTP o adaptador serial).',
    'Variable de entorno JWT_SECRET definida en el servidor de producción (no en el repositorio).',
    'Conexión estable del servidor VPS a la red local del gimnasio (LAN o VLAN dedicada).',
    'Certificado SSL/TLS válido instalado en Nginx.',
]:
    add_bullet(doc, d)

doc.add_paragraph('El siguiente diagrama muestra los estados posibles del token QR durante la validación:')
add_diagram(doc, UML_ESTADO_QR,
            'Figura 3.3: Diagrama de Estado — Ciclo de Vida del Token QR (TEC-002, RF-06, RNF-02, RNF-06)', 5.8)

doc.add_paragraph('Definition of Ready (DoR):')
dor2_t = doc.add_table(rows=1, cols=2)
dor2_t.style = 'Table Grid'
header_row(dor2_t, ['Criterio', 'Estado requerido'])
for row in [
    ('Diagramas UML disponibles',    'Diagramas de secuencia, estado y actividades completados'),
    ('Base de datos modelada',       'Tablas socios, membresias y logs_acceso con índices validados'),
    ('Endpoints definidos',          'POST /api/acceso/validar documentado con esquema de request/response'),
    ('Validaciones definidas',       'Reglas de membresía, expiración QR y penalizaciones documentadas'),
    ('Entorno de desarrollo',        'Backend y PostgreSQL configurados en VPS o entorno local equivalente'),
    ('Casos de prueba definidos',    'Escenarios de éxito, QR expirado, membresía vencida y penalización documentados'),
    ('Seguridad documentada',        'JWT_SECRET y bcrypt rounds especificados. Variables de entorno definidas.'),
    ('Integración hardware validada','Spike de conexión con torniquete completado (protocolo confirmado)'),
]:
    r = dor2_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

doc.add_paragraph('Definition of Done (DoD):')
dod2_t = doc.add_table(rows=1, cols=2)
dod2_t.style = 'Table Grid'
header_row(dod2_t, ['Criterio', 'Verificación'])
for row in [
    ('Código implementado',       'Funcionalidad integrada y desplegada en el VPS de producción'),
    ('Validación QR operativa',   'Acceso PERMITIDO y DENEGADO funcionando con logs correctos'),
    ('Pruebas Jest/Supertest OK', 'Todos los escenarios definidos en el DoR pasan exitosamente'),
    ('Logs implementados',        'tabla logs_acceso populada correctamente en pruebas de integración'),
    ('Manejo de errores',         'Timeout y BD no disponible manejados con retry + HTTP 503'),
    ('Seguridad validada',        'JWT verificado, bcrypt funcional, rutas sin JWT bloqueadas con 401'),
    ('Documentación técnica',     'Endpoint documentado con curl/Postman. Variables de entorno documentadas.'),
    ('Revisión de código',        'Code review aprobado por al menos 1 integrante del equipo'),
]:
    r = dod2_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

# 3.4 TEC-003 (nuevo)
doc.add_heading('3.4. Requisito Técnico TEC-003', 2)
two_col_table(doc, [
    ('ID',     'TEC-003'),
    ('Título', 'Configuración del pipeline CI/CD y entorno de producción (VPS + Nginx + PM2)'),
    ('Tipo',   'Tarea técnica (DevOps + infraestructura)'),
    ('RF asociados', 'Todos los RF — es la infraestructura que los soporta'),
    ('RNF asociados','RNF-11 (Disponibilidad 99%), RNF-12 (Control de versiones), RNF-13 (Índices BD)'),
    ('RQ-CAL asociados','RQ-CAL-OLY-11 (Uptime), RQ-CAL-OLY-12 (Git), RD-DEV-07, RD-DEV-10'),
])

doc.add_paragraph(
    'Configuración completa del entorno de producción en un VPS Linux y del pipeline '
    'de integración y despliegue continuo (CI/CD) con GitHub Actions. El objetivo es '
    'garantizar que cada push a la rama main dispare automáticamente lint, tests, build '
    'y deploy con zero-downtime, incluyendo health check y rollback automático en caso '
    'de fallo.'
)

doc.add_paragraph('Criterios de aceptación:')
for c in [
    'El pipeline se ejecuta automáticamente en cada push a la rama main.',
    'Un fallo en lint, tests o integración detiene el pipeline antes del deploy.',
    'El deploy utiliza SSH + PM2 reload para garantizar zero-downtime (el sistema no deja de responder durante el despliegue).',
    'Tras el deploy, el pipeline verifica GET /api/health → HTTP 200 antes de marcar el deploy como exitoso.',
    'Si el health check falla, el pipeline ejecuta rollback automático al commit anterior.',
    'Los certificados SSL/TLS se gestionan con Let\'s Encrypt (certbot) y se renuevan automáticamente.',
    'Los backups de PostgreSQL se ejecutan diariamente a las 3:00 AM con pg_dump.',
]:
    add_bullet(doc, c)

doc.add_paragraph('El siguiente diagrama muestra el pipeline CI/CD completo:')
add_diagram(doc, UML_CICD, 'Figura 3.4: Diagrama de Actividad — Pipeline CI/CD con GitHub Actions (TEC-003)', 5.0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. VERIFICACIÓN DE REQUISITOS DE DESARROLLO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Verificación de Requisitos de Desarrollo', 1)

doc.add_heading('4.1. Estrategias Generales de Validación', 2)
estrategias = [
    ('Revisiones de Pares (Peer Reviews)',
     'Se realizarán inspecciones técnicas estáticas mediante grupos de desarrolladores para '
     'validar que los requisitos documentados (DoR y DoD) sean comprendidos unánimemente. '
     'Cada PR en GitHub requiere aprobación de al menos 1 integrante antes del merge.'),
    ('Prototipos Arquitectónicos (Spikes)',
     'Antes de programar el sistema completo, se desarrollará un prototipo rápido del escaneo '
     'del código QR y su comunicación con el torniquete físico para validar la compatibilidad '
     'del protocolo (HTTP o serial RS-232). El spike se completará en la Semana 1 del proyecto.'),
    ('Pruebas de Estrés y Rendimiento',
     'Se ejecutarán simulaciones de carga concurrente con k6 en el entorno de desarrollo para '
     'asegurar que el módulo de reservas soporte 100 VUs sin sobreventa (TEC-001) y que la '
     'validación QR responda en < 300 ms (TEC-002). Las pruebas se corren antes del go-live.'),
    ('Auditorías de Seguridad (OWASP)',
     'Se ejecutarán pruebas de vulnerabilidades basadas en las guías OWASP Top 10 para validar '
     'que los datos personales (ficha médica, DNI), contraseñas y tokens estén correctamente '
     'protegidos. Se probará específicamente SQL Injection y XSS (RQ-CAL-OLY-07).'),
    ('Verificación Automatizada (CI/CD)',
     'Se auditará el cumplimiento de los estándares de codificación insertando ESLint y Jest '
     'directamente en el pipeline de GitHub Actions. Ningún código puede llegar a producción '
     'sin pasar lint y pruebas automatizadas.'),
    ('Pruebas de Disponibilidad',
     'Se configurará UptimeRobot para hacer ping al endpoint /api/health cada 5 minutos durante '
     '30 días. La disponibilidad calculada debe ser >= 99% en horario operativo (L-S 6AM-10PM) '
     'para cumplir RQ-CAL-OLY-11.'),
]
for nombre, desc in estrategias:
    p = doc.add_paragraph()
    r = p.add_run(f'{nombre}: ')
    r.bold = True
    p.add_run(desc)

# 4.2 Inspección técnica
doc.add_heading('4.2. Inspección Técnica', 2)
doc.add_paragraph('Requisito Original (ambiguo — extraído del documento original):')
p_cita = doc.add_paragraph(
    '"El sistema backend debe registrar las reservas de las clases de spinning instantáneamente '
    'y bloquear la base de datos de forma segura para que no haya sobrecupos, soportando a '
    'todos los usuarios conectados a las 6:00 p.m."'
)
p_cita.paragraph_format.left_indent = Inches(0.5)
p_cita.runs[0].italic = True

doc.add_paragraph('Evaluación del requisito original:')
insp_t = doc.add_table(rows=1, cols=4)
insp_t.style = 'Table Grid'
header_row(insp_t, ['Atributo', 'Evaluación', 'Observaciones', 'Recomendación'])
for row in [
    ('Atomicidad',   'No atómico',          'Mezcla dos preocupaciones distintas: (a) el tiempo de respuesta del sistema ("instantáneamente") y (b) el control de concurrencia y sobrecupos en la base de datos. Deben especificarse en requisitos separados.',
                                             'Separar ambos conceptos en TEC-001 (concurrencia) y TEC-002 (latencia QR).'),
    ('Testabilidad', 'No testable',          'Palabras como "instantáneamente" y "todos los usuarios" no son cuantificables. El equipo de pruebas no puede medir "todos"; necesita un número máximo de peticiones simultáneas (ej. 100 req/seg) para configurar la prueba de estrés con k6.',
                                             'Definir métricas cuantificables: 100 VUs concurrentes, P95 < 500 ms.'),
    ('Factibilidad', 'Parcialmente factible','Node.js y PostgreSQL pueden manejar concurrencia, pero "bloquear la base de datos entera" no es factible ni eficiente. Bloquearía todos los módulos del sistema durante la reserva.',
                                             'Especificar bloqueo a nivel de fila con SELECT FOR UPDATE NOWAIT en la tabla clases_grupales.'),
    ('Ambigüedad',   'Ambiguo',              'Términos como "de forma segura" son vagos. ¿Se refiere a seguridad contra ataques (OWASP — RQ-CAL-OLY-07) o a seguridad transaccional (ACID)?',
                                             'Especificar explícitamente: transacciones ACID de PostgreSQL para integridad y JWT/bcrypt para seguridad de acceso.'),
]:
    r = insp_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

doc.add_paragraph('Requisito Técnico Reescrito y Validado (TEC-003 reescrito):')
doc.add_paragraph(
    'Al detectar las deficiencias y la falta de parámetros matemáticos para las pruebas '
    'automatizadas, el requisito original se reescribe y divide en dos requisitos técnicos '
    'medibles:'
)

two_col_table(doc, [
    ('ID reescrito',   'TEC-001 (Control de concurrencia en reservas)'),
    ('Descripción',    'El endpoint POST /api/reservas debe implementar transacciones PostgreSQL con SELECT FOR UPDATE NOWAIT para prevenir condiciones de carrera (Race Conditions) al verificar y actualizar el aforo máximo de una clase grupal.'),
    ('Criterio matemático de aceptación',
     'Una prueba de carga con k6 simulará 150 peticiones concurrentes en el mismo segundo apuntando a una clase con aforo máximo de 20 socios. Condición de éxito: exactamente 20 respuestas HTTP 201 y 130 respuestas HTTP 409, con P95 < 500 ms para todo el lote.'),
])

doc.add_paragraph('Validaciones técnicas esperadas para TEC-001:')
val_t = doc.add_table(rows=1, cols=2)
val_t.style = 'Table Grid'
header_row(val_t, ['Validación', 'Resultado esperado'])
for row in [
    ('Control concurrente',           'No existe sobreventa de cupos (exactamente 20 reservas creadas)'),
    ('Integridad de datos',           'No existen reservas duplicadas para el mismo socio+clase'),
    ('Transacciones PostgreSQL ACID', 'SELECT FOR UPDATE NOWAIT operativo — verificado con EXPLAIN ANALYZE'),
    ('Respuestas HTTP diferenciadas', 'HTTP 201 (éxito) y HTTP 409 (sin cupos) correctamente devueltos'),
    ('Rendimiento bajo carga',        'P95 <= 500 ms para el lote completo de 150 peticiones concurrentes'),
    ('Rollback en caso de error',     'Ninguna reserva queda en estado inconsistente si la transacción falla'),
]:
    r = val_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

doc.add_paragraph('Riesgos mitigados por TEC-001:')
for r in [
    'Pérdida de integridad del aforo (sobreventa de cupos de clase).',
    'Inconsistencias concurrentes entre múltiples socios que reservan simultáneamente.',
    'Corrupción de datos en la tabla clases_grupales.aforo_actual.',
    'Errores de sincronización entre usuarios simultáneos en hora pico (6:00 PM).',
]:
    add_bullet(doc, r)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Conclusiones', 1)
conclusiones = [
    'Los requisitos de desarrollo permitieron identificar restricciones técnicas, metodológicas, operativas y legales críticas para el proyecto OLYMPUS CORE, documentándolas de forma sistemática y alineándolas con los 34 RF, 13 RNF y 13 RQ-CAL del sistema.',
    'La selección de un stack tecnológico realista (Node.js 20 + PostgreSQL 16 + React 18 + Nginx en VPS Linux) garantiza que el sistema pueda operar en producción con alta disponibilidad, seguridad y rendimiento, sin depender de plataformas con limitaciones de recursos o tiempo de inactividad.',
    'La aplicación de metodologías ágiles (XP) combinadas con CI/CD automatizado (GitHub Actions) mejora la calidad, trazabilidad y velocidad del desarrollo, permitiendo entregables funcionales al final de cada sprint semanal.',
    'La verificación técnica aplicada al requisito ambiguo del documento original permitió transformar un requisito no testable en dos requisitos técnicos medibles (TEC-001 y TEC-002) con métricas matemáticas precisas verificables con herramientas de prueba de carga (k6).',
    'La implementación de control concurrente con SELECT FOR UPDATE NOWAIT en PostgreSQL, integración de código QR con JWT de expiración dinámica y seguridad basada en bcrypt + HTTPS fortalece la confiabilidad, seguridad y escalabilidad del sistema de cara a los 500 socios activos del gimnasio.',
    'El enfoque incremental basado en MVP1 facilita la construcción progresiva del sistema y reduce riesgos durante el desarrollo, sentando las bases técnicas necesarias para las versiones 2.0 (pagos online, WhatsApp) y 3.0 (app móvil, módulo nutricional) identificadas en el proceso de gestión de cambios.',
]
for c in conclusiones:
    add_bullet(doc, c)

# ── GUARDAR ─────────────────────────────────────────────────────────────────────
output = (
    r'c:/Users/ACER/OneDrive/Desktop/PROGRAMACION/UNIVERSITY PROJECTS/'
    r'SIstemaGestorGimnasio/Documentacion/Documentos Actualizados/'
    r'REQUISITOS_DESARROLLO_actualizado.docx'
)
doc.save(output)
print(f'Documento guardado en: {output}')
