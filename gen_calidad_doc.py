"""
gen_calidad_doc.py
Genera el documento actualizado: Especificacion de Requisitos de Calidad
OLYMPUS CORE — con matrices de trazabilidad y diagramas UML embebidos.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests, zlib, io

# ══════════════════════════════════════════════════════════════════════════════
# PLANTUMB — Encoding y fetch de imagen PNG desde la API pública
# ══════════════════════════════════════════════════════════════════════════════
def _encode6bit(b):
    if b < 10: return chr(48 + b)
    b -= 10
    if b < 26: return chr(65 + b)
    b -= 26
    if b < 26: return chr(97 + b)
    b -= 26
    return '-' if b == 0 else '_'

def _append3bytes(b1, b2, b3):
    return (
        _encode6bit(b1 >> 2) +
        _encode6bit(((b1 & 0x3) << 4) | (b2 >> 4)) +
        _encode6bit(((b2 & 0xF) << 2) | (b3 >> 6)) +
        _encode6bit(b3 & 0x3F)
    )

def plantuml_encode(text):
    data = zlib.compress(text.encode('utf-8'))[2:-4]
    res = ""
    for i in range(0, len(data), 3):
        if i + 2 < len(data): res += _append3bytes(data[i], data[i+1], data[i+2])
        elif i + 1 < len(data): res += _append3bytes(data[i], data[i+1], 0)
        else: res += _append3bytes(data[i], 0, 0)
    return res

def fetch_diagram(uml_code, width_inches=6.2):
    """Devuelve (BytesIO_png, width_inches) o None si falla la red."""
    try:
        enc = plantuml_encode(uml_code)
        url = f"http://www.plantuml.com/plantuml/png/{enc}"
        r = requests.get(url, timeout=20)
        if r.status_code == 200 and r.content[:4] == b'\x89PNG':
            return io.BytesIO(r.content), width_inches
    except Exception:
        pass
    return None

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, cols, bg='1E3A5F', fg='FFFFFF'):
    for cell, txt in zip(table.rows[0].cells, cols):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        set_cell_bg(cell, bg)

def add_diagram(doc, uml_code, caption, width=5.8):
    result = fetch_diagram(uml_code, width)
    if result:
        img_stream, w = result
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(w))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        # Fallback: bloque de código
        doc.add_paragraph(f'[Diagrama: {caption}]').italic = True
        pre = doc.add_paragraph(uml_code)
        pre.style.font.size = Pt(8)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# DATOS DE LOS 13 REQUISITOS DE CALIDAD
# ══════════════════════════════════════════════════════════════════════════════
RQ_DATA = [
    {
        'id': 'RQ-CAL-OLY-01',
        'nombre': 'Tiempo de Validacion de Acceso por QR',
        'stakeholder': 'Socio / Recepcionista',
        'necesidad': 'Validar el ingreso fisico al gimnasio rapidamente, sin generar colas en recepcion.',
        'descripcion': 'Procesar la lectura del QR y emitir resultado PERMITIDO/DENEGADO en <= 300 ms.',
        'prioridad': 'Alta',
        'criterio': 'El tiempo de respuesta del endpoint POST /acceso/validar debe ser <= 300 ms en el percentil 95, medido con 50 usuarios concurrentes.',
        'iso_caracteristica': 'Eficiencia de Desempeno',
        'iso_subcaracteristica': 'Comportamiento Temporal',
        'metrica_nombre': 'TiempoLatenciaQR',
        'metrica_unidad': 'Milisegundos (ms)',
        'metrica_objetivo': '<= 300 ms',
        'verificacion_metodo': 'Pruebas de latencia extremo a extremo con k6 o Apache JMeter',
        'verificacion_evidencia': 'Logs del endpoint con marcas de tiempo en milisegundos',
        'tecnica_origen': 'Observacion (flujo as-is en recepcion)',
        'rfs_asociados': 'RF-06, RF-07',
        'rnf_asociado': 'RNF-02',
        'diagrama_tipo': 'clase',
    },
    {
        'id': 'RQ-CAL-OLY-02',
        'nombre': 'Tiempo de Respuesta General de la API REST',
        'stakeholder': 'Administrador / Recepcionista',
        'necesidad': 'Operar el sistema de forma fluida durante la jornada laboral sin esperas perceptibles.',
        'descripcion': 'Todos los endpoints de la API deben responder en <= 500 ms bajo carga normal de hasta 50 usuarios concurrentes.',
        'prioridad': 'Alta',
        'criterio': 'Percentil 95 de tiempos de respuesta < 500 ms con 50 usuarios concurrentes. Medido con herramienta de carga.',
        'iso_caracteristica': 'Eficiencia de Desempeno',
        'iso_subcaracteristica': 'Utilizacion de Recursos',
        'metrica_nombre': 'TiempoRespuestaAPI',
        'metrica_unidad': 'Milisegundos (ms)',
        'metrica_objetivo': '<= 500 ms (P95)',
        'verificacion_metodo': 'Prueba de carga con k6 — 50 VUs durante 5 minutos',
        'verificacion_evidencia': 'Reporte de k6 con percentiles de latencia',
        'tecnica_origen': 'Benchmarking (analisis de sistemas similares)',
        'rfs_asociados': 'RF-01, RF-02, RF-03, RF-04, RF-05, RF-15',
        'rnf_asociado': 'RNF-01',
        'diagrama_tipo': 'secuencia',
    },
    {
        'id': 'RQ-CAL-OLY-03',
        'nombre': 'Capacidad de Usuarios Concurrentes',
        'stakeholder': 'Administrador del Gimnasio',
        'necesidad': 'El sistema debe sostenerse durante las horas pico sin caidas ni degradacion del servicio.',
        'descripcion': 'Soportar al menos 100 usuarios concurrentes con tasa de error < 1% y tiempo medio de respuesta < 800 ms.',
        'prioridad': 'Alta',
        'criterio': 'Prueba de estres con 100 VUs simultaneos: tasa de error < 1%, tiempo medio < 800 ms.',
        'iso_caracteristica': 'Eficiencia de Desempeno',
        'iso_subcaracteristica': 'Capacidad',
        'metrica_nombre': 'MaxUsuariosConcurrentes',
        'metrica_unidad': 'Numero de usuarios virtuales (VU)',
        'metrica_objetivo': '>= 100 VUs sin degradacion',
        'verificacion_metodo': 'Prueba de carga escalonada (ramp-up 0 → 100 VUs en 60s)',
        'verificacion_evidencia': 'Informe de k6 con throughput, errores y latencias',
        'tecnica_origen': 'Observacion (horas pico 7AM-9AM y 6PM-8PM)',
        'rfs_asociados': 'RF-07, RF-09, RF-15',
        'rnf_asociado': 'RNF-03',
        'diagrama_tipo': 'estado',
    },
    {
        'id': 'RQ-CAL-OLY-04',
        'nombre': 'Almacenamiento Seguro de Contrasenas',
        'stakeholder': 'Todos los usuarios del sistema',
        'necesidad': 'Proteger las credenciales de los usuarios ante una eventual brecha de seguridad en la base de datos.',
        'descripcion': 'Todas las contrasenas deben almacenarse en formato bcrypt con factor de costo >= 10. Nunca en texto plano.',
        'prioridad': 'Alta',
        'criterio': 'Inspeccion de la columna password_hash en la BD: todos los registros deben iniciar con $2b$10$ (formato bcrypt).',
        'iso_caracteristica': 'Seguridad',
        'iso_subcaracteristica': 'Confidencialidad',
        'metrica_nombre': 'FormatoHash',
        'metrica_unidad': 'Formato de cadena (patron regex)',
        'metrica_objetivo': '100% de registros con patron ^\\$2b\\$1[0-9]\\$',
        'verificacion_metodo': 'Inspeccion de BD + prueba de autenticacion con credencial conocida',
        'verificacion_evidencia': 'Query SELECT password_hash FROM usuarios — verificacion de patron',
        'tecnica_origen': 'Brainstorming del equipo de desarrollo',
        'rfs_asociados': 'RF-01, RF-10',
        'rnf_asociado': 'RNF-04',
        'diagrama_tipo': 'clase',
    },
    {
        'id': 'RQ-CAL-OLY-05',
        'nombre': 'Comunicacion HTTPS Obligatoria',
        'stakeholder': 'Todos los usuarios del sistema',
        'necesidad': 'Garantizar que los datos sensibles (credenciales, datos de socios) no puedan ser interceptados.',
        'descripcion': 'Todo el trafico entre el cliente y el servidor debe realizarse por HTTPS con TLS valido en produccion.',
        'prioridad': 'Alta',
        'criterio': 'Las URLs de produccion deben iniciar con https://. Las peticiones HTTP deben ser rechazadas o redirigidas.',
        'iso_caracteristica': 'Seguridad',
        'iso_subcaracteristica': 'Integridad',
        'metrica_nombre': 'PropocionTraficoCifrado',
        'metrica_unidad': 'Porcentaje (%)',
        'metrica_objetivo': '100% del trafico via HTTPS',
        'verificacion_metodo': 'Inspeccion de configuracion del servidor + prueba con navegador (candado)',
        'verificacion_evidencia': 'Captura de headers HTTP mostrando Strict-Transport-Security',
        'tecnica_origen': 'Brainstorming + Benchmarking (estandar del sector)',
        'rfs_asociados': 'RF-01, RF-06, RF-07',
        'rnf_asociado': 'RNF-05',
        'diagrama_tipo': 'actividad',
    },
    {
        'id': 'RQ-CAL-OLY-06',
        'nombre': 'Expiracion de Tokens JWT y QR',
        'stakeholder': 'Sistema / Equipo de Seguridad',
        'necesidad': 'Limitar la ventana de uso de tokens robados o compartidos, reduciendo el riesgo de acceso no autorizado.',
        'descripcion': 'Los tokens JWT de sesion deben expirar en 24 horas. Los tokens QR deben expirar en 12 horas (configurable).',
        'prioridad': 'Alta',
        'criterio': 'Decodificar cualquier JWT emitido: el campo exp debe corresponder a <= 24h desde la emision. Para QR: <= 12h.',
        'iso_caracteristica': 'Seguridad',
        'iso_subcaracteristica': 'Autenticidad',
        'metrica_nombre': 'VidaMaximaToken',
        'metrica_unidad': 'Horas (h)',
        'metrica_objetivo': 'JWT: <= 24h / QR: <= 12h',
        'verificacion_metodo': 'Prueba de acceso con token expirado — debe rechazarse con HTTP 401',
        'verificacion_evidencia': 'Decodificacion del JWT mostrando campo exp correcto',
        'tecnica_origen': 'Benchmarking (analisis de GymMaster y Mindbody)',
        'rfs_asociados': 'RF-01, RF-06, RF-07, RF-10',
        'rnf_asociado': 'RNF-06',
        'diagrama_tipo': 'secuencia',
    },
    {
        'id': 'RQ-CAL-OLY-07',
        'nombre': 'Prevencion de Inyeccion SQL y XSS',
        'stakeholder': 'Sistema / Equipo de Seguridad',
        'necesidad': 'Proteger la integridad de la base de datos y la experiencia del usuario ante ataques maliciosos.',
        'descripcion': 'Todos los endpoints deben sanitizar las entradas, rechazando payloads de inyeccion SQL y XSS sin generar errores internos.',
        'prioridad': 'Alta',
        'criterio': 'Intentar enviar payloads: OR 1=1 en campos de busqueda y <script>alert(1)</script> en campos de texto. El sistema no debe ejecutarlos.',
        'iso_caracteristica': 'Seguridad',
        'iso_subcaracteristica': 'Resistencia a Ataques',
        'metrica_nombre': 'TasaRechazoPayloadsInvalidos',
        'metrica_unidad': 'Porcentaje (%)',
        'metrica_objetivo': '100% de payloads maliciosos rechazados o ignorados',
        'verificacion_metodo': 'Pruebas de penetracion basicas (OWASP Top 10)',
        'verificacion_evidencia': 'Reporte de pruebas con lista de payloads probados y resultados',
        'tecnica_origen': 'Brainstorming del equipo + estandar OWASP',
        'rfs_asociados': 'RF-02, RF-05, RF-08, RF-12',
        'rnf_asociado': 'RNF-07',
        'diagrama_tipo': 'clase',
    },
    {
        'id': 'RQ-CAL-OLY-08',
        'nombre': 'Compatibilidad de Navegadores Web',
        'stakeholder': 'Todos los usuarios del sistema',
        'necesidad': 'Garantizar que el sistema funcione correctamente desde cualquier computadora del gimnasio sin instalar software adicional.',
        'descripcion': 'La interfaz debe ser funcional en Chrome, Firefox, Edge y Safari (ultimas 2 versiones de cada uno).',
        'prioridad': 'Media',
        'criterio': 'Ejecutar el flujo completo (login, registrar socio, validar QR) en cada navegador. Sin errores criticos de consola.',
        'iso_caracteristica': 'Compatibilidad',
        'iso_subcaracteristica': 'Coexistencia',
        'metrica_nombre': 'NavegadoresCompatibles',
        'metrica_unidad': 'Numero de navegadores soportados',
        'metrica_objetivo': '>= 4 navegadores (Chrome, Firefox, Edge, Safari)',
        'verificacion_metodo': 'Prueba manual del flujo completo en cada navegador listado',
        'verificacion_evidencia': 'Checklist de pruebas por navegador con capturas de pantalla',
        'tecnica_origen': 'Entrevista con el cliente (requisito explicito)',
        'rfs_asociados': 'RF-01, RF-04, RF-07, RF-15, RF-16',
        'rnf_asociado': 'RNF-08',
        'diagrama_tipo': 'componente',
    },
    {
        'id': 'RQ-CAL-OLY-09',
        'nombre': 'Diseno Responsive (Adaptabilidad de Pantalla)',
        'stakeholder': 'Recepcionista / Socio',
        'necesidad': 'Poder operar el sistema desde tablets de recepcion y laptops sin problemas de visualizacion.',
        'descripcion': 'La interfaz debe ser completamente operable en pantallas con ancho >= 768 px sin scroll horizontal.',
        'prioridad': 'Media',
        'criterio': 'Redimensionar el navegador a 768 px de ancho: todos los elementos deben ser visibles y operables sin scroll horizontal.',
        'iso_caracteristica': 'Usabilidad',
        'iso_subcaracteristica': 'Accesibilidad',
        'metrica_nombre': 'AnchoMinimoSoportado',
        'metrica_unidad': 'Pixeles (px)',
        'metrica_objetivo': '>= 768 px sin degradacion de UI',
        'verificacion_metodo': 'Prueba manual con DevTools (modo responsive) a 768px y 1024px',
        'verificacion_evidencia': 'Capturas de pantalla en cada breakpoint probado',
        'tecnica_origen': 'Benchmarking (estandar de aplicaciones web modernas)',
        'rfs_asociados': 'RF-15, RF-16, RF-07',
        'rnf_asociado': 'RNF-09',
        'diagrama_tipo': 'componente',
    },
    {
        'id': 'RQ-CAL-OLY-10',
        'nombre': 'Eficiencia de Flujos Criticos (Max 3 Interacciones)',
        'stakeholder': 'Recepcionista',
        'necesidad': 'Atender a los socios en recepcion rapidamente, especialmente en horas pico, sin que el sistema sea un cuello de botella.',
        'descripcion': 'Las acciones criticas de recepcion (validar QR, registrar pago) deben completarse con un maximo de 3 interacciones desde la pantalla principal.',
        'prioridad': 'Media',
        'criterio': 'Contar los pasos (clic o inputs) para validar un QR y para registrar un pago. No debe superar 3 pasos en ninguno de los dos flujos.',
        'iso_caracteristica': 'Usabilidad',
        'iso_subcaracteristica': 'Eficiencia de Uso',
        'metrica_nombre': 'NumeroInteraccionesCriticas',
        'metrica_unidad': 'Numero de clics / pasos',
        'metrica_objetivo': '<= 3 interacciones para flujos criticos',
        'verificacion_metodo': 'Prueba de usabilidad con usuario recepcionista (cronometro + conteo de clics)',
        'verificacion_evidencia': 'Registro de sesion de prueba de usabilidad (video/pantalla)',
        'tecnica_origen': 'Observacion directa del proceso as-is en recepcion',
        'rfs_asociados': 'RF-05, RF-07',
        'rnf_asociado': 'RNF-10',
        'diagrama_tipo': 'actividad',
    },
    {
        'id': 'RQ-CAL-OLY-11',
        'nombre': 'Disponibilidad del Servicio en Horario Operativo',
        'stakeholder': 'Administrador del Gimnasio',
        'necesidad': 'Garantizar que el sistema este disponible en todo el horario de atencion al publico del gimnasio.',
        'descripcion': 'El sistema debe mantener disponibilidad >= 99% de lunes a sabado entre 6:00 AM y 10:00 PM.',
        'prioridad': 'Alta',
        'criterio': 'Monitoreo de uptime durante 30 dias. La disponibilidad calculada en el horario operativo debe ser >= 99%.',
        'iso_caracteristica': 'Fiabilidad',
        'iso_subcaracteristica': 'Disponibilidad',
        'metrica_nombre': 'UptimeHorarioOperativo',
        'metrica_unidad': 'Porcentaje (%)',
        'metrica_objetivo': '>= 99% en horario operativo',
        'verificacion_metodo': 'Monitoreo con UptimeRobot (ping cada 5 min) durante 30 dias',
        'verificacion_evidencia': 'Reporte de UptimeRobot con historial de disponibilidad',
        'tecnica_origen': 'Entrevista con el cliente (necesidad critica del negocio)',
        'rfs_asociados': 'RF-01, RF-07, RF-09, RF-15',
        'rnf_asociado': 'RNF-11',
        'diagrama_tipo': 'estado',
    },
    {
        'id': 'RQ-CAL-OLY-12',
        'nombre': 'Control de Versiones y Trazabilidad del Codigo',
        'stakeholder': 'Equipo de Desarrollo',
        'necesidad': 'Mantener un historial de cambios que permita revertir errores y auditar el desarrollo del sistema.',
        'descripcion': 'El codigo fuente debe estar versionado en Git con ramas por modulo, commits descriptivos y etiquetas de version para cada entrega.',
        'prioridad': 'Baja',
        'criterio': 'Verificar en el repositorio GitHub: existencia de ramas feature/, historial de commits descriptivos y etiquetas v1.0.',
        'iso_caracteristica': 'Mantenibilidad',
        'iso_subcaracteristica': 'Capacidad de Modificacion',
        'metrica_nombre': 'PorcentajeCommitsDescriptivos',
        'metrica_unidad': 'Porcentaje (%)',
        'metrica_objetivo': '>= 80% de commits con mensaje descriptivo (> 10 caracteres)',
        'verificacion_metodo': 'Revision del historial de commits en GitHub + verificacion de etiquetas',
        'verificacion_evidencia': 'Captura del repositorio con ramas y etiquetas visibles',
        'tecnica_origen': 'Brainstorming del equipo (buenas practicas de ingenieria)',
        'rfs_asociados': 'Todos los RF',
        'rnf_asociado': 'RNF-12',
        'diagrama_tipo': 'componente',
    },
    {
        'id': 'RQ-CAL-OLY-13',
        'nombre': 'Optimizacion de Consultas con Indices en BD',
        'stakeholder': 'Sistema / Equipo de Desarrollo',
        'necesidad': 'Garantizar que las busquedas mas frecuentes (socios por DNI, membresias activas) sean rapidas incluso con muchos datos.',
        'descripcion': 'La base de datos debe tener indices en las columnas: dni, socio_id, estado, fecha_fin (membresias) y fecha_hora (clases_grupales).',
        'prioridad': 'Media',
        'criterio': 'Ejecutar EXPLAIN ANALYZE en las consultas criticas. El plan de ejecucion debe mostrar "Index Scan", no "Seq Scan".',
        'iso_caracteristica': 'Eficiencia de Desempeno',
        'iso_subcaracteristica': 'Comportamiento Temporal',
        'metrica_nombre': 'TipoEscaneoConsultasCriticas',
        'metrica_unidad': 'Tipo de escaneo (Index Scan / Seq Scan)',
        'metrica_objetivo': '100% de consultas criticas usan Index Scan',
        'verificacion_metodo': 'Ejecutar EXPLAIN ANALYZE en las 5 consultas mas frecuentes del sistema',
        'verificacion_evidencia': 'Salida de EXPLAIN ANALYZE mostrando planes de ejecucion',
        'tecnica_origen': 'Benchmarking + Observacion (cuellos de botella anticipados)',
        'rfs_asociados': 'RF-02, RF-04, RF-07, RF-08, RF-12',
        'rnf_asociado': 'RNF-13',
        'diagrama_tipo': 'clase',
    },
]

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMAS UML POR TIPO
# ══════════════════════════════════════════════════════════════════════════════

# Diagrama tipo CLASE (para RQ individuales)
def uml_clase_rq(rq):
    return f"""@startuml
hide circle
skinparam classAttributeIconSize 0
skinparam class {{
  BackgroundColor #F8F9FA
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
  FontName Calibri
}}

class "{rq['stakeholder']}" as Actor {{
  necesidad = "{rq['necesidad'][:60]}..."
}}

class "{rq['id']} : RequisitoCalidad" as RQ {{
  ID = "{rq['id']}"
  descripcion = "{rq['descripcion'][:65]}..."
  prioridad = "{rq['prioridad']}"
  criterioAceptacion = "{rq['criterio'][:65]}..."
}}

class "{rq['iso_caracteristica']} : CaracteristicaISO25010" as ISO {{
  nombre = "{rq['iso_caracteristica']}"
  subcaracteristica = "{rq['iso_subcaracteristica']}"
}}

class "{rq['metrica_nombre']} : Metrica" as MET {{
  nombre = "{rq['metrica_nombre']}"
  unidad = "{rq['metrica_unidad']}"
  valorObjetivo = "{rq['metrica_objetivo']}"
}}

class "PruebaVerificacion" as VER {{
  metodo = "{rq['verificacion_metodo'][:55]}..."
  evidencia = "{rq['verificacion_evidencia'][:55]}..."
}}

Actor "1" --> "1" RQ : origina
RQ "1" --> "1" ISO : se clasifica en
RQ "1" --> "1..*" MET : se mide con
RQ "1" --> "1" VER : se valida mediante

note bottom of RQ
  RF Asociados: {rq['rfs_asociados']}
  RNF Asociado: {rq['rnf_asociado']}
  Tecnica de origen: {rq['tecnica_origen']}
end note
@enduml"""

# Diagrama tipo SECUENCIA (validacion QR)
UML_SECUENCIA_QR = """@startuml
skinparam sequenceArrowThickness 2
skinparam roundcorner 10
skinparam maxmessagesize 100
skinparam sequenceParticipant underline

title Flujo de Validacion QR — RQ-CAL-OLY-01 (RF-07)\\nObjetivo: respuesta <= 300 ms

actor "Recepcionista" as REC
participant "Frontend\\nReact" as UI
participant "API Backend\\nNode/Express" as API
database "PostgreSQL\\nNeon" as DB

REC -> UI: Escanea / ingresa QR del socio
note right of REC: [RQ-CAL-OLY-01]\\nInicio medicion de tiempo

UI -> API: POST /acceso/validar\\n{token: "<jwt_qr>"}
activate API #DDAAFF

API -> API: 1) Verificar firma y expiracion JWT
note right of API: Si expirado → 401 QR_EXPIRADO

API -> DB: SELECT s.*, m.fecha_fin, m.estado\\nFROM socios s\\nJOIN membresias m ON s.id = m.socio_id\\nWHERE s.id = :socioId [INDEX SCAN]
DB --> API: Datos del socio y membresia

API -> DB: SELECT bloqueado_hasta\\nFROM penalizaciones_bloqueo\\nWHERE socio_id = :id\\nAND bloqueado_hasta > NOW()
DB --> API: Sin bloqueo activo

alt Acceso PERMITIDO (todas las validaciones OK)
  API -> DB: INSERT INTO logs_acceso\\n(resultado='PERMITIDO', ...)
  API --> UI: 200 OK {resultado: "PERMITIDO"}\\n[<= 300 ms de RQ-CAL-OLY-01]
  UI --> REC: Pantalla VERDE — Acceso Concedido
else Acceso DENEGADO
  API -> DB: INSERT INTO logs_acceso\\n(resultado='DENEGADO', motivo)
  API --> UI: 403 {resultado: "DENEGADO", motivo}
  UI --> REC: Pantalla ROJA — Motivo del Rechazo
end

deactivate API
note right of REC: [RQ-CAL-OLY-01] Fin medicion\\nObjetivo: todo el flujo < 300 ms
@enduml"""

# Diagrama tipo SECUENCIA (expiración de token JWT)
UML_SECUENCIA_JWT = """@startuml
skinparam sequenceArrowThickness 2
skinparam roundcorner 10

title Ciclo de Vida de Token JWT — RQ-CAL-OLY-06\\n(RF-01, RF-10)

actor "Usuario" as USR
participant "Frontend" as UI
participant "API Backend" as API

USR -> UI: Login con email y contrasena
UI -> API: POST /auth/login {email, password}
API -> API: Verificar bcrypt [RQ-CAL-OLY-04]
API -> UI: 200 OK {token: JWT, exp: +24h}
note right: [RQ-CAL-OLY-06]\\nJWT valido por 24 horas

loop Peticiones normales (< 24h)
  UI -> API: GET /ruta-protegida\\nAuthorization: Bearer <token>
  API -> API: Verificar firma JWT
  API --> UI: 200 OK {datos}
end

UI -> API: GET /ruta-protegida\\nAuthorization: Bearer <token_expirado>
API -> API: Verificar firma JWT → EXPIRADO
API --> UI: 401 TOKEN_EXPIRED
UI -> USR: Redirigir al login
note right: Token expirado correcto.\\n[RQ-CAL-OLY-06] cumplido.
@enduml"""

# Diagrama tipo ESTADO (disponibilidad del sistema)
UML_ESTADO_DISPONIBILIDAD = """@startuml
skinparam state {{
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}}
skinparam roundcorner 15

title Estados de Disponibilidad del Sistema — RQ-CAL-OLY-11\\n(RF-07, RF-15 | RNF-11)

[*] --> Operativo : Deploy exitoso

state Operativo {{
  [*] --> NormalLoad
  NormalLoad : Respuesta < 500ms
  NormalLoad : Usuarios: 1-50
  NormalLoad --> PeakLoad : Hora pico\\n(7AM o 6PM)
  PeakLoad : Respuesta < 800ms
  PeakLoad : Usuarios: 51-100
  PeakLoad --> NormalLoad : Fuera de hora pico
}}

Operativo --> Degradado : Latencia > 800ms\\no error rate > 1%%
Degradado : [RQ-CAL-OLY-02] y [RQ-CAL-OLY-03] en riesgo
Degradado : Alerta al equipo de devops
Degradado --> Operativo : Resolucion del problema

Degradado --> Inactivo : Falla critica\\no caida del servidor
Inactivo : [RQ-CAL-OLY-11] VIOLADO
Inactivo : Afecta RF-07, RF-09, RF-15

Inactivo --> Operativo : Reinicio y\\nrecuperacion

note bottom of Inactivo
  [RQ-CAL-OLY-11] exige que el tiempo
  en este estado sea < 1%% del horario
  operativo (L-S, 6AM-10PM)
end note
@enduml"""

# Diagrama tipo ESTADO (ciclo vida membresía)
UML_ESTADO_MEMBRESIA = """@startuml
skinparam state {{
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}}
skinparam roundcorner 15

title Ciclo de Vida de Membresia — RF-04, RF-05, RF-14\\n(Vinculado a RQ-CAL-OLY-01 y RQ-CAL-OLY-03)

[*] --> Activa : Pago registrado (RF-05)

Activa : fecha_fin >= HOY
Activa : Acceso QR permitido
Activa : Reservas habilitadas

Activa --> Vencida : fecha_fin < HOY
Vencida : QR → DENEGADO (MEMBRESIA_VENCIDA)
Vencida : Sin reservas permitidas

Activa --> Bloqueada : 3 strikes en 30 dias (RF-14)
Bloqueada : QR → DENEGADO (PENALIZACION_ACTIVA)
Bloqueada : Sin reservas permitidas

Bloqueada --> Activa : Pasan 7 dias\\nO admin justifica strike (RF-17)
Vencida --> Activa : Socio renueva membresia (RF-05)

Vencida --> [*] : Socio se da de baja
@enduml"""

# Diagrama tipo ACTIVIDAD (flujo HTTPS)
UML_ACTIVIDAD_HTTPS = """@startuml
skinparam activity {{
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}}
skinparam roundcorner 15

title Proceso de Comunicacion Segura HTTPS — RQ-CAL-OLY-05\\n(RF-01, RF-06, RF-07 | RNF-05)

start

:Usuario abre la aplicacion en el navegador;

if (La URL es http:// ?) then (si)
  :El servidor redirige a https://;
  note right: Redireccion 301\\nStrict-Transport-Security
else (ya es https://)
endif

:Negociacion TLS (TLS 1.2 o 1.3);
:Validacion del certificado SSL/TLS;

if (Certificado valido y no expirado?) then (no)
  :Navegador muestra error de seguridad;
  stop
else (si)
endif

:Conexion cifrada establecida;
:Usuario envia credenciales\\n(POST /auth/login) por canal cifrado;
note right: [RQ-CAL-OLY-04]\\nbcrypt en el backend

:Backend procesa y devuelve JWT;
:Todo el trafico posterior es cifrado;

note right
  [RQ-CAL-OLY-05]: 100%% del
  trafico va por HTTPS.
  Certificado gestionado por
  Vercel y Koyeb (Let's Encrypt).
end note

stop
@enduml"""

# Diagrama tipo ACTIVIDAD (flujo de reserva)
UML_ACTIVIDAD_RESERVA = """@startuml
skinparam activity {{
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}}
skinparam roundcorner 15

title Flujo de Reserva de Clase — RF-09\\n(Relacionado con RQ-CAL-OLY-01, RQ-CAL-OLY-03)

start

:Socio selecciona clase en el sistema;
:POST /reservas {socio_id, clase_id};

fork
  :Verificar membresia activa (RF-04);
fork again
  :Verificar aforo disponible (RF-08);
fork again
  :Verificar sin penalizacion activa (RF-14);
fork again
  :Verificar sin reserva duplicada;
end fork

if (Todas las validaciones OK?) then (si)
  :Crear reserva con estado 'confirmada';
  :Decrementar aforo en 1;
  :Respuesta 201 CREATED;
  note right: [RQ-CAL-OLY-02]\\nTodo en < 500ms
else (no)
  :Identificar motivo de rechazo;
  if (Sin cupos?) then (si)
    :409 CLASE_SIN_CUPOS;
  elseif (Membresia vencida?) then (si)
    :400 MEMBRESIA_VENCIDA;
  elseif (Penalizacion?) then (si)
    :400 PENALIZACION_ACTIVA;
  else (duplicada)
    :409 RESERVA_DUPLICADA;
  endif
endif

stop
@enduml"""

# Diagrama tipo DESPLIEGUE (Compatibilidad)
UML_COMPATIBILIDAD = """@startuml
skinparam node {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}
skinparam componentStyle uml2

title Compatibilidad de Navegadores — RQ-CAL-OLY-08

node "Vercel CDN" as CDN {
  [Aplicación React SPA] as APP
}

node "Dispositivos Cliente" {
  node "Google Chrome" as C1
  node "Mozilla Firefox" as C2
  node "Safari (iOS/macOS)" as C3
  node "Microsoft Edge" as C4
}

C1 -down-> APP : Renderizado Web
C2 -down-> APP : Renderizado Web
C3 -up-> APP : Renderizado Web
C4 -up-> APP : Renderizado Web

note bottom of CDN
  RQ-CAL-OLY-08: El código React/Vite
  garantiza compatibilidad cruzada sin
  depender de motores específicos.
end note
@enduml"""

# Diagrama tipo ESTADO (Responsive)
UML_RESPONSIVE = """@startuml
skinparam state {
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}

title Estados de Interfaz Responsive — RQ-CAL-OLY-09

[*] --> PantallaChica : Ancho < 768px
[*] --> PantallaTablet : 768px <= Ancho < 1024px
[*] --> PantallaDesktop : Ancho >= 1024px

state PantallaChica {
  PantallaChica : Menú hamburguesa colapsado
  PantallaChica : Layout de 1 columna
  note right: Estado no ideal\npero funcional
}

state PantallaTablet {
  PantallaTablet : Menú lateral reducido
  PantallaTablet : Layout de 2 columnas
  note right: RQ-CAL-OLY-09\n(Mínimo aceptable sin scroll)
}

state PantallaDesktop {
  PantallaDesktop : Menú lateral expandido
  PantallaDesktop : Layout completo
}

PantallaChica --> PantallaTablet : Resize (+)
PantallaTablet --> PantallaDesktop : Resize (+)
PantallaDesktop --> PantallaTablet : Resize (-)
PantallaTablet --> PantallaChica : Resize (-)
@enduml"""

# Diagrama tipo COMPONENTE (arquitectura de calidad)
UML_COMPONENTE = """@startuml
skinparam componentStyle uml2
skinparam component {{
  BackgroundColor #E8F4F8
  BorderColor #1E3A5F
}}

title Distribucion de Responsabilidades de Calidad — Vista de Componentes\\n(RQ-CAL-OLY-05, RQ-CAL-OLY-08, RQ-CAL-OLY-09, RQ-CAL-OLY-12)

node "Dispositivo Cliente" {{
  [Navegador Web] as NAV
  note right of NAV
    [RQ-CAL-OLY-08] Chrome/Firefox/Edge/Safari
    [RQ-CAL-OLY-09] Responsive >= 768px
    [RQ-CAL-OLY-10] Max 3 clics flujos criticos
  end note
}}

node "Vercel (CDN + Hosting)" {{
  component "Frontend SPA\\nReact + Vite" as FE {{
    [Auth Guard (RBAC)]
    [React Router DOM]
    [Componentes UI]
  }}
  note right of FE
    [RQ-CAL-OLY-05] HTTPS automatico
    Certificado Let's Encrypt
  end note
}}

node "Koyeb (PaaS — Backend)" {{
  component "API REST\\nNode.js + Express" as BE {{
    [Middleware JWT — RQ-CAL-OLY-06]
    [Middleware RBAC — RF-11]
    [Sanitizacion inputs — RQ-CAL-OLY-07]
    [Controladores]
  }}
  note right of BE
    [RQ-CAL-OLY-01] Latencia QR
    [RQ-CAL-OLY-02] Latencia API
    [RQ-CAL-OLY-03] 100 VUs
  end note
}}

node "Neon (DBaaS — PostgreSQL)" {{
  database "Base de Datos" as DB
  note right of DB
    [RQ-CAL-OLY-04] bcrypt
    [RQ-CAL-OLY-13] Indices
  end note
}}

node "GitHub" {{
  [Repositorio Git] as GIT
  note right of GIT
    [RQ-CAL-OLY-12] Control\\nde versiones
  end note
}}

NAV <--> FE : HTTPS [RQ-CAL-OLY-05]
FE <--> BE : HTTPS REST API
BE <--> DB : TCP/SSL (Pool conexiones)
BE --> GIT : CI/CD Pipeline
@enduml"""

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCION DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── PORTADA ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('OLYMPUS CORE')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Sistema Gestor de Gimnasio')
r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x44, 0x4A, 0x57)

doc.add_paragraph()
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('ESPECIFICACION DE REQUISITOS DE CALIDAD')
r3.bold = True; r3.font.size = Pt(20); r3.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.add_paragraph()
for line in ['Entregable: Iteracion 3 — Requisitos de Calidad', 'Grupo: 5', 'Julio 2026']:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run(line).font.size = Pt(12)

doc.add_page_break()

# ── 1. INTRODUCCION ──────────────────────────────────────────────────────────
doc.add_heading('1. Introduccion', 1)
doc.add_paragraph(
    'El presente documento especifica los Requisitos de Calidad del sistema OLYMPUS CORE '
    'siguiendo el estandar ISO/IEC 25010:2011 (Systems and software Quality Requirements '
    'and Evaluation — SQuaRE). Se identifican 13 Requisitos de Calidad (RQ-CAL) derivados '
    'directamente de los 13 Requisitos No Funcionales del sistema y de los hallazgos de '
    'las cuatro tecnicas de elicitacion aplicadas.\n\n'
    'Cada requisito se especifica con su actor de origen, clasificacion ISO 25010, metrica '
    'de medicion, metodo de verificacion y criterio de aceptacion. El documento incluye '
    'matrices de trazabilidad hacia el proceso de elicitacion y hacia los Requisitos '
    'Funcionales asociados, ademas de diagramas UML de distintos tipos para ilustrar el '
    'comportamiento del sistema en cada dimension de calidad.'
)

# ── 2. MARCO ISO 25010 ────────────────────────────────────────────────────────
doc.add_heading('2. Marco de Referencia: ISO/IEC 25010', 1)
doc.add_paragraph(
    'Los requisitos de calidad del sistema se clasifican segun las caracteristicas y '
    'subcaracteristicas del modelo de calidad de producto ISO/IEC 25010:'
)
iso_t = doc.add_table(rows=1, cols=3)
iso_t.style = 'Table Grid'
header_row(iso_t, ['Caracteristica ISO 25010', 'Subcaracteristica', 'RQ-CAL Asociados'])
iso_data = [
    ('Eficiencia de Desempeno', 'Comportamiento Temporal\nUtilizacion de Recursos\nCapacidad', 'RQ-CAL-OLY-01, RQ-CAL-OLY-02, RQ-CAL-OLY-03, RQ-CAL-OLY-13'),
    ('Seguridad', 'Confidencialidad\nIntegridad\nAutenticidad\nResistencia a Ataques', 'RQ-CAL-OLY-04, RQ-CAL-OLY-05, RQ-CAL-OLY-06, RQ-CAL-OLY-07'),
    ('Usabilidad', 'Accesibilidad\nEficiencia de Uso', 'RQ-CAL-OLY-08, RQ-CAL-OLY-09, RQ-CAL-OLY-10'),
    ('Fiabilidad', 'Disponibilidad', 'RQ-CAL-OLY-11'),
    ('Mantenibilidad', 'Capacidad de Modificacion', 'RQ-CAL-OLY-12'),
]
for row in iso_data:
    r = iso_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
doc.add_paragraph()

# ── 3. CATALOGO DE REQUISITOS DE CALIDAD ──────────────────────────────────────
doc.add_heading('3. Catalogo de Requisitos de Calidad', 1)
doc.add_paragraph(
    'A continuacion se presenta la especificacion formal de cada Requisito de Calidad '
    'utilizando la plantilla estandar del proyecto.'
)

# Mapa de tipos de diagrama a UML code (para variar)
diagrama_map = {}
for rq in RQ_DATA:
    rid = rq['id']
    tipo = rq['diagrama_tipo']
    if tipo == 'clase':
        diagrama_map[rid] = (uml_clase_rq(rq), f"Figura: Diagrama de Clases — {rid}", 6.2)
    elif tipo == 'secuencia':
        if rid == 'RQ-CAL-OLY-02':
            diagrama_map[rid] = (UML_SECUENCIA_QR, f"Figura: Diagrama de Secuencia — Flujo de Validacion QR ({rid})", 6.0)
        else:
            diagrama_map[rid] = (UML_SECUENCIA_JWT, f"Figura: Diagrama de Secuencia — Ciclo de Vida JWT ({rid})", 6.0)
    elif tipo == 'estado':
        if rid == 'RQ-CAL-OLY-03':
            diagrama_map[rid] = (UML_ESTADO_DISPONIBILIDAD, f"Figura: Diagrama de Estado — Disponibilidad del Sistema ({rid})", 5.5)
        else:
            diagrama_map[rid] = (UML_ESTADO_MEMBRESIA, f"Figura: Diagrama de Estado — Ciclo de Vida Membresia ({rid})", 5.5)
    elif tipo == 'actividad':
        if rid == 'RQ-CAL-OLY-05':
            diagrama_map[rid] = (UML_ACTIVIDAD_HTTPS, f"Figura: Diagrama de Actividad — Flujo HTTPS ({rid})", 4.5)
        else:
            diagrama_map[rid] = (UML_ACTIVIDAD_RESERVA, f"Figura: Diagrama de Actividad — Flujo de Reserva ({rid})", 4.5)
    elif tipo == 'componente':
        if rid == 'RQ-CAL-OLY-08':
            diagrama_map[rid] = (UML_COMPATIBILIDAD, f"Figura: Diagrama de Despliegue — Compatibilidad de Navegadores ({rid})", 5.5)
        elif rid == 'RQ-CAL-OLY-09':
            diagrama_map[rid] = (UML_RESPONSIVE, f"Figura: Diagrama de Estados — Diseño Responsive ({rid})", 5.0)
        else:
            diagrama_map[rid] = (UML_COMPONENTE, f"Figura: Diagrama de Componentes — Arquitectura de Calidad ({rid})", 6.2)

PRIORIDAD_BG = {'Alta': 'FFD7D7', 'Media': 'FFF3CD', 'Baja': 'D4EDDA'}

for rq in RQ_DATA:
    doc.add_heading(f"{rq['id']} — {rq['nombre']}", 2)

    tbl = doc.add_table(rows=11, cols=2)
    tbl.style = 'Table Grid'
    campos = [
        ('Identificador',          rq['id']),
        ('Nombre del RQ',          rq['nombre']),
        ('Actor / Stakeholder',    rq['stakeholder']),
        ('Necesidad del negocio',  rq['necesidad']),
        ('Descripcion formal',     rq['descripcion']),
        ('Prioridad',              rq['prioridad']),
        ('Criterio de Aceptacion', rq['criterio']),
        ('Caracteristica ISO 25010', f"{rq['iso_caracteristica']} > {rq['iso_subcaracteristica']}"),
        ('Metrica',                f"Nombre: {rq['metrica_nombre']}\nUnidad: {rq['metrica_unidad']}\nValor objetivo: {rq['metrica_objetivo']}"),
        ('Metodo de Verificacion', f"{rq['verificacion_metodo']}\nEvidencia: {rq['verificacion_evidencia']}"),
        ('RF Asociados',           rq['rfs_asociados']),
    ]
    for i, (key, val) in enumerate(campos):
        cell_key = tbl.rows[i].cells[0]
        cell_val = tbl.rows[i].cells[1]
        cell_key.text = key
        cell_key.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell_key, 'D0D8E4')
        cell_val.text = val
        if key == 'Prioridad':
            bg = PRIORIDAD_BG.get(val, 'FFFFFF')
            set_cell_bg(cell_val, bg)
            cell_val.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Diagrama UML correspondiente
    if rq['id'] in diagrama_map:
        uml_code, caption, width = diagrama_map[rq['id']]
        add_diagram(doc, uml_code, caption, width)

    doc.add_paragraph()

# ── 4. MATRICES DE TRAZABILIDAD ───────────────────────────────────────────────
doc.add_heading('4. Matrices de Trazabilidad', 1)

# 4.1 RQ → Proceso de Elicitacion
doc.add_heading('4.1 Trazabilidad: Proceso de Elicitacion → Requisito de Calidad', 2)
doc.add_paragraph(
    'La siguiente matriz identifica de que tecnica de elicitacion se derivo cada Requisito '
    'de Calidad, asegurando que todos los RQ tienen respaldo en el proceso de levantamiento '
    'de requisitos documentado en el documento de Descubrimiento y Elicitacion.'
)
mt1 = doc.add_table(rows=1, cols=4)
mt1.style = 'Table Grid'
header_row(mt1, ['ID RQ-CAL', 'Nombre', 'Tecnica de Elicitacion', 'Hallazgo que la origino'])
traza_elicit = [
    ('RQ-CAL-OLY-01', 'Latencia QR',         'Observacion (as-is)',        'El tiempo de verificacion manual era 60-90 seg; el QR debe hacerlo en < 300 ms'),
    ('RQ-CAL-OLY-02', 'Latencia API',         'Benchmarking',              'GymMaster y PerfectGym prometen respuestas < 500 ms en sus SLAs publicos'),
    ('RQ-CAL-OLY-03', 'Concurrencia',         'Observacion (as-is)',        'Se observaron colas en hora pico (7AM y 6PM) con uso manual'),
    ('RQ-CAL-OLY-04', 'Bcrypt passwords',     'Brainstorming',             'El equipo identifico el riesgo de brecha de BD durante la sesion de brainstorming'),
    ('RQ-CAL-OLY-05', 'HTTPS obligatorio',    'Brainstorming + Benchmark', 'Todos los sistemas analizados (GymMaster, Mindbody) usan HTTPS exclusivo'),
    ('RQ-CAL-OLY-06', 'Expiracion tokens',    'Benchmarking',              'Mindbody usa QR dinamico; se adopto expiracion de token como alternativa factible'),
    ('RQ-CAL-OLY-07', 'Anti SQL/XSS',         'Brainstorming',             'El equipo menciono OWASP Top 10 como estandar minimo de seguridad en sesion 1'),
    ('RQ-CAL-OLY-08', 'Compat. navegadores',  'Entrevista',                'Cliente: "quiero que funcione en cualquier computadora del gym sin instalar nada"'),
    ('RQ-CAL-OLY-09', 'Responsive',           'Benchmarking',              'Todos los sistemas analizados soportan tablets; el cliente tiene tablets en recepcion'),
    ('RQ-CAL-OLY-10', 'Max 3 clics',          'Observacion (as-is)',        'Se midio que el recepcionista tarda mas por procesos manuales complejos en hora pico'),
    ('RQ-CAL-OLY-11', 'Disponibilidad 99%%',  'Entrevista',                'Cliente: "si el sistema cae en la manana del lunes es un desastre para nosotros"'),
    ('RQ-CAL-OLY-12', 'Control de versiones', 'Brainstorming',             'El equipo decidio en sesion 2 seguir Git Flow para facilitar el trabajo en paralelo'),
    ('RQ-CAL-OLY-13', 'Indices en BD',        'Benchmarking + Observacion','Se anticipo degradacion de busqueda por DNI a mayor volumen de socios'),
]
for row in traza_elicit:
    r = mt1.add_row()
    for c, t in zip(r.cells, row): c.text = t

doc.add_paragraph()

# 4.2 RQ → RF Asociados
doc.add_heading('4.2 Trazabilidad: Requisito de Calidad → Requisitos Funcionales', 2)
doc.add_paragraph(
    'La siguiente matriz vincula cada Requisito de Calidad con los Requisitos Funcionales '
    'cuya implementacion debe satisfacer dicho requisito de calidad. Esta trazabilidad '
    'garantiza cobertura total y permite identificar el impacto de cambios en los RF.'
)
mt2 = doc.add_table(rows=1, cols=5)
mt2.style = 'Table Grid'
header_row(mt2, ['ID RQ-CAL', 'Nombre', 'RNF Origen', 'RF Asociados', 'Impacto si no se cumple'])
traza_rf = [
    ('RQ-CAL-OLY-01', 'Latencia QR',         'RNF-02', 'RF-06, RF-07',                       'Colas en recepcion; socios insatisfechos; imagen negativa'),
    ('RQ-CAL-OLY-02', 'Latencia API',         'RNF-01', 'RF-01 a RF-08, RF-15',              'Experiencia lenta; rechazo del sistema por parte del personal'),
    ('RQ-CAL-OLY-03', 'Concurrencia',         'RNF-03', 'RF-07, RF-09, RF-15',              'Caidas en hora pico; perdida de confianza en el sistema'),
    ('RQ-CAL-OLY-04', 'Bcrypt passwords',     'RNF-04', 'RF-01, RF-10',                     'Exposicion de credenciales ante brecha de seguridad'),
    ('RQ-CAL-OLY-05', 'HTTPS obligatorio',    'RNF-05', 'RF-01, RF-06, RF-07',              'Intercepcion de datos sensibles (credentials, QR tokens)'),
    ('RQ-CAL-OLY-06', 'Expiracion tokens',    'RNF-06', 'RF-01, RF-06, RF-07, RF-10',      'Accesos no autorizados con tokens robados o compartidos'),
    ('RQ-CAL-OLY-07', 'Anti SQL/XSS',         'RNF-07', 'RF-02, RF-05, RF-08, RF-12',      'Corrupcion de datos o ejecucion de codigo malicioso en BD'),
    ('RQ-CAL-OLY-08', 'Compat. navegadores',  'RNF-08', 'RF-01, RF-04, RF-07, RF-15, RF-16', 'Partes del sistema inaccesibles desde ciertos navegadores'),
    ('RQ-CAL-OLY-09', 'Responsive',           'RNF-09', 'RF-15, RF-16, RF-07',             'UI inutilizable en tablets de recepcion'),
    ('RQ-CAL-OLY-10', 'Max 3 clics',          'RNF-10', 'RF-05, RF-07',                    'Ralentizacion del flujo de recepcion; errores del operador'),
    ('RQ-CAL-OLY-11', 'Disponibilidad 99%%',  'RNF-11', 'RF-01, RF-07, RF-09, RF-15',     'Gimnasio sin sistema durante horario de atencion'),
    ('RQ-CAL-OLY-12', 'Control versiones',    'RNF-12', 'Todos los RF',                   'Dificultad para revertir bugs; perdida de historial de cambios'),
    ('RQ-CAL-OLY-13', 'Indices en BD',        'RNF-13', 'RF-02, RF-04, RF-07, RF-08, RF-12', 'Consultas lentas a mayor volumen de datos; degrada RQ-01 y 02'),
]
for row in traza_rf:
    r = mt2.add_row()
    for c, t in zip(r.cells, row): c.text = t

doc.add_paragraph()

# 4.3 Resumen de cobertura
doc.add_heading('4.3 Resumen de Cobertura de Trazabilidad', 2)
doc.add_paragraph(
    'La siguiente tabla resume el porcentaje de RF del sistema que tienen al menos un '
    'Requisito de Calidad asociado, garantizando cobertura del modelo de calidad.'
)
cob_t = doc.add_table(rows=1, cols=4)
cob_t.style = 'Table Grid'
header_row(cob_t, ['Categoria', 'Total RF', 'RF con RQ-CAL asociado', 'Cobertura'])
cob_data = [
    ('Must Have',   '11', '11', '100%'),
    ('Should Have', '8',  '8',  '100%'),
    ('Could Have',  '8',  '3',  '37.5% (los implementados en MVP1)'),
    ("Won't Have",  '7',  '0',  '— (fuera de alcance)'),
    ('TOTAL MVP1',  '19', '19', '100%'),
]
for row in cob_data:
    r = cob_t.add_row()
    for c, t in zip(r.cells, row): c.text = t
    if row[0] == 'TOTAL MVP1':
        for c in r.cells:
            c.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# ── 5. DIAGRAMAS UML DE CALIDAD DEL SISTEMA ──────────────────────────────────
doc.add_heading('5. Diagramas UML de Calidad del Sistema', 1)
doc.add_paragraph(
    'Esta seccion consolida los principales diagramas UML que ilustran el comportamiento '
    'del sistema en sus dimensiones de calidad. Cada diagrama incluye anotaciones que '
    'referencian los RQ-CAL y RF asociados.'
)

# 5.1 Diagrama de Clases - Modelo de Calidad Completo
doc.add_heading('5.1 Diagrama de Clases — Modelo de Calidad del Sistema (ISO 25010)', 2)
doc.add_paragraph(
    'Vista general del modelo de calidad aplicado a OLYMPUS CORE: muestra como los actores, '
    'las caracteristicas ISO 25010 y los RQ-CAL se relacionan entre si. '
    '[Asociado a todos los RQ-CAL | RF-01 a RF-19]'
)
UML_CLASE_MODELO = """@startuml
hide circle
skinparam classAttributeIconSize 0
skinparam class {
  BackgroundColor #F0F4F8
  BorderColor #1E3A5F
  ArrowColor #1E3A5F
}

abstract class "CaracteristicaISO25010" {
  nombre : String
}

class "EficienciaDesempeno" extends "CaracteristicaISO25010" {
  comportamientoTemporal
  utilizacionRecursos
  capacidad
  -- RQ-CAL: OLY-01, 02, 03, 13 --
}

class "Seguridad" extends "CaracteristicaISO25010" {
  confidencialidad
  integridad
  autenticidad
  resistenciaAuataques
  -- RQ-CAL: OLY-04, 05, 06, 07 --
}

class "Usabilidad" extends "CaracteristicaISO25010" {
  accesibilidad
  eficienciaDeUso
  -- RQ-CAL: OLY-08, 09, 10 --
}

class "Fiabilidad" extends "CaracteristicaISO25010" {
  disponibilidad
  -- RQ-CAL: OLY-11 --
}

class "Mantenibilidad" extends "CaracteristicaISO25010" {
  capacidadDeModificacion
  -- RQ-CAL: OLY-12 --
}

class "RequisitoCalidad" {
  id : String
  descripcion : String
  prioridad : Enum
  criterioAceptacion : String
  rnfAsociado : String
  rfsAsociados : List<String>
}

class "Metrica" {
  nombre : String
  unidad : String
  valorObjetivo : String
}

class "Verificacion" {
  metodo : String
  evidencia : String
}

class "Sistema OLYMPUS CORE" {
  nombre = "OLYMPUS CORE"
  version = "MVP1"
  stack = "React + Node + PostgreSQL"
}

"CaracteristicaISO25010" "1" <-- "1..*" RequisitoCalidad : clasifica
RequisitoCalidad "1" --> "1..*" Metrica : se mide con
RequisitoCalidad "1" --> "1" Verificacion : se valida mediante
"Sistema OLYMPUS CORE" "1" --> "13" RequisitoCalidad : debe satisfacer
@enduml"""
add_diagram(doc, UML_CLASE_MODELO, "Figura 5.1: Diagrama de Clases — Modelo de Calidad del Sistema (ISO 25010)", 6.2)

# 5.2 Diagrama de Secuencia - Flujo QR completo
doc.add_heading('5.2 Diagrama de Secuencia — Validacion de Acceso QR', 2)
doc.add_paragraph(
    'Especifica el flujo de mensajes exacto entre actores y componentes durante la '
    'validacion de QR, con anotaciones de tiempo para verificar el cumplimiento de '
    '[RQ-CAL-OLY-01 (<= 300ms) | RF-07]'
)
add_diagram(doc, UML_SECUENCIA_QR, "Figura 5.2: Diagrama de Secuencia — Validacion de QR (RQ-CAL-OLY-01, RF-07)", 6.0)

# 5.3 Diagrama de Estado - Disponibilidad
doc.add_heading('5.3 Diagrama de Estado — Disponibilidad del Sistema', 2)
doc.add_paragraph(
    'Muestra los estados de disponibilidad del sistema y las transiciones entre ellos, '
    'formalizando el requisito de 99%% de uptime en horario operativo. '
    '[RQ-CAL-OLY-11 | RF-07, RF-15 | RNF-11]'
)
add_diagram(doc, UML_ESTADO_DISPONIBILIDAD, "Figura 5.3: Diagrama de Estado — Disponibilidad del Sistema (RQ-CAL-OLY-11)", 5.5)

# 5.4 Diagrama de Estado - Membresia
doc.add_heading('5.4 Diagrama de Estado — Ciclo de Vida de la Membresia', 2)
doc.add_paragraph(
    'Formaliza los estados validos de una membresia y los eventos que producen transiciones. '
    'Este diagrama impacta directamente en la validacion de acceso QR. '
    '[RF-04, RF-05, RF-14, RF-17 | RQ-CAL-OLY-01, RQ-CAL-OLY-03]'
)
add_diagram(doc, UML_ESTADO_MEMBRESIA, "Figura 5.4: Diagrama de Estado — Ciclo de Vida de la Membresia (RF-04, RF-14)", 5.5)

# 5.5 Diagrama de Actividad - HTTPS
doc.add_heading('5.5 Diagrama de Actividad — Flujo de Comunicacion Segura HTTPS', 2)
doc.add_paragraph(
    'Detalla el proceso de establecimiento de la conexion segura entre cliente y servidor, '
    'cubriendo el requisito de trafico 100%% cifrado. '
    '[RQ-CAL-OLY-05 | RF-01, RF-06, RF-07 | RNF-05]'
)
add_diagram(doc, UML_ACTIVIDAD_HTTPS, "Figura 5.5: Diagrama de Actividad — Comunicacion Segura HTTPS (RQ-CAL-OLY-05)", 4.5)

# 5.6 Diagrama de Actividad - Reserva
doc.add_heading('5.6 Diagrama de Actividad — Flujo de Reserva de Clase', 2)
doc.add_paragraph(
    'Muestra las validaciones paralelas que se ejecutan al realizar una reserva, '
    'ilustrando como multiples RQ-CAL afectan a un mismo flujo de RF. '
    '[RF-09 | RQ-CAL-OLY-02 (latencia), RQ-CAL-OLY-03 (concurrencia)]'
)
add_diagram(doc, UML_ACTIVIDAD_RESERVA, "Figura 5.6: Diagrama de Actividad — Reserva de Clase (RF-09)", 4.5)

# 5.7 Diagrama de Secuencia - JWT
doc.add_heading('5.7 Diagrama de Secuencia — Ciclo de Vida del Token JWT', 2)
doc.add_paragraph(
    'Especifica como el sistema gestiona la autenticacion basada en tokens y aplica '
    'la politica de expiracion para minimizar el riesgo de acceso no autorizado. '
    '[RQ-CAL-OLY-06 | RF-01, RF-10 | RNF-06]'
)
add_diagram(doc, UML_SECUENCIA_JWT, "Figura 5.7: Diagrama de Secuencia — Ciclo de Vida JWT (RQ-CAL-OLY-06)", 6.0)

# 5.8 Diagrama de Componentes
doc.add_heading('5.8 Diagrama de Componentes — Arquitectura de Calidad del Sistema', 2)
doc.add_paragraph(
    'Muestra como las responsabilidades de calidad se distribuyen entre los componentes '
    'de la arquitectura (frontend, backend, BD, repositorio). Cada componente se anota '
    'con los RQ-CAL que debe satisfacer. '
    '[RQ-CAL-OLY-05, 08, 09, 12 | Arquitectura general del MVP1]'
)
add_diagram(doc, UML_COMPONENTE, "Figura 5.8: Diagrama de Componentes — Distribucion de Calidad por Componente", 6.2)

# ── 6. GLOSARIO ──────────────────────────────────────────────────────────────
doc.add_heading('6. Glosario', 1)
glosario = [
    ('RQ-CAL', 'Requisito de Calidad. Especifica un atributo de calidad del sistema medible y verificable.'),
    ('ISO/IEC 25010', 'Estandar internacional que define el modelo de calidad del producto software con 8 caracteristicas y 31 subcaracteristicas.'),
    ('SLA', 'Service Level Agreement. Acuerdo de nivel de servicio que define los umbrales de calidad comprometidos.'),
    ('Percentil 95 (P95)', 'El 95%% de las mediciones registradas se encuentran por debajo de este valor. Se usa para evaluar latencia excluyendo casos extremos.'),
    ('VU', 'Usuario Virtual. Simulacion de un usuario real en una prueba de carga.'),
    ('bcrypt', 'Algoritmo de hashing de contrasenas con factor de costo configurable para resistir ataques de fuerza bruta.'),
    ('TLS', 'Transport Layer Security. Protocolo de cifrado que asegura la comunicacion en red sobre HTTPS.'),
    ('JWT', 'JSON Web Token. Token de autenticacion firmado digitalmente que contiene la identidad y rol del usuario.'),
    ('RBAC', 'Role-Based Access Control. Control de acceso basado en roles.'),
    ('Index Scan', 'Tipo de consulta en PostgreSQL que utiliza un indice para acceder a datos sin recorrer toda la tabla (mas eficiente que Seq Scan).'),
    ('OWASP Top 10', 'Lista de las diez vulnerabilidades de seguridad web mas criticas segun el Open Web Application Security Project.'),
    ('UptimeRobot', 'Herramienta de monitoreo de disponibilidad que verifica periodicamente si un servicio esta en linea.'),
    ('k6', 'Herramienta de pruebas de carga de codigo abierto para APIs y microservicios.'),
]
g_t = doc.add_table(rows=1, cols=2)
g_t.style = 'Table Grid'
header_row(g_t, ['Termino', 'Definicion'])
for term, defin in glosario:
    r = g_t.add_row()
    r.cells[0].text = term
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[1].text = defin

doc.add_paragraph()

# ── GUARDAR ───────────────────────────────────────────────────────────────────
output_path = r'c:/Users/ACER/OneDrive/Desktop/PROGRAMACION/UNIVERSITY PROJECTS/SIstemaGestorGimnasio/Documentacion/Especificacion_Requisitos_Calidad_actualizado.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
